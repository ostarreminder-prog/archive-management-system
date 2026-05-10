import os
import json
import io
import re
import secrets
import tempfile
import hashlib
import html
import base64
import mimetypes
import zipfile
import struct
import zlib
import posixpath
import threading
from urllib.parse import quote
import xml.etree.ElementTree as ET
from copy import deepcopy
from datetime import datetime
from flask import (Flask, render_template, request, jsonify,
                   session, redirect, url_for, send_file, after_this_request)
from dotenv import load_dotenv
from werkzeug.utils import secure_filename
load_dotenv()

try:
    from docx import Document as DocxDocument
    from docx.shared import Inches
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except Exception:
    DocxDocument = None
    Inches = None
    OxmlElement = None
    qn = None
    DOCX_AVAILABLE = False

from modules.database import init_db, get_db, log_action, log_login
from modules.auth import (
    hash_password, check_password, generate_otp, verify_otp,
    generate_serial, generate_temp_password, verify_sign_password,
    generate_sign_hash, login_required, role_required,
    create_user, get_user_by_email, get_user_by_id, get_all_users,
    get_all_archive_sections, create_archive_section,
    update_password, update_sign_password, update_signature, can_use_stamp,
    get_device_hash, is_trusted_device, register_pending_device,
    trust_device_by_token, block_user, get_trusted_devices
)
from modules.email import (
    send_welcome, send_otp, send_sign_used, send_sign_request,
    send_approved, send_rejected, send_weekly_password, send_new_device
)
from modules.storage import storage

app = Flask(__name__)
app.secret_key = os.getenv("SECRET_KEY", "najm-secret-change-in-production")
app.config['MAX_CONTENT_LENGTH'] = 10 * 1024 * 1024

UPLOAD_FOLDER = os.path.join(os.path.dirname(__file__), 'static', 'uploads')
os.makedirs(os.path.join(UPLOAD_FOLDER, 'signatures'), exist_ok=True)
os.makedirs(os.path.join(UPLOAD_FOLDER, 'templates'), exist_ok=True)
os.makedirs(os.path.join(UPLOAD_FOLDER, 'stamps'), exist_ok=True)

LOCALHOST_BASE_URL = os.getenv("LOCALHOST_BASE_URL", os.getenv("VERIFY_BASE_URL", "http://localhost:5000")).rstrip("/")

def _get_verify_base_url():
    """Returns appropriate base URL for QR verification links — uses current request host if available (supports mobile)."""
    try:
        from flask import has_request_context, request as _req
        if has_request_context():
            host_url = _req.host_url.rstrip('/')
            # Ignore localhost only if we have a real alternative
            if host_url and '127.0.0.1' not in host_url and 'localhost' not in host_url:
                return host_url
    except Exception as e:
        print(f"[DEBUG] has_request_context check failed: {e}", file=__import__('sys').stderr)
    
    # Fallback to env variables or default
    result = os.getenv("VERIFY_BASE_URL", os.getenv("SERVER_BASE_URL", LOCALHOST_BASE_URL)).rstrip("/")
    if not result or result == "":
        result = "http://localhost:5000"
    return result


def _build_qr_payload(serial_text='', archive_number=''):
    """Build QR payload URL — returns verify URL with proper serial/archive value, or None."""
    try:
        serial_value = str(serial_text or '').strip().upper()
        if serial_value:
            base_url = _get_verify_base_url()
            if not base_url:
                print(f"[ERROR] _get_verify_base_url() returned empty/None", file=__import__('sys').stderr)
                # Fallback to environment variable or default
                base_url = os.getenv("VERIFY_BASE_URL", os.getenv("SERVER_BASE_URL", "http://localhost:5000")).rstrip("/")
            
            qr_url = f"{base_url}/verify/{quote(serial_value)}"
            print(f"[DEBUG] Generated QR URL for serial {serial_value}: {qr_url}", file=__import__('sys').stderr)
            return qr_url

        # Try archive_number as fallback if serial doesn't exist
        archive_value = str(archive_number or '').strip().upper()
        if archive_value:
            base_url = _get_verify_base_url()
            if not base_url:
                base_url = os.getenv("VERIFY_BASE_URL", os.getenv("SERVER_BASE_URL", "http://localhost:5000")).rstrip("/")
            
            qr_url = f"{base_url}/verify/{quote(archive_value)}"
            print(f"[DEBUG] Generated QR URL for archive {archive_value}: {qr_url}", file=__import__('sys').stderr)
            return qr_url

        print(f"[WARNING] No serial or archive value provided to _build_qr_payload()", file=__import__('sys').stderr)
        return None
    except Exception as e:
        print(f"[ERROR] Exception in _build_qr_payload(): {str(e)}", file=__import__('sys').stderr)
        import traceback
        traceback.print_exc(file=__import__('sys').stderr)
        return None


def _png_response(blob, status=200):
    return app.response_class(blob, status=status, mimetype="image/png")
DEFAULT_ADMIN_EMAIL = os.getenv("DEFAULT_ADMIN_EMAIL", os.getenv("SMTP_USER", "admin@najm.sa")).strip().lower()
ARCHIVE_REMOTE_PREFIX = os.getenv("ARCHIVE_REMOTE_PREFIX", "documents").strip().strip("/") or "documents"
PDF_RENDER_LOCK = threading.Lock()
DEFAULT_DOCX_TEMPLATE_SETTING_KEY = "default_docx_template"


def _safe_int(value, default=0):
    try:
        return int(value)
    except (TypeError, ValueError):
        return default


def _normalize_relative_path(path=""):
    raw = str(path or "").replace("\\", "/").strip().strip("/")
    if not raw:
        return ""
    safe_parts = []
    for part in raw.split("/"):
        clean = part.strip()
        if not clean or clean == ".":
            continue
        if clean == "..":
            if safe_parts:
                safe_parts.pop()
            continue
        safe_parts.append(clean)
    return "/".join(safe_parts)


def _resolve_document_file_path(file_path):
    if not file_path:
        return None

    normalized = str(file_path).strip().replace("\\", "/")
    if not normalized:
        return None

    candidates = []
    if os.path.isabs(normalized):
        candidates.append(normalized)

    candidates.append(os.path.join(app.root_path, normalized))

    if normalized.startswith("uploads/"):
        candidates.append(os.path.join(app.root_path, "static", normalized))
    if normalized.startswith("static/"):
        candidates.append(os.path.join(app.root_path, normalized))

    for candidate in candidates:
        if os.path.exists(candidate):
            return candidate

    if not normalized.startswith(('uploads/', 'static/')):
        cache_root = os.path.join(app.root_path, 'static', 'uploads', 'archive_storage_cache')
        cache_path = os.path.join(cache_root, *[part for part in normalized.split('/') if part])
        if os.path.exists(cache_path):
            return cache_path

        os.makedirs(os.path.dirname(cache_path), exist_ok=True)
        downloaded = storage.download(normalized, cache_path)
        if downloaded and os.path.exists(cache_path):
            return cache_path

    return None


def _sha256_hex_from_bytes(blob):
    if blob is None:
        return None
    hasher = hashlib.sha256()
    hasher.update(blob)
    return hasher.hexdigest()


def _sha256_hex_from_file(path):
    if not path or not os.path.exists(path):
        return None

    hasher = hashlib.sha256()
    with open(path, 'rb') as handle:
        while True:
            chunk = handle.read(1024 * 1024)
            if not chunk:
                break
            hasher.update(chunk)
    return hasher.hexdigest()


def _build_text_document_fingerprint(content_html='', template_name=''):
    normalized_payload = json.dumps(
        {
            'mode': 'text',
            'template_name': str(template_name or '').strip(),
            'content_html': str(content_html or '').strip(),
        },
        ensure_ascii=False,
        sort_keys=True,
    )
    return hashlib.sha256(normalized_payload.encode('utf-8')).hexdigest()


def _compute_document_fingerprint(doc):
    if not doc:
        return None

    existing = str(doc.get('file_sha256') or '').strip().lower()
    if existing:
        return existing

    source_path = _resolve_document_file_path(doc.get('file_path'))
    if not source_path:
        source_path = _resolve_document_file_path(doc.get('archive_storage_path'))
    if source_path:
        try:
            return _sha256_hex_from_file(source_path)
        except Exception:
            return None

    content_html = _extract_content_html(doc.get('content_json') or '')
    if content_html:
        return _build_text_document_fingerprint(
            content_html=content_html,
            template_name=doc.get('template_name') or '',
        )

    return None


def _find_approved_duplicate_document(conn, fingerprint):
    normalized = str(fingerprint or '').strip().lower()
    if not normalized:
        return None

    exact = conn.execute(
        """
        SELECT id, title, archive_number, status, file_sha256
        FROM documents
        WHERE status='approved' AND LOWER(COALESCE(file_sha256, ''))=?
        ORDER BY id DESC
        LIMIT 1
        """,
        (normalized,)
    ).fetchone()
    if exact:
        return dict(exact)

    legacy_rows = conn.execute(
        """
        SELECT id, title, archive_number, status, file_path, archive_storage_path,
               content_json, template_name, file_sha256
        FROM documents
        WHERE status='approved' AND TRIM(COALESCE(file_sha256, ''))=''
        ORDER BY id DESC
        """
    ).fetchall()
    for row in legacy_rows:
        doc = dict(row)
        doc_fingerprint = _compute_document_fingerprint(doc)
        if not doc_fingerprint:
            continue
        try:
            conn.execute(
                "UPDATE documents SET file_sha256=? WHERE id=? AND TRIM(COALESCE(file_sha256, ''))=''",
                (doc_fingerprint, doc['id'])
            )
        except Exception:
            pass
        if doc_fingerprint == normalized:
            return doc

    return None


def _persist_document_fingerprint(conn, doc_id):
    if not conn or not doc_id:
        return None

    row = conn.execute("SELECT * FROM documents WHERE id=? LIMIT 1", (doc_id,)).fetchone()
    if not row:
        return None

    fingerprint = _compute_document_fingerprint(dict(row))
    if fingerprint and fingerprint != str(row['file_sha256'] or '').strip().lower():
        conn.execute("UPDATE documents SET file_sha256=? WHERE id=?", (fingerprint, doc_id))
    return fingerprint


def _extract_docx_text_lines(xml_bytes, limit=120):
    if not xml_bytes:
        return []

    try:
        root = ET.fromstring(xml_bytes)
    except ET.ParseError:
        return []

    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    lines = []
    seen = set()
    for paragraph in root.findall('.//w:p', ns):
        chunks = []
        for node in paragraph.findall('.//w:t', ns):
            text = (node.text or '').strip()
            if text:
                chunks.append(text)
        if not chunks:
            continue

        line = ' '.join(chunks).strip()
        if not line:
            continue

        if line in seen:
            continue
        seen.add(line)

        lines.append(line)
        if len(lines) >= limit:
            break

    return lines


def _crop_header_img(blob, mime):
    """اقتصاص الصورة لإظهار هيدر الكليشة فقط (أعلى 22%)"""
    try:
        import io as _cio
        from PIL import Image as _CPIL
        _img = _CPIL.open(_cio.BytesIO(blob))
        _iw, _ih = _img.size
        if _ih > 300:
            _ch = max(150, int(_ih * 0.22))
            _cr = _img.crop((0, 0, _iw, _ch))
            _buf = _cio.BytesIO()
            _fmt = 'JPEG' if 'jpeg' in mime else 'PNG'
            _cr.save(_buf, format=_fmt, quality=92)
            return _buf.getvalue()
    except Exception:
        pass
    return blob

def _build_docx_template_preview(file_path):
    import zipfile, base64, mimetypes, posixpath
    import xml.etree.ElementTree as ET
 
    header_images = []
    header_lines  = []
    body_lines    = []
    seen_images   = set()
 
    with zipfile.ZipFile(file_path, 'r') as archive:
        archive_names = set(archive.namelist())
 
        # ─── جمع كل صور media أولاً (fallback شامل) ──────────────
        all_media = sorted(n for n in archive_names if n.startswith('word/media/'))
 
        # ─── معالجة الهيدر ─────────────────────────────────────────
        header_files = sorted(n for n in archive_names
                               if n.startswith('word/header') and n.endswith('.xml'))
 
        for header_name in header_files:
            try:
                header_xml = archive.read(header_name)
                root = ET.fromstring(header_xml)
                ns   = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                for p in root.findall('.//w:p', ns):
                    text = ''.join(t.text or '' for t in p.findall('.//w:t', ns)).strip()
                    if text and len(text) > 1 and text not in header_lines:
                        header_lines.append(text)
            except Exception:
                pass
 
            # جرب rels أولاً
            rels_name = f"word/_rels/{posixpath.basename(header_name)}.rels"
            found_via_rels = False
 
            if rels_name in archive_names:
                try:
                    rels_root = ET.fromstring(archive.read(rels_name))
                    ns_r = 'http://schemas.openxmlformats.org/package/2006/relationships'
                    for rel in rels_root.findall(f'{{{ns_r}}}Relationship'):
                        if not (rel.attrib.get('Type') or '').lower().endswith('/image'):
                            continue
                        target = (rel.attrib.get('Target') or '').replace('\\', '/')
                        if target.startswith('/'):
                            img_path = target.lstrip('/')
                        else:
                            img_path = posixpath.normpath(
                                posixpath.join(posixpath.dirname(header_name), target)
                            )
                        if img_path in seen_images or img_path not in archive_names:
                            continue
                        blob = archive.read(img_path)
                        if blob and len(blob) <= 3 * 1024 * 1024:
                            mime = mimetypes.guess_type(img_path)[0] or 'image/jpeg'
                            header_images.append(
                                f"data:{mime};base64,{base64.b64encode(_crop_header_img(blob, mime)).decode('ascii')}"
                            )
                            seen_images.add(img_path)
                            found_via_rels = True
                        if len(header_images) >= 2:
                            break
                except Exception:
                    pass
 
            # لو ما جاب عبر rels — استخدم الصور من media مباشرة
            if not found_via_rels and all_media:
                for img_path in all_media:
                    if img_path in seen_images:
                        continue
                    try:
                        blob = archive.read(img_path)
                    except Exception:
                        continue
                    if not blob or len(blob) > 3 * 1024 * 1024:
                        continue
                    mime = mimetypes.guess_type(img_path)[0] or 'image/jpeg'
                    header_images.append(
                        f"data:{mime};base64,{base64.b64encode(_crop_header_img(blob, mime)).decode('ascii')}"
                    )
                    seen_images.add(img_path)
                    if len(header_images) >= 2:
                        break
 
        # ─── نصوص جسم الوثيقة ──────────────────────────────────────
        if 'word/document.xml' in archive_names:
            doc_root = ET.fromstring(archive.read('word/document.xml'))
            ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
            seen_b = set()
            for p in doc_root.findall('.//w:p', ns):
                text = ''.join(t.text or '' for t in p.findall('.//w:t', ns)).strip()
                if text and len(text) > 1 and text not in seen_b:
                    seen_b.add(text)
                    body_lines.append(text)
                    if len(body_lines) >= 50:
                        break
 
    return {
        'header_images': header_images,
        'header_lines':  header_lines[:5],
        'body_lines':    body_lines[:50],
    }

def _get_archive_scope_for_current_user():
    user = get_user_by_id(session.get('user_id'))
    role = (session.get('user_role') or '').lower()

    section_name = (user['archive_section'] if user and user['archive_section'] else 'عام')
    section_code = (user['archive_section_code'] if user and user['archive_section_code'] else 'GN')
    section_code = str(section_code).strip().upper() or 'GN'

    is_scoped = role not in ("admin", "sys_admin")

    if is_scoped:
        # جلب كل الأقسام المُصرح بها من user_section_permissions
        uid = session.get('user_id')
        conn = get_db()
        perm_rows = conn.execute(
            """
            SELECT s.section_code, s.section_name
            FROM user_section_permissions usp
            JOIN archive_sections s ON s.id = usp.section_id
            WHERE usp.user_id = ? AND usp.can_view_archive = 1
            """,
            (uid,)
        ).fetchall()
        conn.close()
        extra_codes = {str(r['section_code']).strip().upper() for r in perm_rows if r['section_code']}
        extra_codes.add(section_code)  # القسم الأساسي دائماً مُدرج
        section_codes = sorted(extra_codes)
    else:
        section_codes = []  # فارغ = كل الأقسام بدون قيود

    return {
        "is_scoped": is_scoped,
        "section_name": section_name,
        "section_code": section_code,
        "section_codes": section_codes,
    }


def _decorate_doc(conn, doc):
    if not doc:
        return doc

    role = (session.get('user_role') or '').lower()
    uid = session.get('user_id')

    pending_sign = conn.execute(
        "SELECT id FROM signature_requests WHERE document_id=? AND requested_from=? AND status='pending' ORDER BY id DESC LIMIT 1",
        (doc['id'], uid)
    ).fetchone()

    is_privileged = role in ('admin', 'manager', 'sys_admin')
    is_owner = bool(uid and uid == doc.get('created_by'))

    # التوقيع المباشر: admin/sys_admin دائمًا، manager فقط إذا كان لديه can_stamp=1 في قسم الوثيقة
    can_sign = False
    if role in ('admin', 'sys_admin'):
        can_sign = True
    elif role == 'manager' and uid:
        section_code = doc.get('archive_section_code') or ''
        perm_row = conn.execute(
            """SELECT 1 FROM user_section_permissions usp
               JOIN archive_sections s ON s.id = usp.section_id
               WHERE usp.user_id=? AND s.section_code=? AND usp.can_stamp=1
               LIMIT 1""",
            (uid, section_code)
        ).fetchone()
        can_sign = bool(perm_row)

    doc['needs_sign'] = bool(pending_sign)
    doc['pending_sign_request_id'] = pending_sign['id'] if pending_sign else None
    doc['can_sign_direct'] = bool(doc.get('status') != 'approved' and can_sign)
    doc['can_delete'] = bool(uid and (is_owner or is_privileged))
    doc['can_edit_text'] = bool(
        uid
        and uid == doc.get('created_by')
        and not doc.get('file_path')
        and doc.get('status') != 'approved'
    )
    doc['can_approve'] = bool(doc.get('status') == 'pending' and is_privileged)
    doc['can_download'] = bool(doc.get('file_path') or doc.get('archive_storage_path') or doc.get('content_json'))
    return doc


def _svg_response(svg_markup, status=200):
    return app.response_class(svg_markup, status=status, mimetype="image/svg+xml")


def _svg_placeholder(text_lines, width=1000, height=1400, subtitle="معاينة مستند"):
    safe_lines = [html.escape(str(line or "")) for line in text_lines if str(line or "").strip()]
    line_markup = "".join(
        f"<text x='50%' y='{220 + (idx * 54)}' text-anchor='middle' fill='#334155' font-size='34' font-family='Arial'>{line}</text>"
        for idx, line in enumerate(safe_lines[:6])
    )
    safe_subtitle = html.escape(subtitle)
    return f"""
    <svg xmlns='http://www.w3.org/2000/svg' width='{width}' height='{height}' viewBox='0 0 {width} {height}'>
      <rect width='100%' height='100%' fill='#f8fafc'/>
      <rect x='48' y='48' width='{width-96}' height='{height-96}' rx='16' ry='16' fill='white' stroke='#cbd5e1' stroke-width='2'/>
      <text x='50%' y='130' text-anchor='middle' fill='#1e40af' font-size='44' font-family='Arial' font-weight='700'>SignMy</text>
      <text x='50%' y='178' text-anchor='middle' fill='#64748b' font-size='24' font-family='Arial'>{safe_subtitle}</text>
      {line_markup}
    </svg>
    """


def _extract_content_html(content_json):
    raw = str(content_json or '').strip()
    if not raw:
        return ''

    try:
        parsed = json.loads(raw)
        if isinstance(parsed, dict):
            html_value = parsed.get('html')
            if html_value is not None:
                return str(html_value)
            text_value = parsed.get('text')
            if text_value is not None:
                return html.escape(str(text_value)).replace('\n', '<br>')
        if isinstance(parsed, str):
            return parsed
    except Exception:
        pass

    return raw


def _docx_templates_folder():
    folder = os.path.join(UPLOAD_FOLDER, 'templates')
    os.makedirs(folder, exist_ok=True)
    return folder


def _sanitize_docx_template_name(template_name=''):
    safe_name = os.path.basename(str(template_name or '').strip())
    if not safe_name or not safe_name.lower().endswith('.docx'):
        return ''
    return safe_name


def _sanitize_any_template_name(template_name=''):
    """Accept both .docx and .pdf template names."""
    safe_name = os.path.basename(str(template_name or '').strip())
    low = safe_name.lower()
    if not safe_name or not (low.endswith('.docx') or low.endswith('.pdf')):
        return ''
    return safe_name


def _resolve_docx_template_file(template_name=''):
    safe_name = _sanitize_docx_template_name(template_name)
    if not safe_name:
        return None

    folder = _docx_templates_folder()
    folder_abs = os.path.abspath(folder)
    target_abs = os.path.abspath(os.path.join(folder_abs, safe_name))
    if os.path.commonpath([folder_abs, target_abs]) != folder_abs:
        return None
    return target_abs


def _resolve_docx_template_path(template_name=''):
    target_abs = _resolve_docx_template_file(template_name)
    if not target_abs:
        return None
    if not os.path.isfile(target_abs):
        return None
    return target_abs


def _docx_template_meta(template_name=''):
    safe_name = _sanitize_docx_template_name(template_name)
    if not safe_name:
        return None

    target_path = _resolve_docx_template_file(safe_name)
    if not target_path or not os.path.isfile(target_path):
        return None

    stat_info = os.stat(target_path)
    return {
        "file_name": safe_name,
        "size_kb": int((stat_info.st_size or 0) / 1024),
        "updated_at": datetime.fromtimestamp(stat_info.st_mtime).strftime('%Y-%m-%d %H:%M')
    }


def _list_docx_templates():
    folder = _docx_templates_folder()
    items = []
    for name in os.listdir(folder):
        low = name.lower()
        if low.endswith('.json'):
            continue
        if not (low.endswith('.docx') or low.endswith('.pdf')):
            continue
        full_path = os.path.join(folder, name)
        if not os.path.isfile(full_path):
            continue
        stat_info = os.stat(full_path)
        items.append({
            "file_name": name,
            "file_type": "pdf" if low.endswith('.pdf') else "docx",
            "size_kb": int((stat_info.st_size or 0) / 1024),
            "updated_at": datetime.fromtimestamp(stat_info.st_mtime).strftime('%Y-%m-%d %H:%M'),
            "_sort_ts": float(stat_info.st_mtime or 0),
        })

    items.sort(key=lambda item: item.get('_sort_ts', 0), reverse=True)
    for item in items:
        item.pop('_sort_ts', None)
    return items


# ── PDF template helpers ─────────────────────────────────────────────────────

def _sanitize_pdf_template_name(template_name=''):
    safe_name = os.path.basename(str(template_name or '').strip())
    if not safe_name or not safe_name.lower().endswith('.pdf'):
        return ''
    return safe_name


def _resolve_pdf_template_file(template_name=''):
    safe_name = _sanitize_pdf_template_name(template_name)
    if not safe_name:
        return None
    folder = _docx_templates_folder()
    folder_abs = os.path.abspath(folder)
    target_abs = os.path.abspath(os.path.join(folder_abs, safe_name))
    if os.path.commonpath([folder_abs, target_abs]) != folder_abs:
        return None
    return target_abs


_PDF_TEMPLATE_DEFAULT_CONFIG = {
    "archive_number": {"x": 0.92, "y": 0.068, "font_size": 11},
    "date":           {"x": 0.92, "y": 0.112, "font_size": 10},
    "content":        {"x_right": 0.92, "x_left": 0.08, "y_start": 0.22, "y_end": 0.83, "font_size": 12, "line_height": 20},
    "qr":             {"x": 0.04, "y": 0.875, "size": 0.11},
}


def _get_pdf_template_config(template_name=''):
    safe_name = _sanitize_pdf_template_name(template_name)
    if not safe_name:
        return dict(_PDF_TEMPLATE_DEFAULT_CONFIG)
    config_path = _resolve_pdf_template_file(safe_name)
    if config_path:
        config_path = config_path + '.json'
    if config_path and os.path.isfile(config_path):
        try:
            with open(config_path, 'r', encoding='utf-8') as _f:
                data = json.load(_f)
            if isinstance(data, dict):
                merged = {k: dict(v) for k, v in _PDF_TEMPLATE_DEFAULT_CONFIG.items()}
                for k, v in data.items():
                    if k in merged and isinstance(v, dict):
                        merged[k].update(v)
                    else:
                        merged[k] = v
                return merged
        except Exception:
            pass
    return dict(_PDF_TEMPLATE_DEFAULT_CONFIG)


def _save_pdf_template_config(template_name='', config=None):
    safe_name = _sanitize_pdf_template_name(template_name)
    if not safe_name:
        return False
    pdf_path = _resolve_pdf_template_file(safe_name)
    if not pdf_path:
        return False
    config_path = pdf_path + '.json'
    try:
        with open(config_path, 'w', encoding='utf-8') as _f:
            json.dump(config or {}, _f, ensure_ascii=False, indent=2)
        return True
    except Exception:
        return False


def _get_app_setting(setting_key, default_value=''):
    key = str(setting_key or '').strip()
    if not key:
        return str(default_value or '')

    conn = get_db()
    row = conn.execute(
        "SELECT setting_value FROM app_settings WHERE setting_key=? LIMIT 1",
        (key,)
    ).fetchone()
    conn.close()

    if not row:
        return str(default_value or '')
    return str(row['setting_value'] or '').strip()


def _set_app_setting(setting_key, setting_value=''):
    key = str(setting_key or '').strip()
    if not key:
        return

    value = str(setting_value or '').strip()
    conn = get_db()
    conn.execute(
        """
        INSERT INTO app_settings (setting_key, setting_value, updated_at)
        VALUES (?, ?, datetime('now'))
        ON CONFLICT(setting_key)
        DO UPDATE SET setting_value=excluded.setting_value, updated_at=datetime('now')
        """,
        (key, value)
    )
    conn.commit()
    conn.close()


def _get_default_docx_template_name():
    configured = os.path.basename(
        _get_app_setting(DEFAULT_DOCX_TEMPLATE_SETTING_KEY, '')
    ).strip()
    if not configured:
        return ''

    # Support both DOCX and PDF templates
    if configured.lower().endswith('.pdf'):
        resolved = _resolve_pdf_template_file(configured)
        if resolved and os.path.isfile(resolved):
            return configured
    else:
        if _resolve_docx_template_path(configured):
            return configured

    _set_app_setting(DEFAULT_DOCX_TEMPLATE_SETTING_KEY, '')
    return ''


def _set_default_docx_template_name(template_name=''):
    safe_name = os.path.basename(str(template_name or '').strip())
    if safe_name:
        is_pdf = safe_name.lower().endswith('.pdf')
        is_docx = safe_name.lower().endswith('.docx')
        if not is_pdf and not is_docx:
            return False, "اسم القالب غير صالح"
        if is_pdf:
            resolved = _resolve_pdf_template_file(safe_name)
            if not resolved or not os.path.isfile(resolved):
                return False, "القالب غير موجود"
        else:
            if not _resolve_docx_template_path(safe_name):
                return False, "القالب غير موجود"

    _set_app_setting(DEFAULT_DOCX_TEMPLATE_SETTING_KEY, safe_name)
    return True, safe_name


def _path_to_data_uri(path_value='', max_bytes=4 * 1024 * 1024):
    resolved = _resolve_document_file_path(path_value)
    if not resolved or not os.path.exists(resolved):
        return None

    try:
        with open(resolved, 'rb') as handle:
            blob = handle.read()
    except OSError:
        return None

    if not blob or len(blob) > max_bytes:
        return None

    mime_type = mimetypes.guess_type(resolved)[0] or 'application/octet-stream'
    return f"data:{mime_type};base64,{base64.b64encode(blob).decode('ascii')}"


def _stamp_text_to_data_uri(stamp_text='STAMP', owner_name='—'):
    text_value = str(stamp_text or 'STAMP').strip() or 'STAMP'
    name_value = str(owner_name or '—').strip() or '—'

    # PNG works reliably across PIL-based preview and final document generation.
    try:
        from PIL import Image, ImageDraw, ImageFont

        img = Image.new('RGBA', (220, 220), (0, 0, 0, 0))
        draw = ImageDraw.Draw(img)

        draw.ellipse((26, 26, 194, 194), fill=(79, 126, 248, 20), outline=(79, 126, 248, 235), width=4)
        draw.ellipse((46, 46, 174, 174), outline=(79, 126, 248, 200), width=2)

        def _ar(text):
            try:
                import arabic_reshaper as _arm
                from bidi.algorithm import get_display as _gd
                return _gd(_arm.reshape(str(text or '')))
            except Exception:
                return str(text or '')

        # Prefer Arabic-capable fonts for stamp text to avoid square glyphs.
        font_path = None
        for fp in (
            r'C:\Windows\Fonts\tahomabd.ttf',
            r'C:\Windows\Fonts\tahoma.ttf',
            r'C:\Windows\Fonts\arialbd.ttf',
            r'C:\Windows\Fonts\arial.ttf',
            '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf',
            '/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf',
            '/usr/share/fonts/truetype/noto/NotoSansArabic-Regular.ttf',
            '/usr/share/fonts/truetype/noto/NotoNaskhArabic-Regular.ttf',
            os.path.join(app.root_path, 'static', 'fonts', 'NotoSansArabic-Regular.ttf'),
        ):
            if os.path.exists(fp):
                font_path = fp
                break

        def _fnt(size):
            try:
                return ImageFont.truetype(font_path, size) if font_path else ImageFont.load_default()
            except Exception:
                return ImageFont.load_default()

        text_font = _fnt(20)
        name_font = _fnt(12)

        text_render = _ar(text_value)
        name_render = _ar(name_value)

        text_bb = draw.textbbox((0, 0), text_render, font=text_font)
        text_w = text_bb[2] - text_bb[0]
        draw.text(((220 - text_w) // 2, 98), text_render, fill=(30, 64, 175, 255), font=text_font)

        name_bb = draw.textbbox((0, 0), name_render, font=name_font)
        name_w = name_bb[2] - name_bb[0]
        draw.text(((220 - name_w) // 2, 130), name_render, fill=(71, 85, 105, 255), font=name_font)

        buf = io.BytesIO()
        img.save(buf, 'PNG')
        buf.seek(0)
        return f"data:image/png;base64,{base64.b64encode(buf.read()).decode('ascii')}"
    except Exception:
        safe_text = html.escape(text_value)
        safe_name = html.escape(name_value)
        svg = f"""
        <svg xmlns='http://www.w3.org/2000/svg' width='220' height='220' viewBox='0 0 220 220'>
          <rect width='100%' height='100%' fill='transparent'/>
          <circle cx='110' cy='110' r='84' fill='rgba(79,126,248,0.08)' stroke='#4f7ef8' stroke-width='4'/>
          <circle cx='110' cy='110' r='64' fill='none' stroke='#4f7ef8' stroke-width='2' stroke-dasharray='4 4'/>
          <text x='110' y='108' text-anchor='middle' fill='#1e40af' font-family='Arial' font-size='20' font-weight='700'>{safe_text}</text>
          <text x='110' y='138' text-anchor='middle' fill='#475569' font-family='Arial' font-size='12'>{safe_name}</text>
        </svg>
        """.strip()
        return f"data:image/svg+xml;base64,{base64.b64encode(svg.encode('utf-8')).decode('ascii')}"


# ── قوالب الختم (stamp templates) ───────────────────────────────────────────

def _get_default_stamp_tpl(conn):
    """يرجع صف القالب الافتراضي أو الأحدث من جدول stamp_templates."""
    row = conn.execute(
        "SELECT * FROM stamp_templates WHERE is_default=1 LIMIT 1"
    ).fetchone()
    if not row:
        row = conn.execute(
            "SELECT * FROM stamp_templates ORDER BY id DESC LIMIT 1"
        ).fetchone()
    return row


def _create_default_stamp_asset_for_user(user_id, conn=None):
    """Backward-compatible wrapper for manager stamp auto-sync."""
    return _sync_manager_section_stamp_assets(user_id, conn=conn)


def _sync_manager_section_stamp_assets(user_id, conn=None):
    """Ensure manager/admin users have one auto-stamp per can_stamp section permission.

    - Creates missing auto stamps for active can_stamp sections.
    - Deactivates removed auto stamps when section permission is revoked.
    - Keeps manually uploaded stamps untouched.
    """
    uid = _safe_int(user_id, 0)
    if uid <= 0:
        return False

    own_conn = conn is None
    db = conn or get_db()
    tmp_path = None
    try:
        user_row = db.execute(
            """
            SELECT u.id, u.name, u.role, u.employee_id, s.section_code
            FROM users u
            LEFT JOIN archive_sections s ON s.id = u.archive_section_id
            WHERE u.id=? LIMIT 1
            """,
            (uid,)
        ).fetchone()
        if not user_row:
            return False

        role_value = str(user_row['role'] or '').strip().lower()
        auto_prefix = 'AUTO_SECTION:'
        privileged_roles = ('manager', 'admin', 'sys_admin')
        if role_value not in privileged_roles:
            db.execute(
                "UPDATE stamp_assets SET is_active=0 WHERE user_id=? AND is_active=1 AND stamp_name LIKE ?",
                (uid, f"{auto_prefix}%")
            )
            if own_conn:
                db.commit()
            return False

        tpl_row = _get_default_stamp_tpl(db)
        if not tpl_row:
            return False

        section_rows = db.execute(
            """
            SELECT UPPER(TRIM(s.section_code)) AS section_code
            FROM user_section_permissions usp
            JOIN archive_sections s ON s.id = usp.section_id
            WHERE usp.user_id=? AND usp.can_stamp=1 AND s.is_active=1
            ORDER BY s.section_name ASC
            """,
            (uid,)
        ).fetchall()
        target_sections = []
        for row in section_rows:
            code = str(row['section_code'] or '').strip().upper()
            if re.fullmatch(r'[A-Z]{2}', code) and code not in target_sections:
                target_sections.append(code)

        primary_code = str(user_row['section_code'] or '').strip().upper()
        if not target_sections and re.fullmatch(r'[A-Z]{2}', primary_code):
            target_sections.append(primary_code)

        if not target_sections:
            target_sections.append('GN')

        existing_auto_rows = db.execute(
            """
            SELECT id, stamp_name
            FROM stamp_assets
            WHERE user_id=? AND is_active=1 AND stamp_name LIKE ?
            """,
            (uid, f"{auto_prefix}%")
        ).fetchall()
        existing_by_code = {}
        for row in existing_auto_rows:
            raw_name = str(row['stamp_name'] or '')
            section_code = raw_name.split(':', 1)[1].strip().upper() if ':' in raw_name else ''
            if re.fullmatch(r'[A-Z]{2}', section_code):
                existing_by_code[section_code] = _safe_int(row['id'], 0)

        emp = str(user_row['employee_id'] or '').strip()
        created_any = False
        os.makedirs(os.path.join(UPLOAD_FOLDER, 'stamps'), exist_ok=True)
        for sec in target_sections:
            if sec in existing_by_code:
                continue

            tmp_path = _render_stamp_from_tpl(
                tpl_row['file_name'],
                tpl_row['text_x_ratio'],
                tpl_row['text_y_ratio'],
                sec,
                emp,
                '',
                as_file=True,
            )
            if not tmp_path or not os.path.exists(tmp_path):
                continue

            final_name = f"stamp_auto_{uid}_{sec}_{int(datetime.utcnow().timestamp())}.png"
            final_abs = os.path.join(UPLOAD_FOLDER, 'stamps', final_name)

            with open(tmp_path, 'rb') as src, open(final_abs, 'wb') as dst:
                dst.write(src.read())

            db.execute(
                "INSERT INTO stamp_assets (user_id, stamp_name, stamp_path, visibility_scope, is_active) VALUES (?,?,?,?,1)",
                (uid, f"{auto_prefix}{sec}", f"uploads/stamps/{final_name}", 'self')
            )
            created_any = True
            try:
                os.remove(tmp_path)
            except OSError:
                pass
            tmp_path = None

        removed_codes = [code for code in existing_by_code.keys() if code not in target_sections]
        for code in removed_codes:
            db.execute(
                "UPDATE stamp_assets SET is_active=0 WHERE id=?",
                (existing_by_code[code],)
            )

        if own_conn:
            db.commit()
        return created_any or bool(target_sections)
    except Exception:
        if own_conn:
            db.rollback()
        return False
    finally:
        if tmp_path and os.path.exists(tmp_path):
            try:
                os.remove(tmp_path)
            except OSError:
                pass
        if own_conn:
            db.close()


def _render_stamp_from_tpl(file_name, text_x_ratio=0.25, text_y_ratio=0.08,
                            section_code='', employee_id='', serial_short='',
                            as_file=False):
    """
    يفتح صورة قالب الختم ويكتب البيانات بصيغة:
    رقم القسم - الرقم الوظيفي - رقم الأرشيف/المرجع
    دون تغيير أي شيء آخر في الصورة.
    - as_file=False  → data-uri PNG
    - as_file=True   → مسار ملف مؤقت
    """
    stamps_dir = os.path.join(UPLOAD_FOLDER, 'stamps')
    tpl_path = os.path.join(stamps_dir, os.path.basename(str(file_name or '')))
    if not os.path.isfile(tpl_path):
        return None
    try:
        from PIL import Image, ImageDraw, ImageFont
    except ImportError:
        return None
    try:
        img = Image.open(tpl_path).convert('RGBA')
        w, h = img.size

        # Draw on a larger canvas first for sharper text, then downsample.
        render_scale = 2 if max(w, h) >= 260 else 3
        work_img = img.resize((w * render_scale, h * render_scale), Image.LANCZOS)
        draw = ImageDraw.Draw(work_img)
        sw, sh = work_img.size

        # Prefer Arabic-capable fonts for stamp metadata text.
        font_path = None
        for fp in [
            r'C:\Windows\Fonts\tahomabd.ttf',
            r'C:\Windows\Fonts\tahoma.ttf',
            r'C:\Windows\Fonts\arialbd.ttf',
            r'C:\Windows\Fonts\arial.ttf',
            '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf',
            '/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf',
            '/usr/share/fonts/truetype/noto/NotoSansArabic-Regular.ttf',
            '/usr/share/fonts/truetype/noto/NotoNaskhArabic-Regular.ttf',
            os.path.join(app.root_path, 'static', 'fonts', 'NotoSansArabic-Regular.ttf'),
        ]:
            if os.path.exists(fp):
                font_path = fp
                break

        def _fnt(size):
            try:
                return ImageFont.truetype(font_path, size) if font_path else ImageFont.load_default()
            except Exception:
                return ImageFont.load_default()

        def _ar(text):
            try:
                import arabic_reshaper as _arm
                from bidi.algorithm import get_display as _gd
                return _gd(_arm.reshape(str(text or '')))
            except Exception:
                return str(text or '')

        # Darker blue ink improves legibility while keeping the stamp style.
        ink = (6, 34, 136, 255)
        ink_shadow = (2, 18, 78, 190)
        cx  = int(sw * float(text_x_ratio))
        by  = int(sh * float(text_y_ratio))

        serial_value = serial_short.strip()
        main_parts = [p for p in [section_code.strip().upper(), employee_id.strip()] if p]
        main_text = _ar(' - '.join(main_parts)) if main_parts else ''
        fnt_size = max(16, int(sw * 0.034))

        if main_text:
            # Auto-fit text width while keeping the stamp metadata compact.
            while fnt_size > 12:
                fnt = _fnt(fnt_size)
                bb = draw.textbbox((0, 0), main_text, font=fnt)
                tw = bb[2] - bb[0]
                if tw <= int(sw * 0.90):
                    break
                fnt_size -= 1

            left_nudge = int(sw * 0.008)
            right_shift = int(sw * 0.004)
            txt_x = cx - tw // 2 - left_nudge + right_shift
            shadow_off = max(1, render_scale)
            draw.text((txt_x + shadow_off, by + shadow_off), main_text, font=fnt, fill=ink_shadow)
            draw.text((txt_x, by), main_text, font=fnt, fill=ink)

        # Draw serial on a dedicated line so long references don't pull main text left.
        if serial_value:
            serial_text = _ar(serial_value)
            serial_size = max(14, int(sw * 0.032))
            while serial_size > 10:
                sfnt = _fnt(serial_size)
                sbb = draw.textbbox((0, 0), serial_text, font=sfnt)
                stw = sbb[2] - sbb[0]
                if stw <= int(sw * 0.78):
                    break
                serial_size -= 1

            serial_shift = int(sw * 0.180)
            serial_x = cx - stw // 2 + serial_shift
            serial_y = by + int(4 * render_scale)
            shadow_off = max(1, render_scale)
            draw.text((serial_x + shadow_off, serial_y + shadow_off), serial_text, font=sfnt, fill=ink_shadow)
            draw.text((serial_x, serial_y), serial_text, font=sfnt, fill=ink)

        # Downsample back to original template size while preserving sharp edges.
        img = work_img.resize((w, h), Image.LANCZOS)

        if as_file:
            tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
            tmp.close()
            img.save(tmp.name, 'PNG', optimize=True)
            return tmp.name
        else:
            buf = io.BytesIO()
            img.save(buf, 'PNG', optimize=True)
            buf.seek(0)
            return f"data:image/png;base64,{base64.b64encode(buf.read()).decode('ascii')}"
    except Exception:
        return None


def _manager_round_stamp_to_data_uri(company_name='شركة SignMy', section_code='', employee_id='', serial_short=''):
    """
    ختم دائري رسمي للمديرين يشبه الختم التجاري:
    - النص العلوي الدائري:  company_name
    - النص السفلي الدائري: ذات مسؤولية محدودة
    - الوسط: section_code  |  employee_id
    - أسفل الوسط: رقم مرجعي صغير (serial_short)
    """
    cx, cy, r_outer, r_inner = 110, 110, 96, 76
    sec  = html.escape(str(section_code  or '').strip().upper())
    emp  = html.escape(str(employee_id   or '').strip())
    ser  = html.escape(str(serial_short  or '').strip())
    co   = html.escape(str(company_name  or 'شركة SignMy').strip())

    # mid-line label
    mid_label = sec
    if emp:
        mid_label = f"{sec}  |  {emp}" if sec else emp

    svg = f"""<svg xmlns='http://www.w3.org/2000/svg' width='220' height='220' viewBox='0 0 220 220'>
  <defs>
    <style>
      @font-face{{font-family:'Noto';}}
      text{{font-family:'Tahoma','Arial',sans-serif;}}
    </style>
    <!-- مسار دائري علوي للنص (من اليمين لليسار فوق الدائرة) -->
    <path id='topArc'
      d='M {cx - r_outer + 6},{cy}
         A {r_outer - 6},{r_outer - 6} 0 0,1
           {cx + r_outer - 6},{cy}'/>
    <!-- مسار دائري سفلي للنص -->
    <path id='botArc'
      d='M {cx - r_outer + 6},{cy}
         A {r_outer - 6},{r_outer - 6} 0 0,0
           {cx + r_outer - 6},{cy}'/>
  </defs>

  <!-- خلفية شفافة -->
  <rect width='220' height='220' fill='transparent'/>

  <!-- الدائرة الخارجية (سميكة) -->
  <circle cx='{cx}' cy='{cy}' r='{r_outer}' fill='rgba(15,30,80,0.04)'
          stroke='#0f1e50' stroke-width='3.5'/>
  <!-- الدائرة الداخلية (رفيعة) -->
  <circle cx='{cx}' cy='{cy}' r='{r_inner}' fill='none'
          stroke='#0f1e50' stroke-width='1.4'/>

  <!-- النص العلوي الدائري (اسم الشركة) -->
  <text fill='#0f1e50' font-size='13' font-weight='700' text-anchor='middle'>
    <textPath href='#topArc' startOffset='50%'>{co}</textPath>
  </text>

  <!-- النص السفلي الدائري (ذات مسؤولية محدودة) -->
  <text fill='#0f1e50' font-size='11' font-weight='600' text-anchor='middle'>
    <textPath href='#botArc' startOffset='50%'>ذات مسؤولية محدودة</textPath>
  </text>

  <!-- خط فاصل أفقي صغير في الوسط -->
  <line x1='{cx - 38}' y1='{cy - 2}' x2='{cx + 38}' y2='{cy - 2}'
        stroke='#0f1e50' stroke-width='0.8' opacity='0.5'/>

  <!-- النص الرئيسي (القسم + الرقم الوظيفي) -->
    <text x='{cx + 4}' y='{cy + 14}' text-anchor='middle'
      fill='#0f1e50' font-size='12' font-weight='700'
        letter-spacing='1'>{html.escape(mid_label)}</text>

  <!-- الرقم المرجعي الصغير (serial) -->
    {'<text x="' + str(cx + 8) + '" y="' + str(cy + 32) + '" text-anchor="middle"' +
     ' fill="#334155" font-size="6" font-family="monospace">#' + ser + '</text>' if ser else ''}
</svg>"""
    return f"data:image/svg+xml;base64,{base64.b64encode(svg.encode('utf-8')).decode('ascii')}"


def _html_to_text_lines(content_html):
    raw_html = str(content_html or '').strip()
    if not raw_html:
        return []

    normalized = re.sub(r'(?i)<br\s*/?>', '\n', raw_html)
    normalized = re.sub(r'(?i)</(p|div|h[1-6]|li|tr)>', '\n', normalized)
    normalized = re.sub(r'(?i)<[^>]+>', '', normalized)
    normalized = html.unescape(normalized)

    lines = [line.strip() for line in normalized.splitlines() if line.strip()]
    return lines[:400]


def _short_signature_serial(serial_value, section_code=None, keep=3):
    serial = str(serial_value or '').strip()
    if not serial:
        return '—'

    inferred_section = ''
    serial_upper = serial.upper()

    # Prefer archive-like base when serial is "ARCHIVE-02" to keep display stable.
    base_candidate = serial_upper.split('-', 1)[0].strip()
    base_clean = re.sub(r'[^0-9A-Z]', '', base_candidate)
    # Examples supported: GN001, GN0326001, HR1225007
    base_match = re.match(r'^([A-Z]{2})([0-9]{3,})$', base_clean)
    if base_match:
        inferred_section = base_match.group(1)
        tail_source = base_match.group(2)
    else:
        tail_source = serial_upper.split('-')[-1] if '-' in serial_upper else serial_upper

    tail = re.sub(r'[^0-9A-Z]', '', tail_source) or re.sub(r'[^0-9A-Z]', '', serial_upper) or serial_upper
    keep_count = max(1, _safe_int(keep, 3))
    short_tail = tail[-keep_count:] if len(tail) > keep_count else tail

    section = re.sub(r'[^A-Za-z0-9]', '', str(section_code or '').upper()).strip()
    if not section and inferred_section:
        section = inferred_section
    if section:
        return f"{section}{short_tail}"
    return short_tail


def _build_signature_serial(conn, doc_id):
    """Prefer archive number as signature serial, with safe uniqueness fallback."""
    doc_row = conn.execute(
        "SELECT archive_number FROM documents WHERE id=? LIMIT 1",
        (doc_id,)
    ).fetchone()
    archive_number = str(doc_row['archive_number'] or '').strip().upper() if doc_row else ''
    if not archive_number:
        return generate_serial("SIG", "signature_requests", "serial_number")

    exists = conn.execute(
        """
        SELECT 1
        FROM signature_requests
        WHERE UPPER(TRIM(COALESCE(serial_number, ''))) = UPPER(TRIM(?))
        LIMIT 1
        """,
        (archive_number,)
    ).fetchone()
    if not exists:
        return archive_number

    for idx in range(2, 1000):
        candidate = f"{archive_number}-{idx:02d}"
        conflict = conn.execute(
            """
            SELECT 1
            FROM signature_requests
            WHERE UPPER(TRIM(COALESCE(serial_number, ''))) = UPPER(TRIM(?))
            LIMIT 1
            """,
            (candidate,)
        ).fetchone()
        if not conflict:
            return candidate

    return generate_serial("SIG", "signature_requests", "serial_number")


def _get_document_extension(doc):
    path_hint = str(doc.get('file_path') or doc.get('archive_storage_path') or '').strip()
    if not path_hint:
        return ''
    clean = path_hint.split('?', 1)[0].strip().lower()
    if '.' not in clean:
        return ''
    return clean.rsplit('.', 1)[-1]


def _clamp_ratio(value, default=0.0):
    try:
        number = float(value)
    except (TypeError, ValueError):
        return float(default)
    return max(0.0, min(1.0, number))


def _fit_inside_box(box_w, box_h, src_w, src_h):
    """Return scaled width/height and centered offsets while preserving aspect ratio."""
    try:
        bw = float(box_w)
        bh = float(box_h)
        sw = float(src_w)
        sh = float(src_h)
    except (TypeError, ValueError):
        return float(box_w or 0), float(box_h or 0), 0.0, 0.0

    if bw <= 0 or bh <= 0 or sw <= 0 or sh <= 0:
        return max(0.0, bw), max(0.0, bh), 0.0, 0.0

    scale = min(bw / sw, bh / sh)
    draw_w = max(1.0, sw * scale)
    draw_h = max(1.0, sh * scale)
    off_x = max(0.0, (bw - draw_w) / 2.0)
    off_y = max(0.0, (bh - draw_h) / 2.0)
    return draw_w, draw_h, off_x, off_y


def _normalize_sign_positions(raw_positions, include_qr=False):
    source = raw_positions
    if isinstance(source, str):
        try:
            source = json.loads(source)
        except Exception:
            source = []

    if not isinstance(source, list):
        source = []

    defaults = {
        'sig': {'type': 'sig', 'x': 0.55, 'y': 0.78, 'w': 0.25, 'h': 0.12, 'page': 0},
        'stamp': {'type': 'stamp', 'x': 0.70, 'y': 0.72, 'w': 0.36, 'h': 0.36, 'page': 0},
        'qr': {'type': 'qr', 'x': 0.08, 'y': 0.92, 'w': 0.14, 'h': 0.14, 'page': 0},
    }

    min_size = {
        'sig': (0.05, 0.05),
        'stamp': (0.32, 0.32),
        'qr': (0.08, 0.08),
    }

    normalized = []
    for entry in source:
        if not isinstance(entry, dict):
            continue

        item_type = str(entry.get('type') or '').strip().lower()
        if item_type not in defaults:
            continue

        base = defaults[item_type]
        min_w, min_h = min_size.get(item_type, (0.05, 0.05))

        normalized.append({
            'type': item_type,
            'x': _clamp_ratio(entry.get('x'), base['x']),
            'y': _clamp_ratio(entry.get('y'), base['y']),
            'w': max(min_w, _clamp_ratio(entry.get('w'), base['w'])),
            'h': max(min_h, _clamp_ratio(entry.get('h'), base['h'])),
            'page': max(0, _safe_int(entry.get('page'), base['page'])),
        })

    if include_qr and not any(item.get('type') == 'qr' for item in normalized):
        normalized.append(dict(defaults['qr']))

    return normalized


def _pick_sign_position(positions, item_type):
    for item in positions or []:
        if str(item.get('type') or '').strip().lower() == item_type:
            return item
    default_map = {
        'sig': {'type': 'sig', 'x': 0.55, 'y': 0.78, 'w': 0.25, 'h': 0.12, 'page': 0},
        'stamp': {'type': 'stamp', 'x': 0.70, 'y': 0.72, 'w': 0.36, 'h': 0.36, 'page': 0},
        'qr': {'type': 'qr', 'x': 0.08, 'y': 0.92, 'w': 0.14, 'h': 0.14, 'page': 0},
    }
    return dict(default_map.get(item_type, {'type': item_type, 'x': 0.5, 'y': 0.8, 'w': 0.2, 'h': 0.12, 'page': 0}))


def _position_to_docx_alignment(x_ratio):
    try:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except Exception:
        return None

    x_value = _clamp_ratio(x_ratio, 0.7)
    if x_value <= 0.34:
        return WD_ALIGN_PARAGRAPH.LEFT
    if x_value >= 0.66:
        return WD_ALIGN_PARAGRAPH.RIGHT
    return WD_ALIGN_PARAGRAPH.CENTER


def _position_width_inches(width_ratio, default_inches):
    ratio = _clamp_ratio(width_ratio, 0.0)
    if ratio <= 0:
        return default_inches
    return max(0.8, min(3.8, ratio * 7.2))


def _set_paragraph_runs_color(paragraph, red=0, green=0, blue=0):
    try:
        from docx.shared import RGBColor, Pt
        from docx.oxml.ns import qn as _qn
        from docx.oxml import OxmlElement as _OxmlE
        _W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(red, green, blue)
            if run.font.size is None:
                run.font.size = Pt(12)
            rPr = run._r.get_or_add_rPr()
            for c in rPr.findall(_qn('w:color')): rPr.remove(c)
            ce = _OxmlE('w:color')
            ce.set(_qn('w:val'), f'{red:02X}{green:02X}{blue:02X}')
            rPr.insert(0, ce)
            rFonts = rPr.find(_qn('w:rFonts'))
            if rFonts is None:
                rFonts = _OxmlE('w:rFonts'); rPr.insert(0, rFonts)
            rFonts.set(_qn('w:ascii'), 'Arial')
            rFonts.set(_qn('w:hAnsi'), 'Arial')
            rFonts.set(_qn('w:cs'), 'Arial')
        pPr = paragraph._p.get_or_add_pPr()
        rPr_p = pPr.find(f'{{{_W}}}rPr')
        if rPr_p is None:
            rPr_p = _OxmlE('w:rPr'); pPr.append(rPr_p)
        for c in rPr_p.findall(_qn('w:color')): rPr_p.remove(c)
        ce2 = _OxmlE('w:color')
        ce2.set(_qn('w:val'), f'{red:02X}{green:02X}{blue:02X}')
        rPr_p.append(ce2)
    except Exception:
        return

def _build_qr_cells(serial_text, size=25):
    digest = hashlib.sha256(serial_text.encode('utf-8')).digest()
    bitstream = ''.join(format(b, '08b') for b in digest * 20)

    def finder(x0, y0):
        out = []
        for y in range(7):
            for x in range(7):
                border = x in (0, 6) or y in (0, 6)
                center = 2 <= x <= 4 and 2 <= y <= 4
                if border or center:
                    out.append((x0 + x, y0 + y))
        return out

    reserved = set(finder(0, 0) + finder(size - 7, 0) + finder(0, size - 7))
    cells = set(reserved)

    idx = 0
    for y in range(size):
        for x in range(size):
            if (x, y) in reserved:
                continue
            if bitstream[idx % len(bitstream)] == '1':
                cells.add((x, y))
            idx += 1

    return cells, size


def _png_chunk(tag, data):
    chunk_head = tag + data
    return (
        struct.pack('>I', len(data)) +
        chunk_head +
        struct.pack('>I', zlib.crc32(chunk_head) & 0xffffffff)
    )


def _build_qr_png_bytes(serial, cell=8, margin=12):
    qr_text = str(serial or '').strip() or 'NAJM'
    import qrcode
    from qrcode.constants import ERROR_CORRECT_H

    qr = qrcode.QRCode(
        version=None,
        error_correction=ERROR_CORRECT_H,
        box_size=max(8, int(cell)),
        border=max(4, int(margin // 3) or 4),
    )
    qr.add_data(qr_text)
    qr.make(fit=True)

    img = qr.make_image(fill_color='black', back_color='white')
    out = io.BytesIO()
    img.save(out, format='PNG')
    return out.getvalue()


def _qr_png_data_uri(serial):
    try:
        blob = _build_qr_png_bytes(serial)
    except Exception:
        return None
    return f"data:image/png;base64,{base64.b64encode(blob).decode('ascii')}"


def _save_qr_png_temp(serial):
    try:
        blob = _build_qr_png_bytes(serial)
    except Exception:
        return None

    handle = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
    try:
        handle.write(blob)
        handle.flush()
        return handle.name
    finally:
        handle.close()


def _to_emu(value, fallback=0):
    try:
        return int(value)
    except Exception:
        return int(fallback)


def _inline_to_anchor(inline, x_emu, y_emu):
    if OxmlElement is None or qn is None:
        return None

    anchor = OxmlElement('wp:anchor')
    anchor.set('distT', '0')
    anchor.set('distB', '0')
    anchor.set('distL', '0')
    anchor.set('distR', '0')
    anchor.set('simplePos', '0')
    anchor.set('relativeHeight', '251658240')
    anchor.set('behindDoc', '0')
    anchor.set('locked', '0')
    anchor.set('layoutInCell', '1')
    anchor.set('allowOverlap', '1')

    simple_pos = OxmlElement('wp:simplePos')
    simple_pos.set('x', '0')
    simple_pos.set('y', '0')
    anchor.append(simple_pos)

    position_h = OxmlElement('wp:positionH')
    position_h.set('relativeFrom', 'page')
    pos_h = OxmlElement('wp:posOffset')
    pos_h.text = str(_to_emu(x_emu, 0))
    position_h.append(pos_h)
    anchor.append(position_h)

    position_v = OxmlElement('wp:positionV')
    position_v.set('relativeFrom', 'page')
    pos_v = OxmlElement('wp:posOffset')
    pos_v.text = str(_to_emu(y_emu, 0))
    position_v.append(pos_v)
    anchor.append(position_v)

    extent = OxmlElement('wp:extent')
    extent.set('cx', str(_to_emu(getattr(inline.extent, 'cx', 0), 0)))
    extent.set('cy', str(_to_emu(getattr(inline.extent, 'cy', 0), 0)))
    anchor.append(extent)

    effect_extent = OxmlElement('wp:effectExtent')
    effect_extent.set('l', '0')
    effect_extent.set('t', '0')
    effect_extent.set('r', '0')
    effect_extent.set('b', '0')
    anchor.append(effect_extent)
    anchor.append(OxmlElement('wp:wrapNone'))

    doc_pr = OxmlElement('wp:docPr')
    doc_pr.set('id', str(_to_emu(getattr(inline.docPr, 'id', 1), 1)))
    doc_pr.set('name', str(getattr(inline.docPr, 'name', '') or f"Picture {_to_emu(getattr(inline.docPr, 'id', 1), 1)}"))
    anchor.append(doc_pr)

    cnv = OxmlElement('wp:cNvGraphicFramePr')
    locks = OxmlElement('a:graphicFrameLocks')
    locks.set('noChangeAspect', '1')
    cnv.append(locks)
    anchor.append(cnv)

    anchor.append(deepcopy(inline.graphic))
    return anchor


def _add_docx_floating_image(document, image_path, position, width_inches=1.4, anchor_paragraph=None):
    if not DOCX_AVAILABLE or not image_path or not os.path.exists(image_path):
        return False

    try:
        section = document.sections[0]
    except Exception:
        return False

    page_w = _to_emu(getattr(section, 'page_width', 0), 0)
    page_h = _to_emu(getattr(section, 'page_height', 0), 0)
    left_margin = _to_emu(getattr(section, 'left_margin', 0), 0)
    right_margin = _to_emu(getattr(section, 'right_margin', 0), 0)
    top_margin = _to_emu(getattr(section, 'top_margin', 0), 0)
    bottom_margin = _to_emu(getattr(section, 'bottom_margin', 0), 0)

    content_w = max(1, page_w - left_margin - right_margin)
    content_h = max(1, page_h - top_margin - bottom_margin)

    x_ratio = _clamp_ratio(position.get('x'), 0.5)
    y_ratio = _clamp_ratio(position.get('y'), 0.8)

    x_emu = left_margin + int(x_ratio * content_w)
    y_emu = top_margin + int(y_ratio * content_h)

    paragraph = anchor_paragraph or document.add_paragraph('')
    run = paragraph.add_run()
    run.add_picture(image_path, width=Inches(max(0.7, min(4.2, float(width_inches)))))

    inline_nodes = run._r.xpath('./w:drawing/wp:inline')
    if not inline_nodes:
        return False

    inline = inline_nodes[0]
    anchor = _inline_to_anchor(inline, x_emu, y_emu)
    if anchor is None:
        return False

    drawing = inline.getparent()
    drawing.remove(inline)
    drawing.append(anchor)
    return True



def _force_docx_body_text_black(document):
    from docx.shared import RGBColor, Pt
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement as _OxmlE
    _W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    def _clear_run_shading(rPr):
        # إزالة التظليل الرمادي من خصائص النص (w:shd, w:highlight)
        for tag in (_qn('w:shd'), _qn('w:highlight')):
            for el in rPr.findall(tag): rPr.remove(el)

    def _blk_run(run):
        try:
            run.font.color.rgb = RGBColor(0, 0, 0)
            if run.font.size is None: run.font.size = Pt(12)
            rPr = run._r.get_or_add_rPr()
            for c in rPr.findall(_qn('w:color')): rPr.remove(c)
            ce = _OxmlE('w:color'); ce.set(_qn('w:val'), '000000'); rPr.insert(0, ce)
            rFonts = rPr.find(_qn('w:rFonts'))
            if rFonts is None:
                rFonts = _OxmlE('w:rFonts'); rPr.insert(0, rFonts)
            rFonts.set(_qn('w:ascii'), 'Arial')
            rFonts.set(_qn('w:hAnsi'), 'Arial')
            rFonts.set(_qn('w:cs'), 'Arial')
            _clear_run_shading(rPr)
        except Exception: pass

    def _blk_para(paragraph):
        for run in paragraph.runs: _blk_run(run)
        try:
            pPr = paragraph._p.get_or_add_pPr()
            # إزالة تظليل الفقرة (السبب الرئيسي للون الرمادي)
            for s in pPr.findall(_qn('w:shd')): pPr.remove(s)
            shd = _OxmlE('w:shd')
            shd.set(_qn('w:val'), 'clear')
            shd.set(_qn('w:color'), 'auto')
            shd.set(_qn('w:fill'), 'FFFFFF')
            # إدراج shd قبل rPr لاحترام ترتيب OOXML (shd يجب أن يسبق rPr)
            rPr_p = pPr.find(f'{{{_W}}}rPr')
            if rPr_p is not None:
                rPr_p.addprevious(shd)
            else:
                pPr.append(shd)
            if rPr_p is None:
                rPr_p = _OxmlE('w:rPr'); pPr.append(rPr_p)
            for c in rPr_p.findall(_qn('w:color')): rPr_p.remove(c)
            ce = _OxmlE('w:color'); ce.set(_qn('w:val'), '000000'); rPr_p.append(ce)
            _clear_run_shading(rPr_p)
        except Exception: pass

    # إزالة خلفية الصفحة الداكنة وضمان خلفية بيضاء
    try:
        body = document.element.body
        doc_el = body.getparent()
        if doc_el is not None:
            for bg in doc_el.findall(_qn('w:background')): doc_el.remove(bg)
            new_bg = _OxmlE('w:background')
            new_bg.set(_qn('w:color'), 'FFFFFF')
            doc_el.insert(0, new_bg)
    except Exception: pass
    # إيقاف displayBackgroundShape لمنع الخلفية السوداء
    try:
        _W_SET = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        set_el = document.settings.element
        for dbs in list(set_el.findall(f'{{{_W_SET}}}displayBackgroundShape')):
            set_el.remove(dbs)
    except Exception: pass

    for p in document.paragraphs: _blk_para(p)
    for t in document.tables:
        for row in t.rows:
            for cell in row.cells:
                # إزالة تظليل خلايا الجدول
                try:
                    from docx.oxml.ns import qn as _qn2
                    tc_pr = cell._tc.get_or_add_tcPr()
                    for s in tc_pr.findall(_qn2('w:shd')): tc_pr.remove(s)
                    cell_shd = _OxmlE('w:shd')
                    cell_shd.set(_qn2('w:val'), 'clear')
                    cell_shd.set(_qn2('w:color'), 'auto')
                    cell_shd.set(_qn2('w:fill'), 'FFFFFF')
                    tc_pr.append(cell_shd)
                except Exception: pass
                for p in cell.paragraphs: _blk_para(p)

def _replace_docx_tokens(document, replacements):
    content_tokens = {'{{CONTENT}}', '{{content}}', '[[CONTENT]]', '__CONTENT__'}
    content_replaced = False

    def _clear_para_shading(paragraph):
        """إزالة التظليل الداكن من الفقرة المستبدلة (يمنع اختفاء النص على الجوال)"""
        try:
            from docx.oxml import OxmlElement as _OE2
            from docx.oxml.ns import qn as _qn2
            _W2 = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
            pPr = paragraph._p.get_or_add_pPr()
            for s in pPr.findall(_qn2('w:shd')): pPr.remove(s)
            shd = _OE2('w:shd')
            shd.set(_qn2('w:val'), 'clear')
            shd.set(_qn2('w:color'), 'auto')
            shd.set(_qn2('w:fill'), 'FFFFFF')
            rPr_p = pPr.find(f'{{{_W2}}}rPr')
            if rPr_p is not None:
                rPr_p.addprevious(shd)
            else:
                pPr.append(shd)
        except Exception: pass

    def _replace_in_paragraph(paragraph):
        nonlocal content_replaced
        before = paragraph.text or ''
        after = before
        for token, value in replacements.items():
            if token in after:
                after = after.replace(token, value)
        if after != before:
            if any(token in before for token in content_tokens):
                content_replaced = True
            paragraph.text = after
            _set_paragraph_runs_color(paragraph, 0, 0, 0)
            _clear_para_shading(paragraph)

    for paragraph in document.paragraphs:
        _replace_in_paragraph(paragraph)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_paragraph(paragraph)

    return content_replaced


def _get_latest_signed_assets_for_docx(doc_id, conn):
    doc_id_int = _safe_int(doc_id, 0)
    if doc_id_int <= 0:
        return {}

    sign_row = None
    try:
        sign_row = conn.execute(
            """
            SELECT sign_type, signature_owner_id, stamp_owner_id,
                   signature_asset_id, stamp_asset_id,
                 requested_from,
                   serial_number, signed_at,
                   positions_json, include_qr
            FROM signature_requests
            WHERE document_id=? AND status='signed'
            ORDER BY id DESC
            LIMIT 1
            """,
            (doc_id_int,)
        ).fetchone()
    except Exception:
        sign_row = conn.execute(
            """
            SELECT sign_type, signature_owner_id, stamp_owner_id,
                 requested_from,
                   serial_number, signed_at
            FROM signature_requests
            WHERE document_id=? AND status='signed'
            ORDER BY id DESC
            LIMIT 1
            """,
            (doc_id_int,)
        ).fetchone()

    if not sign_row:
        return {}

    doc_row = conn.execute(
        "SELECT archive_number FROM documents WHERE id=? LIMIT 1",
        (doc_id_int,)
    ).fetchone()
    stamp_reference_number = str(doc_row['archive_number'] or '').strip() if doc_row else ''

    sign_type = str(sign_row['sign_type'] or 'signature').strip().lower()
    if sign_type not in ('signature', 'stamp', 'both'):
        sign_type = 'signature'

    requested_from_id = _safe_int(sign_row['requested_from'], 0) if 'requested_from' in sign_row.keys() else 0
    signature_owner_id = _safe_int(sign_row['signature_owner_id'], 0) or requested_from_id
    stamp_owner_id = _safe_int(sign_row['stamp_owner_id'], 0) or requested_from_id

    signature_asset_id = _safe_int(sign_row['signature_asset_id'], 0) if 'signature_asset_id' in sign_row.keys() else 0
    stamp_asset_id = _safe_int(sign_row['stamp_asset_id'], 0) if 'stamp_asset_id' in sign_row.keys() else 0
    force_text_stamp = stamp_asset_id < 0
    include_qr = bool(_safe_int(sign_row['include_qr'], 1)) if 'include_qr' in sign_row.keys() else True
    positions = _normalize_sign_positions(
        sign_row['positions_json'] if 'positions_json' in sign_row.keys() else None,
        include_qr=include_qr
    )

    has_sig_position = any((item.get('type') == 'sig') for item in positions)
    has_stamp_position = any((item.get('type') == 'stamp') for item in positions)
    if sign_type == 'signature' and has_stamp_position:
        sign_type = 'both'
    if sign_type == 'stamp' and has_sig_position:
        sign_type = 'both'

    signature_path = None
    stamp_path = None
    stamp_text = None

    if sign_type in ('signature', 'both') and (signature_owner_id > 0 or signature_asset_id > 0):
        sig_row = None
        if signature_asset_id > 0:
            sig_row = conn.execute(
                "SELECT signature_path FROM signature_assets WHERE id=? AND is_active=1 LIMIT 1",
                (signature_asset_id,)
            ).fetchone()
        if not sig_row and signature_owner_id > 0:
            sig_row = conn.execute(
                """
                SELECT signature_path
                FROM signature_assets
                WHERE user_id=? AND is_active=1 AND COALESCE(signature_path,'')!=''
                ORDER BY id DESC LIMIT 1
                """,
                (signature_owner_id,)
            ).fetchone()
        if sig_row:
            signature_path = _resolve_document_file_path(sig_row['signature_path'])

    if sign_type in ('stamp', 'both') and (stamp_owner_id > 0 or stamp_asset_id > 0 or force_text_stamp):
        stamp_row = None
        if not force_text_stamp:
            if stamp_asset_id > 0:
                stamp_row = conn.execute(
                    "SELECT stamp_path FROM stamp_assets WHERE id=? AND is_active=1 LIMIT 1",
                    (stamp_asset_id,)
                ).fetchone()
            if not stamp_row and stamp_owner_id > 0:
                stamp_row = conn.execute(
                    """
                    SELECT stamp_path
                    FROM stamp_assets
                    WHERE user_id=? AND is_active=1 AND COALESCE(stamp_path,'')!=''
                    ORDER BY id DESC LIMIT 1
                    """,
                    (stamp_owner_id,)
                ).fetchone()

        if stamp_row:
            stamp_path = _resolve_document_file_path(stamp_row['stamp_path'])

        if not stamp_path and stamp_owner_id > 0:
            user_row = conn.execute(
                """
                SELECT u.name, u.stamp_text, u.role, u.employee_id,
                       s.section_code
                FROM users u
                LEFT JOIN archive_sections s ON s.id = u.archive_section_id
                WHERE u.id=? LIMIT 1
                """,
                (stamp_owner_id,)
            ).fetchone()
            if user_row:
                u_role = str(user_row['role'] or '').lower()
                if u_role in ('manager', 'admin', 'sys_admin'):
                    sec  = str(user_row['section_code'] or '').strip().upper()
                    emp  = str(user_row['employee_id']  or '').strip()
                    ser  = stamp_reference_number or (_short_signature_serial(sign_row['serial_number']) if sign_row['serial_number'] else '')
                    tpl_row = _get_default_stamp_tpl(conn)
                    if tpl_row:
                        stamp_path = _render_stamp_from_tpl(
                            tpl_row['file_name'], tpl_row['text_x_ratio'], tpl_row['text_y_ratio'],
                            sec, emp, ser, as_file=True
                        )
                    # Enforce default template if available; only fallback when no default template exists.
                    if not stamp_path and not tpl_row:
                        parts = [p for p in [sec, emp, ser] if p]
                        stamp_text = ' - '.join(parts)
                        if not stamp_text:
                            stamp_text = str(user_row['stamp_text'] or user_row['name'] or 'مدير')
                elif (user_row['stamp_text'] or '').strip():
                    stamp_text = f"{user_row['stamp_text']} — {user_row['name'] or ''}".strip(' —')

    return {
        'sign_type': sign_type,
        'serial_number': sign_row['serial_number'],
        'serial_short': _short_signature_serial(sign_row['serial_number']),
        'stamp_reference': stamp_reference_number or _short_signature_serial(sign_row['serial_number']),
        'signed_at': sign_row['signed_at'],
        'signature_path': signature_path,
        'stamp_path': stamp_path,
        'stamp_text': stamp_text,
        'positions': positions,
        'include_qr': include_qr,
    }


def _build_text_document_docx_file(doc, conn=None):
    if not DOCX_AVAILABLE:
        return None

    own_conn = conn is None
    db = conn or get_db()
    temp_assets = []
    try:
        template_name = str(doc.get('template_name') or '').strip()
        template_path = _resolve_docx_template_path(template_name) if template_name else None

        try:
            document = DocxDocument(template_path) if template_path else DocxDocument()
        except Exception:
            document = DocxDocument()
        # ─── تثبيت ألوان الأنماط منذ البداية ──────────────
        _W_FIX = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        try:
            from docx.oxml.ns import qn as _qf
            from docx.oxml import OxmlElement as _OEf
            for _st in document.styles:
                try:
                    _el = _st.element
                    # إزالة التظليل من خصائص الفقرة في النمط (pPr/w:shd)
                    _pp = _el.find(f'{{{_W_FIX}}}pPr')
                    if _pp is not None:
                        for _s in _pp.findall(_qf('w:shd')): _pp.remove(_s)
                        _sd = _OEf('w:shd')
                        _sd.set(_qf('w:val'), 'clear')
                        _sd.set(_qf('w:color'), 'auto')
                        _sd.set(_qf('w:fill'), 'FFFFFF')
                        # إدراج shd قبل rPr لاحترام ترتيب OOXML
                        _rp_in_pp = _pp.find(f'{{{_W_FIX}}}rPr')
                        if _rp_in_pp is not None:
                            _rp_in_pp.addprevious(_sd)
                        else:
                            _pp.append(_sd)
                    _rp = _el.find(f'{{{_W_FIX}}}rPr')
                    if _rp is None: _rp = _OEf('w:rPr'); _el.append(_rp)
                    for _c in _rp.findall(_qf('w:color')): _rp.remove(_c)
                    for _sh in _rp.findall(_qf('w:shd')): _rp.remove(_sh)
                    for _hl in _rp.findall(_qf('w:highlight')): _rp.remove(_hl)
                    _ce = _OEf('w:color'); _ce.set(_qf('w:val'), '000000'); _rp.append(_ce)
                    _rf = _rp.find(_qf('w:rFonts'))
                    if _rf is not None: _rf.set(_qf('w:cs'), 'Arial')
                except Exception: pass
            # docDefaults is in styles.xml - access via styles.element
            _styles_el = document.styles.element
            _dd = _styles_el.find(f'.//{{{_W_FIX}}}docDefaults')
            if _dd is not None:
                # إزالة التظليل من docDefaults/pPr
                _dp = _dd.find(f'.//{{{_W_FIX}}}pPr')
                if _dp is not None:
                    for _s in _dp.findall(_qf('w:shd')): _dp.remove(_s)
                _dr = _dd.find(f'.//{{{_W_FIX}}}rPr')
                if _dr is not None:
                    for _c in list(_dr.findall(_qf('w:color'))): _dr.remove(_c)
                    for _s in list(_dr.findall(_qf('w:shd'))): _dr.remove(_s)
                    for _h in list(_dr.findall(_qf('w:highlight'))): _dr.remove(_h)
                    _ce = _OEf('w:color'); _ce.set(_qf('w:val'), '000000'); _dr.insert(0, _ce)
                    _rf = _dr.find(_qf('w:rFonts'))
                    if _rf is not None:
                        # Remove all theme font attrs, set explicit Arial
                        for _ak in list(_rf.attrib.keys()):
                            if 'Theme' in _ak or 'theme' in _ak:
                                del _rf.attrib[_ak]
                        _rf.set(_qf('w:ascii'), 'Arial')
                        _rf.set(_qf('w:hAnsi'), 'Arial')
                        _rf.set(_qf('w:cs'), 'Arial')
            # فرض خلفية بيضاء صريحة على الصفحة (يمنع الخلفية السوداء من الكليشة)
            try:
                _doc_el = document.element
                for _bg in _doc_el.findall(_qf('w:background')): _doc_el.remove(_bg)
                _new_bg = _OEf('w:background')
                _new_bg.set(_qf('w:color'), 'FFFFFF')
                _doc_el.insert(0, _new_bg)
            except Exception: pass
            # إزالة displayBackgroundShape من settings لمنع الخلفية الداكنة
            try:
                _set_el = document.settings.element
                _W_SET = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                for _dbs in list(_set_el.findall(f'{{{_W_SET}}}displayBackgroundShape')):
                    _set_el.remove(_dbs)
            except Exception: pass
        except Exception: pass


        content_html = _extract_content_html(doc.get('content_json') or '')
        content_lines = _html_to_text_lines(content_html)
        content_text = '\n'.join(content_lines)
        created_at_str = str(doc.get('created_at') or '').replace('T', ' ')[:16]

        # ─── جرب استبدال {{CONTENT}} ───────────────────────────
        replacements = {
            '{{CONTENT}}': content_text,
            '{{content}}': content_text,
            '[[CONTENT]]': content_text,
            '__CONTENT__': content_text,
            '{CONTENT}': content_text,
            '{{TITLE}}': str(doc.get('title') or ''),
            '{{ARCHIVE_NUMBER}}': str(doc.get('archive_number') or ''),
            '{{SECTION}}': str(doc.get('archive_section') or ''),
            '{{DATE}}': created_at_str,
        }
        content_replaced = _replace_docx_tokens(document, replacements)

        # ─── لو ما في token: اكتب النص في الجسم ──────────────
        if not content_replaced and content_lines:
            try:
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                from docx.shared import Pt, RGBColor
                from docx.oxml.ns import qn as _qn
                from docx.oxml import OxmlElement as _OxmlE

                _WI = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

                def _run_ex(run, pt=12):
                    try:
                        run.font.color.rgb = RGBColor(0, 0, 0)
                        run.font.size = Pt(pt)
                        rPr = run._r.get_or_add_rPr()
                        for c in rPr.findall(_qn('w:color')): rPr.remove(c)
                        ce = _OxmlE('w:color'); ce.set(_qn('w:val'), '000000'); rPr.insert(0, ce)
                        rFonts = rPr.find(_qn('w:rFonts'))
                        if rFonts is None: rFonts = _OxmlE('w:rFonts'); rPr.insert(0, rFonts)
                        rFonts.set(_qn('w:ascii'), 'Arial')
                        rFonts.set(_qn('w:hAnsi'), 'Arial')
                        rFonts.set(_qn('w:cs'), 'Arial')
                    except Exception: pass

                def _set_para_white_bg(para):
                    try:
                        pPr = para._p.get_or_add_pPr()
                        for s in pPr.findall(_qn('w:shd')): pPr.remove(s)
                        shd = _OxmlE('w:shd')
                        shd.set(_qn('w:val'), 'clear')
                        shd.set(_qn('w:color'), 'auto')
                        shd.set(_qn('w:fill'), 'FFFFFF')
                        # shd يجب أن يأتي قبل rPr في ترتيب OOXML
                        rPr_p = pPr.find(f'{{{_WI}}}rPr')
                        if rPr_p is not None:
                            rPr_p.addprevious(shd)
                        else:
                            pPr.append(shd)
                        if rPr_p is None: rPr_p = _OxmlE('w:rPr'); pPr.append(rPr_p)
                        for c in rPr_p.findall(_qn('w:color')): rPr_p.remove(c)
                        ce = _OxmlE('w:color'); ce.set(_qn('w:val'), '000000'); rPr_p.append(ce)
                    except Exception: pass

                meta_p = document.add_paragraph()
                meta_r = meta_p.add_run(
                    f"رقم الأرشيف: {doc.get('archive_number','—')}    |    {created_at_str}"
                )
                meta_r.bold = True
                _run_ex(meta_r, 10)
                meta_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                _set_para_white_bg(meta_p)

                document.add_paragraph()

                for line in content_lines:
                    p = document.add_paragraph()
                    run = p.add_run(line)
                    _run_ex(run, 12)
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    _set_para_white_bg(p)

                # ─── التوقيع (اختياري) ─────────────────────────────────
                signed = _get_latest_signed_assets_for_docx(doc.get('id'), db)
                if signed and signed.get('serial_number'):
                    serial = str(signed['serial_number'])
                    serial_short = str(_short_signature_serial(serial, section_code=doc.get('archive_section_code')))
                    signed_at = str(signed.get('signed_at') or '').replace('T', ' ')[:16] or '—'
                    positions = signed.get('positions') or []

                    document.add_paragraph('')
                    meta_signed = document.add_paragraph(f"✅ توقيع معتمد | الرقم: {serial} | التاريخ: {signed_at}")
                    meta_signed.alignment = _position_to_docx_alignment(0.9)
                    _set_paragraph_runs_color(meta_signed, 0, 0, 0)

                    anchor_paragraph = document.add_paragraph('')
                    visual_items = []

                    if signed.get('signature_path') and os.path.exists(signed['signature_path']):
                        visual_items.append({
                            'type': 'sig',
                            'label': 'التوقيع',
                            'image_path': signed['signature_path'],
                            'default_inches': 2.2,
                            'position': _pick_sign_position(positions, 'sig'),
                        })

                    if signed.get('stamp_path') and os.path.exists(signed['stamp_path']):
                        visual_items.append({
                            'type': 'stamp',
                            'label': 'الختم',
                            'image_path': signed['stamp_path'],
                            'default_inches': 1.4,
                            'position': _pick_sign_position(positions, 'stamp'),
                        })
                    elif signed.get('stamp_text'):
                        visual_items.append({
                            'type': 'stamp',
                            'label': 'الختم',
                            'text': signed['stamp_text'],
                            'default_inches': 1.4,
                            'position': _pick_sign_position(positions, 'stamp'),
                        })

                    if signed.get('include_qr'):
                        qr_seed = _build_qr_payload(
                            serial_text=str(serial or '').strip(),
                            archive_number=str(doc.get('archive_number') or '').strip(),
                        )
                        qr_temp = _save_qr_png_temp(qr_seed) if qr_seed else None
                        if qr_temp and os.path.exists(qr_temp):
                            temp_assets.append(qr_temp)
                            visual_items.append({
                                'type': 'qr',
                                'label': 'الباركود',
                                'image_path': qr_temp,
                                'default_inches': 1.3,
                                'position': _pick_sign_position(positions, 'qr'),
                            })

                    visual_items.sort(key=lambda item: _clamp_ratio((item.get('position') or {}).get('y'), 0.8))

                    previous_y_pt = 0.0
                    for item in visual_items:
                        position = item.get('position') or _pick_sign_position(positions, item['type'])
                        align = _position_to_docx_alignment(position.get('x', 0.7))

                        label_paragraph = document.add_paragraph()
                        label_paragraph.add_run(item.get('label', ''))
                        serial_run = label_paragraph.add_run(f"  #{serial_short}")
                        if align is not None:
                            label_paragraph.alignment = align
                        _set_paragraph_runs_color(label_paragraph, 0, 0, 0)

                        try:
                            from docx.shared import Pt, RGBColor
                            serial_run.font.size = Pt(8)
                            serial_run.font.color.rgb = RGBColor(0x33, 0x44, 0x66)
                        except Exception:
                            pass

                        try:
                            from docx.shared import Pt
                            target_pt = _clamp_ratio(position.get('y'), 0.8) * 460.0
                            label_paragraph.paragraph_format.space_before = Pt(max(0.0, target_pt - previous_y_pt))
                            previous_y_pt = target_pt
                        except Exception:
                            pass

                        if item.get('text'):
                            value_paragraph = document.add_paragraph(str(item['text']))
                            if align is not None:
                                value_paragraph.alignment = align
                            _set_paragraph_runs_color(value_paragraph, 0, 0, 0)
                            continue

                image_path = item.get('image_path')
                if image_path and os.path.exists(image_path):
                    added_anchor = _add_docx_floating_image(
                        document,
                        image_path,
                        position,
                        width_inches=_position_width_inches(position.get('w'), item.get('default_inches', 1.4)),
                        anchor_paragraph=anchor_paragraph
                    )
                    if not added_anchor:
                        image_paragraph = document.add_paragraph()
                        if align is not None:
                            image_paragraph.alignment = align
                        run = image_paragraph.add_run()
                        run.add_picture(
                            image_path,
                            width=Inches(_position_width_inches(position.get('w'), item.get('default_inches', 1.4)))
                        )
            except Exception:
                pass

        _force_docx_body_text_black(document)

        tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        tmp.close()
        document.save(tmp.name)
        return tmp.name

    except Exception as e:
        print(f"[DOCX] ERROR: {e}")
        return None
    finally:
        for asset_path in temp_assets:
            try:
                os.remove(asset_path)
            except OSError:
                pass
        if own_conn:
            db.close()

def _get_latest_signed_assets_preview(doc_id, conn=None):
    doc_id_int = _safe_int(doc_id, 0)
    if doc_id_int <= 0:
        print(f"[WARNING] Invalid doc_id for signed assets: {doc_id}", file=__import__('sys').stderr)
        return {}

    own_conn = conn is None
    db = conn or get_db()
    try:
        sign_row = None
        try:
            sign_row = db.execute(
                """
                SELECT sign_type, signature_owner_id, stamp_owner_id,
                       signature_asset_id, stamp_asset_id,
                      requested_from,
                       serial_number, signed_at,
                       positions_json, include_qr
                FROM signature_requests
                WHERE document_id=? AND status='signed'
                ORDER BY id DESC
                LIMIT 1
                """,
                (doc_id_int,)
            ).fetchone()
        except Exception:
            sign_row = db.execute(
                """
                SELECT sign_type, signature_owner_id, stamp_owner_id,
                      requested_from,
                       serial_number, signed_at
                FROM signature_requests
                WHERE document_id=? AND status='signed'
                ORDER BY id DESC
                LIMIT 1
                """,
                (doc_id_int,)
            ).fetchone()

        if not sign_row:
            print(f"[INFO] No signed record found for doc {doc_id_int}", file=__import__('sys').stderr)
            return {}

        doc_row = db.execute(
            "SELECT archive_number FROM documents WHERE id=? LIMIT 1",
            (doc_id_int,)
        ).fetchone()
        stamp_reference_number = str(doc_row['archive_number'] or '').strip() if doc_row else ''

        sign_type = str(sign_row['sign_type'] or 'signature').strip().lower()
        if sign_type not in ('signature', 'stamp', 'both'):
            sign_type = 'signature'

        signature_data_uri = None
        stamp_data_uri = None
        include_qr = bool(_safe_int(sign_row['include_qr'], 1)) if 'include_qr' in sign_row.keys() else True
        positions = _normalize_sign_positions(
            sign_row['positions_json'] if 'positions_json' in sign_row.keys() else None,
            include_qr=include_qr
        )

        signature_asset_id = _safe_int(sign_row['signature_asset_id'], 0) if 'signature_asset_id' in sign_row.keys() else 0
        stamp_asset_id = _safe_int(sign_row['stamp_asset_id'], 0) if 'stamp_asset_id' in sign_row.keys() else 0
        force_text_stamp = stamp_asset_id < 0
        requested_from_id = _safe_int(sign_row['requested_from'], 0) if 'requested_from' in sign_row.keys() else 0

        signature_owner_id = _safe_int(sign_row['signature_owner_id'], 0) or requested_from_id
        if sign_type in ('signature', 'both') and (signature_owner_id > 0 or signature_asset_id > 0):
            sig_row = None
            if signature_asset_id > 0:
                sig_row = db.execute(
                    "SELECT signature_path FROM signature_assets WHERE id=? AND is_active=1 LIMIT 1",
                    (signature_asset_id,)
                ).fetchone()
            if not sig_row:
                sig_row = db.execute(
                    """
                    SELECT signature_path
                    FROM signature_assets
                    WHERE user_id=? AND is_active=1 AND COALESCE(signature_path,'')!=''
                    ORDER BY id DESC
                    LIMIT 1
                    """,
                    (signature_owner_id,)
                ).fetchone()
            if sig_row:
                signature_data_uri = _path_to_data_uri(sig_row['signature_path'])

        stamp_owner_id = _safe_int(sign_row['stamp_owner_id'], 0) or requested_from_id
        if sign_type in ('stamp', 'both') and (stamp_owner_id > 0 or stamp_asset_id > 0 or force_text_stamp):
            stamp_row = None
            if not force_text_stamp:
                if stamp_asset_id > 0:
                    stamp_row = db.execute(
                        "SELECT stamp_path FROM stamp_assets WHERE id=? AND is_active=1 LIMIT 1",
                        (stamp_asset_id,)
                    ).fetchone()
                if not stamp_row:
                    stamp_row = db.execute(
                        """
                        SELECT stamp_path
                        FROM stamp_assets
                        WHERE user_id=? AND is_active=1 AND COALESCE(stamp_path,'')!=''
                        ORDER BY id DESC
                        LIMIT 1
                        """,
                        (stamp_owner_id,)
                    ).fetchone()

            # Always load user info to decide rendering strategy
            user_row = db.execute(
                """
                SELECT u.name, u.stamp_text, u.role, u.employee_id,
                       s.section_code
                FROM users u
                LEFT JOIN archive_sections s ON s.id = u.archive_section_id
                WHERE u.id=? LIMIT 1
                """,
                (stamp_owner_id,)
            ).fetchone()

            u_role = str(user_row['role'] or '').lower() if user_row else ''
            if user_row and u_role in ('manager', 'admin', 'sys_admin'):
                # Managers: render from template using archive number as stamp reference.
                serial_val = sign_row['serial_number'] if sign_row['serial_number'] else ''
                ser_short = stamp_reference_number or (_short_signature_serial(serial_val) if serial_val else '')
                tpl_row = _get_default_stamp_tpl(db)
                if tpl_row:
                    stamp_data_uri = _render_stamp_from_tpl(
                        tpl_row['file_name'], tpl_row['text_x_ratio'], tpl_row['text_y_ratio'],
                        user_row['section_code'] or '', user_row['employee_id'] or '', ser_short
                    )
                # Enforce default template if available; fallback to pre-generated PNG only when no template is configured.
                if not stamp_data_uri and stamp_row and not tpl_row:
                    stamp_data_uri = _path_to_data_uri(stamp_row['stamp_path'])
            else:
                # Non-managers: use pre-generated PNG stamp
                if stamp_row:
                    stamp_data_uri = _path_to_data_uri(stamp_row['stamp_path'])
                if not stamp_data_uri and user_row and (user_row['stamp_text'] or '').strip():
                    stamp_data_uri = _stamp_text_to_data_uri(user_row['stamp_text'], user_row['name'])

        qr_data_uri = None
        if include_qr:
            qr_payload = _build_qr_payload(
                serial_text=str(sign_row['serial_number'] or '').strip(),
                archive_number=stamp_reference_number,
            )
            if qr_payload:
                qr_data_uri = _qr_png_data_uri(qr_payload)

        return {
            "sign_type": sign_type,
            "serial_number": sign_row['serial_number'],
            "serial_short": _short_signature_serial(sign_row['serial_number']),
            "stamp_reference": stamp_reference_number or _short_signature_serial(sign_row['serial_number']),
            "archive_number": stamp_reference_number,
            "signed_at": sign_row['signed_at'],
            "signature_data_uri": signature_data_uri,
            "stamp_data_uri": stamp_data_uri,
            "positions": positions,
            "include_qr": include_qr,
            "qr_data_uri": qr_data_uri,
        }
    finally:
        if own_conn:
            db.close()


def _build_text_document_archive_markup(doc, conn=None):
    content_html = _extract_content_html(doc.get('content_json'))

    template_name = str(doc.get('template_name') or '').strip()
    template_image = None
    template_lines = []
    template_body_lines = []
    template_path = _resolve_docx_template_path(template_name)
    if template_path:
        try:
            template_preview = _build_docx_template_preview(template_path)
        except Exception:
            template_preview = {}

        preview_images = template_preview.get('header_images') or []
        template_image = preview_images[0] if preview_images else None
        template_lines = (template_preview.get('header_lines') or [])[:2]
        template_body_lines = (template_preview.get('body_lines') or [])[:12]

    if not content_html:
        if template_body_lines:
            content_html = ''.join(f"<p>{html.escape(line)}</p>" for line in template_body_lines)
        else:
            content_html = '<p>—</p>'

    title = html.escape(str(doc.get('title') or 'وثيقة'))
    archive_number = html.escape(str(doc.get('archive_number') or '—'))
    section = html.escape(str(doc.get('archive_section') or 'عام'))
    section_code = html.escape(str(doc.get('archive_section_code') or 'GN'))
    created_at = html.escape(str(doc.get('created_at') or datetime.utcnow().isoformat()))
    template_meta = html.escape(template_name) if template_name else 'بدون كليشة'

    signed_assets = _get_latest_signed_assets_preview(doc.get('id'), conn=conn)
    signed_serial = html.escape(str(signed_assets.get('serial_number') or '—'))
    signed_serial_short = html.escape(str(_short_signature_serial(signed_assets.get('serial_number'), section_code=doc.get('archive_section_code'))))
    signed_stamp_reference = html.escape(str(signed_assets.get('stamp_reference') or doc.get('archive_number') or signed_serial_short).strip())
    signed_at = html.escape(str(signed_assets.get('signed_at') or '').replace('T', ' ')[:16] or '—')

    template_banner = f"<div class='tpl-banner'><img src='{template_image}' alt='قالب DOCX'></div>" if template_image else ""
    template_hint = ""
    if template_lines:
        template_hint = "<div class='tpl-lines'>" + " • ".join(html.escape(line) for line in template_lines) + "</div>"

    overlay_items = []
    sign_block = ""
    if signed_assets:
        signature_img = signed_assets.get('signature_data_uri')
        stamp_img = signed_assets.get('stamp_data_uri')
        qr_img = signed_assets.get('qr_data_uri') if signed_assets.get('include_qr') else None
        positions = signed_assets.get('positions') or []

        def _build_overlay_item(src, item_type, alt_label):
            if not src:
                return None
            pos = _pick_sign_position(positions, item_type)
            x_val = _clamp_ratio(pos.get('x'), 0.5)
            y_val = _clamp_ratio(pos.get('y'), 0.8)
            w_val = max(0.08, _clamp_ratio(pos.get('w'), 0.15))
            h_val = max(0.08, _clamp_ratio(pos.get('h'), 0.12))

            if x_val + w_val > 1.0:
                x_val = max(0.0, 1.0 - w_val)
            if y_val + h_val > 1.0:
                y_val = max(0.0, 1.0 - h_val)

            return (
                f"<div class='ov-item ov-{item_type}' style='left:{x_val * 100:.2f}%;top:{y_val * 100:.2f}%;"
                f"width:{w_val * 100:.2f}%;height:{h_val * 100:.2f}%;'>"
                f"<img src='{src}' alt='{html.escape(alt_label)}'></div>"
            )

        for overlay in (
            _build_overlay_item(signature_img, 'sig', 'التوقيع'),
            _build_overlay_item(stamp_img, 'stamp', 'الختم'),
            _build_overlay_item(qr_img, 'qr', 'الباركود'),
        ):
            if overlay:
                overlay_items.append(overlay)

        if signature_img or stamp_img or qr_img:
            sign_block = f"""
            <div class='signed-block'>
                <div class='signed-meta'>توقيع معتمد • الرقم التسلسلي: {signed_serial} • التاريخ: {signed_at}</div>
                <div class='sig-grid'>
                    {f"<div class='sig-item'><div class='sig-lbl'>التوقيع <span class='sig-mini'>#{signed_serial_short}</span></div><img src='{signature_img}' alt='signature'></div>" if signature_img else ''}
                    {f"<div class='sig-item'><div class='sig-lbl'>الختم <span class='sig-mini'>{signed_stamp_reference}</span></div><img src='{stamp_img}' alt='stamp'></div>" if stamp_img else ''}
                    {f"<div class='sig-item'><div class='sig-lbl'>الباركود <span class='sig-mini'>#{signed_serial_short}</span></div><img src='{qr_img}' alt='qr'></div>" if qr_img else ''}
                </div>
            </div>
            """

    if overlay_items:
        content_markup = (
            f"<div class='content content-with-sign'><div class='content-body'>{content_html}</div>"
            f"<div class='sign-overlay'>{''.join(overlay_items)}</div></div>"
        )
    else:
        content_markup = f"<div class='content'>{content_html}</div>"

    return f"""
    <!DOCTYPE html>
    <html lang='ar' dir='rtl'>
    <head>
        <meta charset='UTF-8'>
        <meta name='viewport' content='width=device-width, initial-scale=1.0'>
        <title>{title}</title>
        <style>
            body{{font-family:'IBM Plex Sans Arabic',Arial,sans-serif;background:#fff;color:#000;line-height:1.9;padding:28px}}
            .head{{border-bottom:2px solid #e5e7eb;margin-bottom:16px;padding-bottom:12px}}
            .title{{font-size:22px;font-weight:700;margin-bottom:6px;color:#0f172a}}
            .meta{{font-size:12px;color:#1f2937;display:flex;gap:14px;flex-wrap:wrap}}
            .content{{font-size:16px;color:#000;direction:rtl;text-align:right}}
            .content p,.content div,.content span,.content h1,.content h2,.content h3,.content li{{color:#000 !important;direction:rtl;text-align:right}}
            .content-with-sign{{position:relative;min-height:320px;padding-bottom:8px}}
            .content-body{{position:relative;z-index:1}}
            .sign-overlay{{position:absolute;inset:0;pointer-events:none;z-index:2}}
            .ov-item{{position:absolute;display:flex;align-items:center;justify-content:center;opacity:.98}}
            .ov-item img{{max-width:100%;max-height:100%;object-fit:contain;display:block}}
            .tpl-banner{{margin:0 0 14px 0;border:1px solid #e5e7eb;border-radius:8px;overflow:hidden;background:#fff}}
            .tpl-banner img{{display:block;width:100%;max-height:180px;object-fit:contain}}
            .tpl-lines{{font-size:12px;color:#334155;margin-bottom:10px}}
            .signed-block{{margin-top:18px;border-top:1px dashed #cbd5e1;padding-top:12px}}
            .signed-meta{{font-size:12px;color:#334155;margin-bottom:8px}}
            .sig-grid{{display:flex;gap:12px;flex-wrap:wrap}}
            .sig-item{{border:1px solid #e5e7eb;border-radius:8px;padding:8px 10px;min-width:170px}}
            .sig-lbl{{font-size:11px;color:#64748b;margin-bottom:4px}}
            .sig-mini{{font-size:9px;color:#475569;margin-right:4px;font-family:monospace}}
            .sig-item img{{max-width:180px;max-height:90px;object-fit:contain;display:block}}
        </style>
    </head>
    <body>
        {template_banner}
        <div class='head'>
            <div class='title'>{title}</div>
            <div class='meta'>
                <span>رقم الأرشيف: {archive_number}</span>
                <span>القسم: {section} ({section_code})</span>
                <span>تاريخ الإنشاء: {created_at}</span>
                <span>الكليشة: {template_meta}</span>
            </div>
        </div>
        {template_hint}
        {content_markup}
        {sign_block}
    </body>
    </html>
    """.strip()


def _build_doc_with_pdf_template(doc, conn=None):
    """
    Generate a PDF document by overlaying content on a PDF template.
    Overlays:
      - رقم الأرشيف + التاريخ/الوقت at configured top positions
      - Document text content in the middle content area
      - Signature / stamp images at selected positions (if signed)
      - QR barcode at configured bottom position
    Returns: temp .pdf file path, or None on failure.
    """
    doc_id = doc.get('id', 'UNKNOWN')
    try:
        from reportlab.pdfgen import canvas as rl_canvas
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.lib.utils import ImageReader
        import arabic_reshaper
        from bidi.algorithm import get_display
        from pypdf import PdfReader, PdfWriter
    except ImportError as e:
        print(f"[ERROR] Import failed for doc {doc_id}: {str(e)}", file=__import__('sys').stderr)
        return None

    template_name = str(doc.get('template_name') or '').strip()
    if not template_name.lower().endswith('.pdf'):
        print(f"[INFO] Doc {doc_id} template not PDF: {template_name}", file=__import__('sys').stderr)
        return None

    template_path = _resolve_pdf_template_file(template_name)
    if not template_path or not os.path.exists(template_path):
        print(f"[ERROR] Template not found for doc {doc_id}: {template_name} -> {template_path}", file=__import__('sys').stderr)
        return None

    config = _get_pdf_template_config(template_name)
    print(f"[INFO] Building PDF for doc {doc_id} with template {template_name}", file=__import__('sys').stderr)

    # ── Page dimensions ──────────────────────────────────────────────
    try:
        _rdr = PdfReader(template_path)
        _pg0 = _rdr.pages[0]
        pw = float(_pg0.mediabox.width)
        ph = float(_pg0.mediabox.height)
    except Exception:
        pw, ph = 595.2, 842.88

    # ── Register Almadinah fonts (Arabic 1) ───────────────
    _FONT_AR = 'AlmadinArabic'      # Font 1 for Arabic
    _FONT_EN = 'Helvetica'          # Use Helvetica for English (better support)
    _FONT_FALLBACK = 'Helvetica'
    
    # Arabic font candidates (Almadinah1 priority)
    _arabic_font_candidates = [
        (_FONT_AR, os.path.join(app.root_path, 'almadinah-font', 'Almadinah1.otf')),
        (_FONT_AR, r'D:\PythonProject2\almadinah-font\Almadinah1.otf'),
        ('TplAr', r'C:\Windows\Fonts\tahoma.ttf'),
        ('TplAr', r'C:\Windows\Fonts\arial.ttf'),
        ('TplAr', '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf'),
        ('TplAr', '/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf'),
        ('TplAr', '/usr/share/fonts/truetype/noto/NotoSansArabic-Regular.ttf'),
        ('TplAr', os.path.join(app.root_path, 'static', 'fonts', 'NotoSansArabic-Regular.ttf')),
    ]
    
    # Register Arabic font
    for _fn, _fpath in _arabic_font_candidates:
        if os.path.exists(_fpath):
            try:
                pdfmetrics.registerFont(TTFont(_fn, _fpath))
                _FONT_AR = _fn
                print(f"[FONTS] ✓ Registered Arabic font: {_fn}", file=__import__('sys').stderr)
                break
            except Exception as _e:
                print(f"[FONTS] Failed to register {_fn}: {_e}", file=__import__('sys').stderr)
                continue
    
    # Fallback if fonts not registered
    if _FONT_AR not in pdfmetrics._fonts:
        _FONT_AR = 'Helvetica'
        print(f"[FONTS] Arabic font not found, using Helvetica fallback", file=__import__('sys').stderr)

    def _is_arabic_text(text):
        """Detect if text contains Arabic characters."""
        arabic_range = range(0x0600, 0x06FF + 1)
        for char in str(text or ''):
            if ord(char) in arabic_range:
                return True
        return False

    def _get_font_for_text(text):
        """Return appropriate font based on text language."""
        if _is_arabic_text(text):
            return _FONT_AR
        return _FONT_EN

    def _reshape(text):
        try:
            import arabic_reshaper as _ar
            from bidi.algorithm import get_display as _gd
            return _gd(_ar.reshape(str(text or '')))
        except Exception:
            return str(text or '')

    # ── Document metadata ────────────────────────────────────────────
    archive_number = str(doc.get('archive_number') or '—')
    created_at_raw = str(doc.get('created_at') or '').replace('T', ' ')
    # Convert UTC stored time to local time
    try:
        import datetime as _dt
        _utc_naive = _dt.datetime.strptime(created_at_raw[:16], '%Y-%m-%d %H:%M')
        _local = _utc_naive + _dt.timedelta(hours=3)   # AST = UTC+3
        date_part = _local.strftime('%Y-%m-%d')
        time_part = _local.strftime('%H:%M')
    except Exception:
        date_part = created_at_raw[:10] if len(created_at_raw) >= 10 else '—'
        time_part = created_at_raw[11:16] if len(created_at_raw) >= 16 else ''
    content_html = _extract_content_html(doc.get('content_json') or '')
    content_lines = _html_to_text_lines(content_html)

    # ── Signed assets ────────────────────────────────────────────────
    signed = _get_latest_signed_assets_preview(doc.get('id'), conn=conn)
    signed_positions = signed.get('positions') if signed else []
    serial_short = _short_signature_serial(
        signed.get('serial_number') if signed else None,
        section_code=doc.get('archive_section_code')
    ) if signed else ''
    stamp_reference = str((signed.get('stamp_reference') if signed else '') or archive_number or '').strip()

    # ── Build overlay PDF ────────────────────────────────────────────
    overlay_buf = io.BytesIO()
    c = rl_canvas.Canvas(overlay_buf, pagesize=(pw, ph))
    c.setFillColorRGB(0.05, 0.05, 0.05)

    def _draw_right(text, x_frac, y_top_frac, font_size):
        reshaped = _reshape(text)
        x_pt = float(x_frac) * pw
        y_pt = ph - (float(y_top_frac) * ph)
        # Select font based on text language
        font_to_use = _get_font_for_text(text)
        # Fallback to Helvetica if font not available
        if font_to_use not in pdfmetrics._fonts:
            font_to_use = 'Helvetica'
        c.setFont(font_to_use, font_size)
        try:
            tw = c.stringWidth(reshaped, font_to_use, font_size)
        except Exception:
            tw = 0
        # clamp left so text doesn't overflow left margin
        if x_pt - tw < 0.04 * pw:
            x_pt = tw + 0.04 * pw
        c.drawRightString(x_pt, y_pt, reshaped)

    an_cfg  = config.get('archive_number', {})
    dt_cfg  = config.get('date', {})
    ct_cfg  = config.get('content', {})
    qr_cfg  = config.get('qr', {})

    def _draw_page_header():
        # رقم الأرشيف — بدون ملصق لأن المكان مخصص في الكليشة
        _draw_right(
            archive_number,
            an_cfg.get('x', 0.92), an_cfg.get('y', 0.068),
            int(an_cfg.get('font_size', 11))
        )

        # التاريخ والوقت — بدون ملصق
        date_label = date_part
        if time_part:
            date_label += f"   {time_part}"
        _draw_right(
            date_label,
            dt_cfg.get('x', 0.92), dt_cfg.get('y', 0.112),
            int(dt_cfg.get('font_size', 10))
        )

    # النص الرئيسي — word-wrap صحيح
    x_left_ratio  = _clamp_ratio(ct_cfg.get('x_left', 0.08), 0.08)
    x_right_ratio = _clamp_ratio(ct_cfg.get('x_right', 0.92), 0.92)
    allow_narrow_content = str(ct_cfg.get('allow_narrow_content', '')).strip().lower() in ('1', 'true', 'yes')
    # إذا كانت مساحة النص ضيقة جداً، نوسعها افتراضياً لتقليل تقطيع الجُمل.
    if x_right_ratio <= x_left_ratio + 0.10:
        x_right_ratio = min(0.94, x_left_ratio + 0.74)
    if (x_right_ratio - x_left_ratio) < 0.66 and not allow_narrow_content:
        x_right_ratio = min(0.94, x_left_ratio + 0.74)

    x_right_pt  = x_right_ratio * pw
    x_left_pt   = x_left_ratio * pw                                # هامش يسار
    y_frac      = float(ct_cfg.get('y_start', 0.22))
    y_end_frac  = float(ct_cfg.get('y_end',   0.83))
    fs          = int(ct_cfg.get('font_size',  12))
    lh_pt       = float(ct_cfg.get('line_height', 20))   # بوجهة ثابتة

    # بعض القوالب تضبط مساحة النص قصيرة جداً (مثل y_end=0.52) مما يسبب تقطيعاً مبكراً.
    if y_end_frac <= y_frac + 0.34:
        y_end_frac = min(0.88, y_frac + 0.62)

    # منع line-height الكبير جداً من استهلاك الصفحة بسرعة.
    if lh_pt > 24:
        lh_pt = 20.0
    elif lh_pt < 14:
        lh_pt = 14.0

    max_width   = x_right_pt - x_left_pt                  # عرض متاح

    def _wrap_line(text, max_w):
        """Break a single logical line into rendered sub-lines that fit max_w.
        Wrapping is done on original text (pre-reshape) so Arabic words split correctly,
        then each sub-line is reshaped individually before width measurement."""
        words = text.split()
        if not words:
            return ['']
        # Select font based on text language
        font_to_use = _get_font_for_text(text)
        # Fallback to Helvetica if font not available
        if font_to_use not in pdfmetrics._fonts:
            font_to_use = 'Helvetica'
        c.setFont(font_to_use, fs)
        sub_lines = []
        current_words = []
        for word in words:
            test_words = current_words + [word]
            test_str = ' '.join(test_words)
            try:
                w = c.stringWidth(_reshape(test_str), font_to_use, fs)
            except Exception:
                w = len(test_str) * fs * 0.6
            if w <= max_w or not current_words:
                current_words = test_words
            else:
                sub_lines.append(' '.join(current_words))
                current_words = [word]
        if current_words:
            sub_lines.append(' '.join(current_words))
        return sub_lines if sub_lines else ['']

    # فراغ بسيط فقط؛ الفراغ الكبير كان يستهلك المساحة ويقصّ النص.
    para_gap_pt = lh_pt * 0.10

    # جهّز الكتل النصية ووزّعها على صفحات تلقائياً عند امتلاء المساحة
    page_blocks = [[]]
    current_y = y_frac
    line_step = lh_pt / ph
    para_step = para_gap_pt / ph

    for logical_line in content_lines:
        sub_lines = _wrap_line(logical_line, max_width)
        if not sub_lines:
            sub_lines = ['']

        required = (len(sub_lines) * line_step) + para_step
        if page_blocks[-1] and (current_y + required) > y_end_frac:
            page_blocks.append([])
            current_y = y_frac

        page_blocks[-1].append((logical_line, sub_lines))
        current_y += required

    # في حال عدم وجود محتوى، أبق صفحة واحدة حتى لا تتأثر مراحل الدمج
    if not page_blocks:
        page_blocks = [[]]

    # QR code — يُجهز مرة واحدة ثم يُرسم على كل صفحة من صفحات الـ overlay.
    qr_temp_path = None
    qr_content = None
    qr_size = float(qr_cfg.get('size', 0.11)) * pw
    qr_x = float(qr_cfg.get('x', 0.04)) * pw
    qr_y_bottom = ph - (float(qr_cfg.get('y', 0.875)) * ph) - qr_size
    if signed and signed.get('include_qr') is not False:
        qr_content = _build_qr_payload(
            serial_text=str(signed.get('serial_number') or '').strip(),
            archive_number=archive_number,
        )
    if qr_content:
        qr_temp_path = _save_qr_png_temp(qr_content)
        if not (qr_temp_path and os.path.exists(qr_temp_path)):
            print(f"[WARNING] QR temp path invalid: path={qr_temp_path}, exists={os.path.exists(qr_temp_path) if qr_temp_path else False}", file=__import__('sys').stderr)
            qr_temp_path = None
    elif signed:
        print(f"[WARNING] No QR content generated for signed doc: serial={signed.get('serial_number')}, archive={archive_number}", file=__import__('sys').stderr)

    for page_idx, blocks in enumerate(page_blocks):
        if page_idx > 0:
            c.showPage()
            c.setFillColorRGB(0.05, 0.05, 0.05)

        _draw_page_header()
        y_cursor = y_frac
        for logical_line, sub_lines in blocks:
            font_to_use = _get_font_for_text(logical_line) if logical_line else _FONT_AR
            if font_to_use not in pdfmetrics._fonts:
                font_to_use = 'Helvetica'
            c.setFont(font_to_use, fs)

            for sub in sub_lines:
                if y_cursor >= y_end_frac:
                    break
                c.drawRightString(x_right_pt, ph - (y_cursor * ph), _reshape(sub))
                y_cursor += line_step

            if sub_lines:
                y_cursor += para_step

        if qr_temp_path:
            try:
                c.drawImage(qr_temp_path, qr_x, qr_y_bottom,
                            width=qr_size, height=qr_size, mask='auto')
            except Exception as e:
                import traceback
                print(f"[ERROR] Failed to draw QR on PDF template: {str(e)}", file=__import__('sys').stderr)
                traceback.print_exc(file=__import__('sys').stderr)

    def _draw_asset_from_data_uri(data_uri, item_type):
        if not data_uri:
            return

        pos = _pick_sign_position(signed_positions, item_type)
        x_ratio = _clamp_ratio(pos.get('x'), 0.6 if item_type != 'qr' else 0.08)
        y_ratio = _clamp_ratio(pos.get('y'), 0.8 if item_type != 'qr' else 0.92)

        if item_type == 'stamp':
            w_ratio = max(0.36, _clamp_ratio(pos.get('w'), 0.36))
            h_ratio = max(0.36, _clamp_ratio(pos.get('h'), 0.36))
        elif item_type == 'qr':
            w_ratio = max(0.12, _clamp_ratio(pos.get('w'), 0.14))
            h_ratio = max(0.12, _clamp_ratio(pos.get('h'), 0.14))
        else:
            w_ratio = max(0.20, _clamp_ratio(pos.get('w'), 0.25))
            h_ratio = max(0.12, _clamp_ratio(pos.get('h'), 0.12))

        ow = max(36.0, float(pw * w_ratio))
        oh = max(36.0, float(ph * h_ratio))

        x_pt = float(pw * x_ratio)
        y_bottom = float(ph - (y_ratio * ph) - oh)
        if x_pt + ow > pw:
            x_pt = max(0.0, pw - ow)
        if y_bottom + oh > ph:
            y_bottom = max(0.0, ph - oh)
        if y_bottom < 0:
            y_bottom = 0.0

        try:
            payload = str(data_uri).split(',', 1)[1]
            blob = base64.b64decode(payload)
            img_reader = ImageReader(io.BytesIO(blob))
            src_w, src_h = img_reader.getSize()
            draw_w, draw_h, off_x, off_y = _fit_inside_box(ow, oh, src_w, src_h)
            draw_x = x_pt + off_x
            draw_y = y_bottom + off_y
            c.drawImage(img_reader, draw_x, draw_y, width=draw_w, height=draw_h, mask='auto')
        except Exception:
            return

        ref_label = '' if item_type == 'stamp' else serial_short
        if item_type in ('sig', 'stamp') and ref_label:
            tiny = str(ref_label) if item_type == 'stamp' else f"#{ref_label}"
            tiny_fs = max(7, int(min(pw, ph) * 0.010))
            tiny_font = _FONT_EN if _FONT_EN in pdfmetrics._fonts else 'Helvetica'
            c.setFont(tiny_font, tiny_fs)
            c.drawString(draw_x + 2, max(6, draw_y - (tiny_fs + 2)), tiny)

    if signed:
        _draw_asset_from_data_uri(signed.get('signature_data_uri'), 'sig')
        _draw_asset_from_data_uri(signed.get('stamp_data_uri'), 'stamp')

    c.save()
    overlay_buf.seek(0)

    # ── Merge template PDF + overlay ─────────────────────────────────
    try:
        template_reader = PdfReader(template_path)
        overlay_reader  = PdfReader(overlay_buf)
        writer = PdfWriter()
        template_pages = list(template_reader.pages)
        overlay_pages = list(overlay_reader.pages)
        total_pages = max(len(template_pages), len(overlay_pages), 1)

        for idx in range(total_pages):
            if idx < len(template_pages):
                tpl_page = template_pages[idx]
            else:
                # لو النص أطول من صفحات القالب، كرر الصفحة الأولى كخلفية.
                tpl_page = PdfReader(template_path).pages[0]

            if idx < len(overlay_pages):
                tpl_page.merge_page(overlay_pages[idx])

            writer.add_page(tpl_page)
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
        tmp.close()
        with open(tmp.name, 'wb') as _fout:
            writer.write(_fout)
        print(f"[INFO] Successfully built PDF for doc {doc_id} at {tmp.name}", file=__import__('sys').stderr)
        return tmp.name
    except Exception as e:
        print(f"[ERROR] PDF merge failed for doc {doc_id}: {str(e)}", file=__import__('sys').stderr)
        import traceback
        traceback.print_exc(file=__import__('sys').stderr)
        return None
    finally:
        if qr_temp_path:
            try:
                os.unlink(qr_temp_path)
            except Exception:
                pass


def _build_text_document_archive_file(doc, conn=None):
    tpl_name = str(doc.get('template_name') or '').strip()

    # ── PDF template (كليشة PDF) ──────────────────────────────────────
    if tpl_name.lower().endswith('.pdf'):
        pdf_result = _build_doc_with_pdf_template(doc, conn=conn)
        if pdf_result and os.path.exists(pdf_result):
            return pdf_result

    # ── DOCX template ─────────────────────────────────────────────────
    if tpl_name.lower().endswith('.docx'):
        generated_docx = _build_text_document_docx_file(doc, conn=conn)
        if generated_docx and os.path.exists(generated_docx):
            return generated_docx

    # ── Fallback: HTML ────────────────────────────────────────────────
    final_html = _build_text_document_archive_markup(doc, conn=conn)
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.html', mode='w', encoding='utf-8')
    try:
        temp_file.write(final_html)
        temp_file.flush()
    finally:
        temp_file.close()
    return temp_file.name


def _build_signed_document_file(doc, conn=None):
    doc_id = doc.get('id', 'UNKNOWN')
    ext = _get_document_extension(doc)
    source_path = _resolve_document_file_path(doc.get('file_path'))
    tpl_name = str(doc.get('template_name') or '').strip()
    
    print(f"[INFO] Building signed file for doc {doc_id}: ext={ext}, has_source={bool(source_path)}, template={tpl_name}", file=__import__('sys').stderr)

    # ── PDF template (كليشة PDF) ─────────────────────────────────────
    # PRIORITY 1: If PDF template exists, use it (even if there's a file_path)
    # This ensures signed/stamped documents use the template with all overlays
    if tpl_name.lower().endswith('.pdf'):
        print(f"[INFO] Attempting to build from PDF template for doc {doc_id}", file=__import__('sys').stderr)
        pdf_result = _build_doc_with_pdf_template(doc, conn=conn)
        if pdf_result and os.path.exists(pdf_result):
            print(f"[INFO] Successfully built PDF template for doc {doc_id}", file=__import__('sys').stderr)
            return pdf_result
        else:
            print(f"[WARNING] PDF template build failed or returned None for doc {doc_id}", file=__import__('sys').stderr)

    if ext == 'docx' or (not ext and tpl_name and not tpl_name.lower().endswith('.pdf')):
        generated_docx = _build_text_document_docx_file(doc, conn=conn)
        if generated_docx and os.path.exists(generated_docx):
            return generated_docx

    if ext in ('png', 'jpg', 'jpeg', 'gif', 'webp') and source_path and os.path.exists(source_path):
        signed = _get_latest_signed_assets_preview(doc.get('id'), conn=conn)
        if not signed:
            return source_path

        serial_short = str(_short_signature_serial(signed.get('serial_number'), section_code=doc.get('archive_section_code')))
        stamp_reference = str(doc.get('archive_number') or signed.get('stamp_reference') or '').strip()

        sign_items = []
        sig_img = signed.get('signature_data_uri')
        stamp_img = signed.get('stamp_data_uri')
        qr_img = signed.get('qr_data_uri') if signed.get('include_qr') else None
        positions = signed.get('positions') or []
        if sig_img:
            sign_items.append(('sig', sig_img, _pick_sign_position(positions, 'sig')))
        if stamp_img:
            sign_items.append(('stamp', stamp_img, _pick_sign_position(positions, 'stamp')))
        if qr_img:
            sign_items.append(('qr', qr_img, _pick_sign_position(positions, 'qr')))

        if not sign_items:
            return source_path

        try:
            from PIL import Image, ImageDraw, ImageFont
        except Exception:
            return source_path

        try:
            base = Image.open(source_path).convert('RGBA')
            width, height = base.size
            draw = ImageDraw.Draw(base)
            for kind, data_uri, pos in sign_items:
                try:
                    payload = data_uri.split(',', 1)[1]
                    blob = base64.b64decode(payload)
                    overlay = Image.open(io.BytesIO(blob)).convert('RGBA')
                except Exception:
                    continue

                x_ratio = _clamp_ratio(pos.get('x'), 0.6)
                y_ratio = _clamp_ratio(pos.get('y'), 0.8)
                if kind == 'stamp':
                    w_ratio = max(0.36, _clamp_ratio(pos.get('w'), 0.36))
                    h_ratio = max(0.36, _clamp_ratio(pos.get('h'), 0.36))
                elif kind == 'qr':
                    w_ratio = max(0.12, _clamp_ratio(pos.get('w'), 0.14))
                    h_ratio = max(0.12, _clamp_ratio(pos.get('h'), 0.14))
                else:
                    w_ratio = max(0.20, _clamp_ratio(pos.get('w'), 0.25))
                    h_ratio = max(0.12, _clamp_ratio(pos.get('h'), 0.12))

                ow = max(48, int(width * w_ratio))
                oh = max(48, int(height * h_ratio))
                draw_w, draw_h, off_x, off_y = _fit_inside_box(ow, oh, overlay.width, overlay.height)
                draw_w = max(1, int(round(draw_w)))
                draw_h = max(1, int(round(draw_h)))
                draw_x_off = int(round(off_x))
                draw_y_off = int(round(off_y))
                overlay = overlay.resize((draw_w, draw_h), Image.LANCZOS)

                x_px = int(width * x_ratio)
                y_px = int(height * y_ratio)
                if x_px + ow > width:
                    x_px = max(0, width - ow)
                if y_px + oh > height:
                    y_px = max(0, height - oh)

                draw_x = x_px + draw_x_off
                draw_y = y_px + draw_y_off
                base.alpha_composite(overlay, dest=(draw_x, draw_y))

                ref_label = '' if kind == 'stamp' else serial_short
                if ref_label:
                    tiny = str(ref_label) if kind == 'stamp' else f"#{ref_label}"
                    font_size = max(8, int(min(width, height) * 0.018))
                    try:
                        font = ImageFont.truetype('arial.ttf', font_size)
                    except Exception:
                        font = ImageFont.load_default()
                    draw.text((draw_x + 2, max(0, draw_y - (font_size + 2))), tiny, fill=(30, 52, 94, 220), font=font)

            suffix = '.png' if ext in ('png', 'webp', 'gif') else '.jpg'
            tmp = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
            tmp.close()
            if suffix == '.jpg':
                base.convert('RGB').save(tmp.name, quality=92)
            else:
                base.save(tmp.name)
            return tmp.name
        except Exception:
            return source_path

    if ext == 'pdf' and source_path and os.path.exists(source_path):
        signed = _get_latest_signed_assets_preview(doc.get('id'), conn=conn)
        if not signed:
            print(f"[INFO] No signed assets for PDF file doc {doc_id}, returning unsigned", file=__import__('sys').stderr)
            return source_path
        else:
            print(f"[INFO] Found signed assets for PDF file doc {doc_id}, should overlay but currently returning unsigned PDF", file=__import__('sys').stderr)

        serial_short = str(_short_signature_serial(signed.get('serial_number'), section_code=doc.get('archive_section_code')))
        stamp_reference = str(doc.get('archive_number') or signed.get('stamp_reference') or '').strip()

        sign_items = []
        sig_img = signed.get('signature_data_uri')
        stamp_img = signed.get('stamp_data_uri')
        qr_img = signed.get('qr_data_uri') if signed.get('include_qr') else None
        positions = signed.get('positions') or []
        if sig_img:
            sign_items.append(('sig', sig_img, _pick_sign_position(positions, 'sig')))
        if stamp_img:
            sign_items.append(('stamp', stamp_img, _pick_sign_position(positions, 'stamp')))
        if qr_img:
            sign_items.append(('qr', qr_img, _pick_sign_position(positions, 'qr')))

        if not sign_items:
            return source_path

        try:
            import pypdfium2 as pdfium
            from PIL import Image, ImageDraw, ImageFont
        except Exception:
            return source_path

        pdf = None
        rendered_pages = []
        try:
            pdf = pdfium.PdfDocument(source_path)
            total_pages = len(pdf)
            if total_pages <= 0:
                return source_path

            for page_index in range(total_pages):
                page = None
                bitmap = None
                try:
                    page = pdf[page_index]
                    bitmap = page.render(scale=1.7)
                    base = bitmap.to_pil().convert('RGBA')

                    if page_index == 0:
                        width, height = base.size
                        draw = ImageDraw.Draw(base)

                        for kind, data_uri, pos in sign_items:
                            try:
                                payload = data_uri.split(',', 1)[1]
                                blob = base64.b64decode(payload)
                                overlay = Image.open(io.BytesIO(blob)).convert('RGBA')
                            except Exception:
                                continue

                            x_ratio = _clamp_ratio(pos.get('x'), 0.6)
                            y_ratio = _clamp_ratio(pos.get('y'), 0.8)
                            if kind == 'stamp':
                                w_ratio = max(0.36, _clamp_ratio(pos.get('w'), 0.36))
                                h_ratio = max(0.36, _clamp_ratio(pos.get('h'), 0.36))
                            elif kind == 'qr':
                                w_ratio = max(0.12, _clamp_ratio(pos.get('w'), 0.14))
                                h_ratio = max(0.12, _clamp_ratio(pos.get('h'), 0.14))
                            else:
                                w_ratio = max(0.20, _clamp_ratio(pos.get('w'), 0.25))
                                h_ratio = max(0.12, _clamp_ratio(pos.get('h'), 0.12))

                            ow = max(48, int(width * w_ratio))
                            oh = max(48, int(height * h_ratio))
                            draw_w, draw_h, off_x, off_y = _fit_inside_box(ow, oh, overlay.width, overlay.height)
                            draw_w = max(1, int(round(draw_w)))
                            draw_h = max(1, int(round(draw_h)))
                            draw_x_off = int(round(off_x))
                            draw_y_off = int(round(off_y))
                            overlay = overlay.resize((draw_w, draw_h), Image.LANCZOS)

                            x_px = int(width * x_ratio)
                            y_px = int(height * y_ratio)
                            if x_px + ow > width:
                                x_px = max(0, width - ow)
                            if y_px + oh > height:
                                y_px = max(0, height - oh)

                            draw_x = x_px + draw_x_off
                            draw_y = y_px + draw_y_off
                            base.alpha_composite(overlay, dest=(draw_x, draw_y))

                            ref_label = '' if kind == 'stamp' else serial_short
                            if ref_label:
                                tiny = str(ref_label) if kind == 'stamp' else f"#{ref_label}"
                                font_size = max(8, int(min(width, height) * 0.012))
                                try:
                                    font = ImageFont.truetype('arial.ttf', font_size)
                                except Exception:
                                    font = ImageFont.load_default()

                                draw.text((draw_x + 2, max(0, draw_y - (font_size + 2))), tiny, fill=(30, 52, 94, 220), font=font)

                    rendered_pages.append(base.convert('RGB'))
                finally:
                    try:
                        if bitmap is not None and hasattr(bitmap, 'close'):
                            bitmap.close()
                    except Exception:
                        pass
                    try:
                        if page is not None and hasattr(page, 'close'):
                            page.close()
                    except Exception:
                        pass

            if not rendered_pages:
                return source_path

            tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
            tmp.close()

            first = rendered_pages[0]
            tail = rendered_pages[1:]
            first.save(tmp.name, format='PDF', resolution=150.0, save_all=bool(tail), append_images=tail)
            return tmp.name
        except Exception:
            return source_path
        finally:
            for image_obj in rendered_pages:
                try:
                    image_obj.close()
                except Exception:
                    pass
            try:
                if pdf is not None and hasattr(pdf, 'close'):
                    pdf.close()
            except Exception:
                pass

    if not source_path and doc.get('content_json'):
        return _build_text_document_archive_file(doc, conn=conn)

    if ext in ('doc', 'docx') and source_path and os.path.exists(source_path):
        return _build_text_document_archive_file(doc, conn=conn)

    return source_path


def _qr_svg(serial):
    serial_text = str(serial or "").strip() or "NAJM"
    digest = hashlib.sha256(serial_text.encode('utf-8')).digest()

    size = 25
    cell = 8
    margin = 12
    width = (size * cell) + (margin * 2)

    bitstream = "".join(format(b, "08b") for b in digest * 20)

    def finder(x0, y0):
        out = []
        for y in range(7):
            for x in range(7):
                border = x in (0, 6) or y in (0, 6)
                center = 2 <= x <= 4 and 2 <= y <= 4
                if border or center:
                    out.append((x0 + x, y0 + y))
        return out

    reserved = set(finder(0, 0) + finder(size - 7, 0) + finder(0, size - 7))
    cells = set(reserved)

    idx = 0
    for y in range(size):
        for x in range(size):
            if (x, y) in reserved:
                continue
            if bitstream[idx % len(bitstream)] == '1':
                cells.add((x, y))
            idx += 1

    rects = "".join(
        f"<rect x='{margin + (x * cell)}' y='{margin + (y * cell)}' width='{cell}' height='{cell}' fill='#111827'/>"
        for (x, y) in sorted(cells)
    )

    safe_serial = html.escape(serial_text)
    return f"""
    <svg xmlns='http://www.w3.org/2000/svg' width='{width}' height='{width + 56}' viewBox='0 0 {width} {width + 56}'>
      <rect width='100%' height='100%' fill='white'/>
      {rects}
      <text x='50%' y='{width + 34}' text-anchor='middle' fill='#334155' font-size='13' font-family='Arial'>{safe_serial}</text>
    </svg>
    """


def _count_pdf_pages(file_path):
    if not file_path:
        return 1
    try:
        from pypdf import PdfReader
        reader = PdfReader(file_path)
        return max(1, len(reader.pages))
    except Exception:
        return 1


def _render_pdf_page_png(file_path, page_index=0, scale=1.5):
    if not file_path or not os.path.exists(file_path):
        return None, 0

    pdf = None
    page = None
    bitmap = None
    image = None

    try:
        import pypdfium2 as pdfium

        with PDF_RENDER_LOCK:
            pdf = pdfium.PdfDocument(file_path)
            total_pages = len(pdf)
            if total_pages <= 0:
                return None, 0

            safe_page_index = min(max(0, _safe_int(page_index, 0)), total_pages - 1)
            page = pdf[safe_page_index]
            bitmap = page.render(scale=max(1.0, float(scale or 1.5)))
            image = bitmap.to_pil()

            stream = io.BytesIO()
            image.save(stream, format='PNG', optimize=True)
            stream.seek(0)
            return stream, total_pages
    except Exception:
        return None, 0
    finally:
        try:
            if image is not None:
                image.close()
        except Exception:
            pass
        try:
            if bitmap is not None and hasattr(bitmap, 'close'):
                bitmap.close()
        except Exception:
            pass
        try:
            if page is not None and hasattr(page, 'close'):
                page.close()
        except Exception:
            pass
        try:
            if pdf is not None and hasattr(pdf, 'close'):
                pdf.close()
        except Exception:
            pass


def _overlay_signed_assets_on_pil_image(base_image, signed_assets, section_code=None):
    if base_image is None or not signed_assets:
        return base_image

    try:
        from PIL import Image, ImageDraw, ImageFont
    except Exception:
        return base_image

    signature_data_uri = signed_assets.get('signature_data_uri')
    stamp_data_uri = signed_assets.get('stamp_data_uri')
    qr_data_uri = signed_assets.get('qr_data_uri') if signed_assets.get('include_qr') else None
    positions = signed_assets.get('positions') or []

    sign_items = []
    if signature_data_uri:
        sign_items.append(('sig', signature_data_uri, _pick_sign_position(positions, 'sig')))
    if stamp_data_uri:
        sign_items.append(('stamp', stamp_data_uri, _pick_sign_position(positions, 'stamp')))
    if qr_data_uri:
        sign_items.append(('qr', qr_data_uri, _pick_sign_position(positions, 'qr')))

    if not sign_items:
        return base_image

    serial_short = str(_short_signature_serial(signed_assets.get('serial_number'), section_code=section_code))
    stamp_reference = str(signed_assets.get('stamp_reference') or signed_assets.get('archive_number') or '').strip()

    output_image = base_image if getattr(base_image, 'mode', '') == 'RGBA' else base_image.convert('RGBA')
    width, height = output_image.size
    draw = ImageDraw.Draw(output_image)

    for kind, data_uri, pos in sign_items:
        overlay = None
        try:
            payload = str(data_uri or '').split(',', 1)[1]
            blob = base64.b64decode(payload)
            overlay = Image.open(io.BytesIO(blob)).convert('RGBA')

            x_ratio = _clamp_ratio(pos.get('x'), 0.6)
            y_ratio = _clamp_ratio(pos.get('y'), 0.8)
            if kind == 'stamp':
                w_ratio = max(0.36, _clamp_ratio(pos.get('w'), 0.36))
                h_ratio = max(0.36, _clamp_ratio(pos.get('h'), 0.36))
            elif kind == 'qr':
                w_ratio = max(0.12, _clamp_ratio(pos.get('w'), 0.14))
                h_ratio = max(0.12, _clamp_ratio(pos.get('h'), 0.14))
            else:
                w_ratio = max(0.20, _clamp_ratio(pos.get('w'), 0.25))
                h_ratio = max(0.12, _clamp_ratio(pos.get('h'), 0.12))

            ow = max(48, int(width * w_ratio))
            oh = max(48, int(height * h_ratio))
            draw_w, draw_h, off_x, off_y = _fit_inside_box(ow, oh, overlay.width, overlay.height)
            draw_w = max(1, int(round(draw_w)))
            draw_h = max(1, int(round(draw_h)))
            draw_x_off = int(round(off_x))
            draw_y_off = int(round(off_y))
            overlay = overlay.resize((draw_w, draw_h), Image.LANCZOS)

            x_px = int(width * x_ratio)
            y_px = int(height * y_ratio)
            if x_px + ow > width:
                x_px = max(0, width - ow)
            if y_px + oh > height:
                y_px = max(0, height - oh)

            draw_x = x_px + draw_x_off
            draw_y = y_px + draw_y_off
            output_image.alpha_composite(overlay, dest=(draw_x, draw_y))

            ref_label = '' if kind == 'stamp' else serial_short
            if ref_label:
                tiny = str(ref_label) if kind == 'stamp' else f"#{ref_label}"
                font_size = max(8, int(min(width, height) * 0.012))
                try:
                    font = ImageFont.truetype('arial.ttf', font_size)
                except Exception:
                    font = ImageFont.load_default()
                draw.text((draw_x + 2, max(0, draw_y - (font_size + 2))), tiny, fill=(30, 52, 94, 220), font=font)
        except Exception:
            continue
        finally:
            try:
                if overlay is not None:
                    overlay.close()
            except Exception:
                pass

    return output_image


def _is_manager_role(role_value):
    return str(role_value or '').strip().lower() in ('manager', 'admin', 'sys_admin')


def _get_signature_assets(conn, viewer_id=None, viewer_role=''):
    rows = conn.execute(
        """
        SELECT sa.id AS asset_id,
               sa.user_id AS owner_id,
               sa.signature_name,
               sa.signature_path,
                         COALESCE(sa.visibility_scope, 'self') AS visibility_scope,
               sa.visible_to_user_id,
               u.name,
               u.role
        FROM signature_assets sa
        LEFT JOIN users u ON u.id = sa.user_id
        WHERE sa.is_active=1
          AND COALESCE(sa.signature_path, '') != ''
          AND COALESCE(u.is_active, 1)=1
        ORDER BY sa.id DESC
        """
    ).fetchall()

    normalized_viewer_role = str(viewer_role or '').strip().lower()
    viewer_id_int = _safe_int(viewer_id, 0)

    filtered = []
    for row in rows:
        item = dict(row)
        owner_role = str(item.get('role') or '').strip().lower()
        if not _is_manager_role(owner_role):
            continue

        scope = str(item.get('visibility_scope') or 'self').strip().lower()
        if scope not in ('all', 'managers', 'specific', 'self'):
            scope = 'self'
            item['visibility_scope'] = 'self'

        visible_to_user_id = _safe_int(item.get('visible_to_user_id'), 0)
        item['visible_to_user_id'] = visible_to_user_id if visible_to_user_id > 0 else None

        owner_id = _safe_int(item.get('owner_id'), 0)
        is_owner = (viewer_id_int > 0 and viewer_id_int == owner_id)
        is_visible = False
        if is_owner:
            is_visible = True
        elif scope == 'all':
            is_visible = True
        elif scope == 'managers':
            is_visible = _is_manager_role(normalized_viewer_role)
        elif scope == 'specific':
            is_visible = (viewer_id_int > 0 and visible_to_user_id == viewer_id_int)
        elif scope == 'self':
            is_visible = False

        if is_visible:
            filtered.append(item)

    return filtered


def _get_signature_visibility_targets(conn):
    rows = conn.execute(
        """
        SELECT id, name, role
        FROM users
        WHERE is_active=1
          AND role IN ('manager', 'admin', 'sys_admin')
        ORDER BY CASE WHEN role='admin' THEN 0 WHEN role='sys_admin' THEN 1 ELSE 2 END, name ASC
        """
    ).fetchall()
    return [dict(r) for r in rows]


def _get_stamp_assets(conn, viewer_id=None, viewer_role=''):
    normalized_viewer_role = str(viewer_role or '').strip().lower()
    viewer_id_int = _safe_int(viewer_id, 0)

    image_rows = conn.execute(
        """
        SELECT st.id AS asset_id,
               st.user_id AS owner_id,
               st.stamp_name,
               st.stamp_path,
             COALESCE(st.visibility_scope, 'self') AS visibility_scope,
               st.visible_to_user_id,
               u.name,
               u.role,
               u.stamp_text
        FROM stamp_assets st
        LEFT JOIN users u ON u.id = st.user_id
        WHERE st.is_active=1
          AND COALESCE(st.stamp_path, '') != ''
          AND COALESCE(u.is_active, 1)=1
          AND COALESCE(u.role, '') IN ('manager', 'admin', 'sys_admin')
        ORDER BY st.id DESC
        """
    ).fetchall()

    def _stamp_visible(scope, visible_to_uid, owner_id):
        scope = str(scope or 'self').strip().lower()
        if scope not in ('all', 'managers', 'specific', 'self'):
            scope = 'self'
        oid = _safe_int(owner_id, 0)
        is_owner = (viewer_id_int > 0 and oid == viewer_id_int)
        if is_owner:
            return True
        if scope == 'all':
            return True
        if scope == 'managers':
            return _is_manager_role(normalized_viewer_role)
        if scope == 'specific':
            uid = _safe_int(visible_to_uid, 0)
            return (viewer_id_int > 0 and uid == viewer_id_int)
        if scope == 'self':
            return False
        return True

    assets = []
    for row in image_rows:
        item = dict(row)
        item['kind'] = 'image'
        vtu = _safe_int(item.get('visible_to_user_id'), 0)
        item['visible_to_user_id'] = vtu if vtu > 0 else None
        if not _stamp_visible(item.get('visibility_scope'), item.get('visible_to_user_id'), item.get('owner_id')):
            continue
        assets.append(item)

    text_rows = conn.execute(
        """
        SELECT u.id AS owner_id,
               u.name,
               u.role,
               u.stamp_text,
               u.employee_id,
               COALESCE(u.stamp_visibility_scope, 'self') AS visibility_scope,
               u.stamp_visible_to_user_id,
               s.section_code
        FROM users u
        LEFT JOIN archive_sections s ON s.id = u.archive_section_id
        WHERE u.is_active=1
          AND COALESCE(u.stamp_text, '') != ''
          AND COALESCE(u.role, '') IN ('manager', 'admin', 'sys_admin')
        ORDER BY u.id DESC
        """
    ).fetchall()

    for row in text_rows:
        owner_id = row['owner_id']
        scope = row['visibility_scope']
        visible_to_uid = _safe_int(row.get('stamp_visible_to_user_id'), 0)
        if not _stamp_visible(scope, visible_to_uid, owner_id):
            continue
        # Use actual stamp text content for better identification in dropdown
        stamp_text_content = row['stamp_text'] or 'ختم المدير'
        assets.append({
            'asset_id': f"text-{owner_id}",
            'owner_id': owner_id,
            'stamp_name': stamp_text_content,
            'stamp_path': None,
            'visibility_scope': scope,
            'visible_to_user_id': visible_to_uid if visible_to_uid > 0 else None,
            'name': row['name'],
            'role': row['role'],
            'stamp_text': row['stamp_text'],
            'employee_id': row['employee_id'],
            'section_code': row['section_code'],
            'kind': 'text',
        })

    return assets

# ════════════════════════════════════════════
# AUTH
# ════════════════════════════════════════════

@app.route("/")
def index():
    return redirect(url_for('dashboard') if 'user_id' in session else url_for('login'))

@app.route("/favicon.ico")
def favicon():
    favicon_path = os.path.join(app.root_path, "static", "favicon.ico")
    if os.path.exists(favicon_path):
        return send_file(favicon_path, mimetype="image/x-icon")
    return "", 204

@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "GET":
        return render_template("login.html")

    data     = request.json
    email    = data.get("email","").strip().lower()
    password = data.get("password","")
    ip       = request.remote_addr
    ua       = request.user_agent.string

    user = get_user_by_email(email)
    if not user or not check_password(password, user['password_hash']):
        if user:
            log_login(user['id'], "failed", ip, ua)
        return jsonify({"success": False, "error": "البريد أو كلمة المرور غير صحيحة"}), 401

    # ── تحقق من الجهاز
    device_hash = get_device_hash(ip, ua)

    if not is_trusted_device(user['id'], device_hash):
        # جهاز جديد — أرسل إيميل تأكيد
        token, device_name = register_pending_device(user['id'], device_hash, ip, ua)
        _base      = _get_verify_base_url()
        confirm_url = f"{_base}/device/confirm/{token}"
        deny_url    = f"{_base}/device/deny/{token}"
        send_new_device(user['email'], user['name'], device_name, ip, confirm_url, deny_url)
        log_login(user['id'], "new_device", ip, ua)
        return jsonify({
            "success":    False,
            "new_device": True,
            "error":      "تم إرسال إيميل تأكيد للجهاز الجديد — تحقق من بريدك"
        })

    # جهاز موثوق — أرسل OTP
    code = generate_otp(user['id'])
    send_otp(user['email'], user['name'], code)
    session['pending_user_id'] = user['id']
    log_login(user['id'], "otp_sent", ip, ua)
    return jsonify({"success": True})

# ── تأكيد الجهاز من الإيميل ──────────────────
@app.route("/device/confirm/<token>")
def device_confirm(token):
    row = trust_device_by_token(token)
    if not row:
        return render_template("device_result.html",
                               success=False,
                               msg="الرابط غير صالح أو انتهت صلاحيته")

    user = get_user_by_id(row['user_id'])
    log_action(row['user_id'], "DEVICE_TRUSTED",
               f"جهاز: {row['device_name']}",
               ip=row['ip_address'])
    return render_template("device_result.html",
                           success=True,
                           msg=f"✅ تم تسجيل الجهاز بنجاح\nيمكنك الآن تسجيل الدخول من هذا الجهاز")

# ── رفض الجهاز + تجميد الحساب ───────────────
@app.route("/device/deny/<token>")
def device_deny(token):
    conn = get_db()
    row  = conn.execute(
        "SELECT * FROM trusted_devices WHERE token=?", (token,)
    ).fetchone()
    if row:
        block_user(row['user_id'])
        conn.execute("DELETE FROM trusted_devices WHERE token=?", (token,))
        conn.commit()
        log_action(row['user_id'], "ACCOUNT_BLOCKED", "المستخدم رفض الجهاز الجديد")
    conn.close()
    return render_template("device_result.html",
                           success=False,
                           msg="🔒 تم تجميد الحساب فوراً — تواصل مع المدير العام")

@app.route("/verify-otp", methods=["POST"])
def verify_otp_route():
    data    = request.json
    code    = data.get("code","").strip()
    user_id = session.get('pending_user_id')

    if not user_id:
        return jsonify({"success": False, "error": "انتهت الجلسة، أعد المحاولة"}), 400

    if not verify_otp(user_id, code):
        log_login(user_id, "otp_fail", request.remote_addr, request.user_agent.string)
        return jsonify({"success": False, "error": "الكود غير صحيح أو انتهت صلاحيته"}), 401

    user = get_user_by_id(user_id)
    session.clear()
    session['user_id']    = user['id']
    session['user_name']  = user['name']
    session['user_role']  = user['role']
    session['user_email'] = user['email']
    session.permanent     = True

    log_login(user['id'], "login", request.remote_addr, request.user_agent.string)
    log_action(user['id'], "LOGIN", ip=request.remote_addr, user_agent=request.user_agent.string)

    if user['first_login']:
        return jsonify({"success": True, "redirect": url_for('change_password_page')})
    return jsonify({"success": True, "redirect": url_for('dashboard')})

@app.route("/change-password", methods=["GET"])
@login_required
def change_password_page():
    return render_template("change_password.html")

@app.route("/api/change-password", methods=["POST"])
@login_required
def api_change_password():
    data     = request.json
    new_pass = data.get("new_password","")
    if len(new_pass) < 8:
        return jsonify({"success": False, "error": "الباسورد يجب أن يكون 8 أحرف على الأقل"})
    update_password(session['user_id'], new_pass)
    log_action(session['user_id'], "CHANGE_PASSWORD", ip=request.remote_addr)
    return jsonify({"success": True, "redirect": url_for('dashboard')})

@app.route("/logout")
def logout():
    if 'user_id' in session:
        log_login(session['user_id'], "logout", request.remote_addr, request.user_agent.string)
        log_action(session['user_id'], "LOGOUT", ip=request.remote_addr)
    session.clear()
    return redirect(url_for('login'))


# ════════════════════════════════════════════
# SELF REGISTRATION
# ════════════════════════════════════════════

# --- بسيط: تحديد معدل التسجيل لكل IP ---
import time as _time
_reg_rate: dict = {}          # ip -> [timestamp, ...]
_reg_rate_lock = threading.Lock()
_REG_WINDOW = 3600            # ساعة واحدة
_REG_MAX    = 5               # أقصى 5 محاولات في الساعة

def _check_reg_rate(ip: str) -> bool:
    """True = مسموح، False = تجاوز الحد."""
    now = _time.time()
    with _reg_rate_lock:
        stamps = [t for t in _reg_rate.get(ip, []) if now - t < _REG_WINDOW]
        if len(stamps) >= _REG_MAX:
            _reg_rate[ip] = stamps
            return False
        stamps.append(now)
        _reg_rate[ip] = stamps
    return True


@app.route("/api/register/sections", methods=["GET"])
def api_register_sections():
    """أقسام الأرشيف النشطة — عامة بلا مصادقة لصفحة التسجيل."""
    sections = get_all_archive_sections(include_inactive=False)
    return jsonify({"success": True, "sections": sections})


@app.route("/api/register", methods=["POST"])
def api_register():
    """تسجيل ذاتي للموظف — الدور user تلقائياً، بدون صلاحية admin."""
    ip = request.remote_addr or "unknown"

    if not _check_reg_rate(ip):
        return jsonify({"success": False,
                        "error": "تجاوزت الحد المسموح به من محاولات التسجيل، حاول مجدداً بعد ساعة"}), 429

    data = request.get_json(silent=True) or {}

    # ── استخراج الحقول ──────────────────────
    name        = str(data.get("name", "") or "").strip()
    email       = str(data.get("email", "") or "").strip().lower()
    phone       = str(data.get("phone", "") or "").strip()
    job_title   = str(data.get("job_title", "") or "").strip() or None
    employee_id = str(data.get("employee_id", "") or "").strip() or None
    section_id  = data.get("section_id")
    password    = str(data.get("password", "") or "")
    confirm     = str(data.get("confirm_password", "") or "")

    # ── تحقق من الحقول المطلوبة ─────────────
    if not name:
        return jsonify({"success": False, "error": "الاسم مطلوب"}), 400
    if len(name) < 2 or len(name) > 100:
        return jsonify({"success": False, "error": "الاسم يجب أن يكون بين 2 و100 حرف"}), 400
    if not re.fullmatch(r'^[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}$', email):
        return jsonify({"success": False, "error": "البريد الإلكتروني غير صالح"}), 400
    if not phone:
        return jsonify({"success": False, "error": "رقم الجوال مطلوب"}), 400
    if not re.fullmatch(r'^[\d\s\+\-]{7,20}$', phone):
        return jsonify({"success": False, "error": "رقم الجوال غير صالح"}), 400
    if not section_id:
        return jsonify({"success": False, "error": "اختر القسم"}), 400
    try:
        section_id = int(section_id)
    except (TypeError, ValueError):
        return jsonify({"success": False, "error": "القسم غير صالح"}), 400
    if len(password) < 8:
        return jsonify({"success": False, "error": "كلمة المرور يجب أن تكون 8 أحرف على الأقل"}), 400
    if not re.search(r'[A-Z]', password):
        return jsonify({"success": False, "error": "كلمة المرور يجب أن تحتوي على حرف كبير على الأقل"}), 400
    if not re.search(r'\d', password):
        return jsonify({"success": False, "error": "كلمة المرور يجب أن تحتوي على رقم على الأقل"}), 400
    if password != confirm:
        return jsonify({"success": False, "error": "كلمة المرور وتأكيدها غير متطابقتين"}), 400

    # ── تحقق من وجود القسم ──────────────────
    conn = get_db()
    try:
        sec_row = conn.execute(
            "SELECT id FROM archive_sections WHERE id=? AND is_active=1 LIMIT 1",
            (section_id,)
        ).fetchone()
        if not sec_row:
            return jsonify({"success": False, "error": "القسم المختار غير موجود"}), 400

        # ── تحقق من تكرار البريد ────────────
        existing = conn.execute(
            "SELECT id FROM users WHERE LOWER(email)=?", (email,)
        ).fetchone()
        if existing:
            return jsonify({"success": False,
                            "error": "هذا البريد الإلكتروني مسجل مسبقاً"}), 409

        # ── إنشاء الحساب ────────────────────
        conn.execute(
            """INSERT INTO users
               (name, email, phone, job_title, role, password_hash,
                first_login, created_by, archive_section_id, employee_id)
               VALUES (?,?,?,?,?,?,0,NULL,?,?)""",
            (name, email, phone, job_title, "user",
             hash_password(password), section_id, employee_id)
        )
        conn.commit()
        new_id = conn.execute("SELECT last_insert_rowid()").fetchone()[0]
    finally:
        conn.close()

    # ── سجّل العملية ────────────────────────
    log_action(new_id, "SELF_REGISTER",
               f"تسجيل ذاتي | IP:{ip}", ip=ip,
               user_agent=request.user_agent.string)

    # ── إيميل ترحيب (اختياري — لن يوقف التسجيل إن فشل) ─
    try:
        user_obj = get_user_by_id(new_id)
        if user_obj:
            send_welcome(user_obj['email'], user_obj['name'], password)
    except Exception:
        pass

    return jsonify({"success": True,
                    "message": "تم إنشاء حسابك بنجاح، يمكنك الآن تسجيل الدخول"})


@app.route("/api/qr/<serial>")
def api_qr(serial):
    qr_payload = _build_qr_payload(serial_text=serial)
    if not qr_payload:
        return _png_response(_build_qr_png_bytes('NAJM'), status=200)
    return _png_response(_build_qr_png_bytes(qr_payload), status=200)


@app.route("/verify", methods=["GET", "POST"])
def verify_lookup():
    serial_value = ''
    if request.method == 'POST':
        serial_value = str(request.form.get('serial') or '').strip().upper()
    else:
        serial_value = str(request.args.get('serial') or request.args.get('q') or '').strip().upper()

    if serial_value:
        return redirect(url_for('verify_serial', serial=serial_value))

    return render_template(
        "verify.html",
        doc=None,
        serial='',
        sign_hash=None,
        signers=[],
        operations=[],
        now=datetime.now().strftime('%Y-%m-%d %H:%M'),
        lookup_performed=False
    )


@app.route("/verify/<serial>")
def verify_serial(serial):
    serial_lookup = str(serial or '').strip().upper()
    conn = get_db()
    sr = conn.execute(
        """
        SELECT sr.serial_number, sr.sign_hash, sr.document_id,
               sr.stamp_owner_id, sr.requested_from, sr.signed_at,
               d.title, d.archive_number, d.status, d.created_at, d.approved_at,
               d.created_by,
               cu.name AS creator_name, cu.job_title AS creator_job_title, cu.role AS creator_role
        FROM signature_requests sr
        LEFT JOIN documents d ON d.id = sr.document_id
        LEFT JOIN users cu ON cu.id = d.created_by
        WHERE UPPER(TRIM(COALESCE(sr.serial_number, '')))=?
        ORDER BY sr.id DESC
        LIMIT 1
        """,
        (serial_lookup,)
    ).fetchone()

    if not sr:
        conn.close()
        return render_template(
            "verify.html",
            doc=None,
            serial=serial_lookup,
            sign_hash=None,
            signers=[],
            operations=[],
            stamp_creator=None,
            now=datetime.now().strftime('%Y-%m-%d %H:%M'),
            lookup_performed=True
        )

    signers = conn.execute(
        """
        SELECT u.name, sr.serial_number, sr.signed_at
        FROM signature_requests sr
        LEFT JOIN users u ON u.id = sr.requested_from
        WHERE sr.document_id=? AND sr.status='signed'
        ORDER BY sr.id ASC
        """,
        (sr['document_id'],)
    ).fetchall()

    operations = conn.execute(
        """
        SELECT al.action, al.details, al.timestamp, al.ip_address,
               u.name as user_name
        FROM audit_log al
        LEFT JOIN users u ON u.id = al.user_id
        WHERE al.document_id=?
        ORDER BY al.id ASC
        """,
        (sr['document_id'],)
    ).fetchall()

    # Show the actual document creator in verify page.
    stamp_creator = None
    creator_id = _safe_int(sr['created_by'], 0)
    if creator_id > 0 and sr['creator_name']:
        stamp_creator = {
            'name': sr['creator_name'],
            'job_title': sr['creator_job_title'],
            'role': sr['creator_role'],
            'created_at': sr['created_at']
        }
    else:
        # Fallback for older data where created_by may be missing.
        stamp_owner_id = _safe_int(sr['stamp_owner_id'], 0)
        if stamp_owner_id <= 0:
            stamp_owner_id = _safe_int(sr['requested_from'], 0)

        if stamp_owner_id > 0:
            stamp_user = conn.execute(
                """
                SELECT name, job_title, role
                FROM users
                WHERE id=? LIMIT 1
                """,
                (stamp_owner_id,)
            ).fetchone()
            if stamp_user:
                stamp_creator = {
                    'name': stamp_user['name'],
                    'job_title': stamp_user['job_title'],
                    'role': stamp_user['role'],
                    'created_at': sr['signed_at']
                }

    conn.close()

    doc = {
        "title": sr['title'],
        "archive_number": sr['archive_number'],
        "status": sr['status'] or 'pending',
        "created_at": sr['created_at'],
        "approved_at": sr['approved_at'],
    }
    return render_template(
        "verify.html",
        doc=doc,
        serial=sr['serial_number'],
        sign_hash=sr['sign_hash'],
        signers=[dict(r) for r in signers],
        operations=[dict(r) for r in operations],
        stamp_creator=stamp_creator,
        now=datetime.now().strftime('%Y-%m-%d %H:%M'),
        lookup_performed=True
    )


def _build_sign_page_payload(doc, req=None):
    conn = get_db()
    viewer_id = session.get('user_id')
    viewer_role = session.get('user_role')
    signature_assets = _get_signature_assets(
        conn,
        viewer_id=viewer_id,
        viewer_role=viewer_role
    )
    stamp_assets = _get_stamp_assets(
        conn,
        viewer_id=viewer_id,
        viewer_role=viewer_role
    )
    stamp_visibility_targets = _get_signature_visibility_targets(conn)
    signature_visibility_targets = _get_signature_visibility_targets(conn)
    conn.close()

    # Defense-in-depth: re-check visibility before exposing assets to UI.
    viewer_id_int = _safe_int(viewer_id, 0)
    viewer_role_norm = str(viewer_role or '').strip().lower()

    def _asset_visible(item):
        owner_id = _safe_int((item or {}).get('owner_id'), 0)
        if viewer_id_int > 0 and owner_id == viewer_id_int:
            return True

        scope = str((item or {}).get('visibility_scope') or 'self').strip().lower()
        if scope not in ('all', 'managers', 'specific', 'self'):
            scope = 'self'

        if scope == 'all':
            return True
        if scope == 'managers':
            return _is_manager_role(viewer_role_norm)
        if scope == 'specific':
            return _safe_int((item or {}).get('visible_to_user_id'), 0) == viewer_id_int and viewer_id_int > 0
        return False

    signature_assets = [item for item in (signature_assets or []) if _asset_visible(item)]
    stamp_assets = [item for item in (stamp_assets or []) if _asset_visible(item)]

    resolved_path = _resolve_document_file_path(doc.get('file_path'))
    if not resolved_path:
        resolved_path = _resolve_document_file_path(doc.get('archive_storage_path'))

    page_count = 1
    signed_preview_path = None
    temp_generated_path = None

    try:
        # Keep sign-page pagination aligned with what preview endpoint renders.
        signed_preview_path = _build_signed_document_file(doc)
        if signed_preview_path and os.path.exists(signed_preview_path):
            resolved_path = signed_preview_path
    except Exception:
        signed_preview_path = None

    if (not resolved_path) and doc.get('content_json') and not doc.get('file_path'):
        try:
            generated = _build_text_document_archive_file(doc)
            if generated and os.path.exists(generated):
                resolved_path = generated
                temp_generated_path = generated
        except Exception:
            pass

    if resolved_path and str(resolved_path).lower().endswith('.pdf'):
        page_count = max(1, _count_pdf_pages(resolved_path))

    # Cleanup temporary generated files used only for page counting.
    for _p in (signed_preview_path, temp_generated_path):
        try:
            if not _p:
                continue
            _abs = os.path.abspath(_p)
            _tmp = os.path.abspath(tempfile.gettempdir())
            if _abs.startswith(_tmp):
                os.remove(_abs)
        except Exception:
            pass

    default_signature_asset_id = signature_assets[0]['asset_id'] if signature_assets else None

    default_stamp_asset_id = stamp_assets[0]['asset_id'] if stamp_assets else None
    default_stamp_owner_id = stamp_assets[0]['owner_id'] if stamp_assets else None

    # Prefer dynamic text stamp so template style changes appear without re-uploading image stamps.
    viewer_id_int = _safe_int(viewer_id, 0)
    preferred_text_stamp = None
    if stamp_assets:
        preferred_text_stamp = next(
            (
                item for item in stamp_assets
                if str(item.get('asset_id') or '').startswith('text-')
                and _safe_int(item.get('owner_id'), 0) == viewer_id_int
            ),
            None
        )
        if not preferred_text_stamp:
            preferred_text_stamp = next(
                (item for item in stamp_assets if str(item.get('asset_id') or '').startswith('text-')),
                None
            )

    if preferred_text_stamp:
        default_stamp_asset_id = preferred_text_stamp.get('asset_id')
        default_stamp_owner_id = preferred_text_stamp.get('owner_id')

    # الملفات المرفوعة: توقيع/ختم فقط بدون QR
    is_uploaded_file = bool(doc.get('file_path'))

    return {
        "doc": doc,
        "req": req,
        "can_stamp": bool(stamp_assets),
        "signature_assets": signature_assets,
        "stamp_assets": stamp_assets,
        "signature_visibility_targets": signature_visibility_targets,
        "stamp_visibility_targets": stamp_visibility_targets,
        "allow_qr_option": not is_uploaded_file,
        "default_signature_asset_id": default_signature_asset_id,
        "default_stamp_asset_id": default_stamp_asset_id,
        "default_stamp_owner_id": default_stamp_owner_id,
        "page_count": max(1, page_count),
    }


@app.route("/sign/<int:req_id>")
@login_required
def sign_request_page(req_id):
    conn = get_db()
    req = conn.execute(
        """
        SELECT sr.*, u.name AS requester_name
        FROM signature_requests sr
        LEFT JOIN users u ON u.id = sr.requested_by
        WHERE sr.id=?
        LIMIT 1
        """,
        (req_id,)
    ).fetchone()

    if not req:
        conn.close()
        return redirect(url_for('documents_page'))

    role = (session.get('user_role') or '').lower()
    if req['requested_from'] != session['user_id'] and role not in ('admin', 'manager', 'sys_admin'):
        conn.close()
        return redirect(url_for('documents_page'))

    doc = conn.execute("SELECT * FROM documents WHERE id=? LIMIT 1", (req['document_id'],)).fetchone()
    conn.close()
    if not doc:
        return redirect(url_for('documents_page'))

    payload = _build_sign_page_payload(dict(doc), dict(req))
    return render_template("sign_document.html", **payload)


@app.route("/sign/doc/<int:doc_id>")
@login_required
def sign_document_page(doc_id):
    role = (session.get('user_role') or '').lower()
    uid  = session.get('user_id')
    conn = get_db()
    doc  = conn.execute("SELECT * FROM documents WHERE id=? LIMIT 1", (doc_id,)).fetchone()
    if not doc:
        conn.close()
        return redirect(url_for('documents_page'))

    # التحقق من الصلاحية قبل عرض الصفحة
    can_sign = False
    if role in ('admin', 'sys_admin'):
        can_sign = True
    elif role == 'manager' and uid:
        section_code = (doc['archive_section_code'] or '')
        perm = conn.execute(
            """SELECT 1 FROM user_section_permissions usp
               JOIN archive_sections s ON s.id = usp.section_id
               WHERE usp.user_id=? AND s.section_code=? AND usp.can_stamp=1
               LIMIT 1""",
            (uid, section_code)
        ).fetchone()
        can_sign = bool(perm)
    conn.close()

    if not can_sign:
        return redirect(url_for('documents_page'))

    payload = _build_sign_page_payload(dict(doc), None)
    return render_template("sign_document.html", **payload)


@app.route("/api/signature/unlock", methods=["POST"])
@login_required
def api_signature_unlock():
    data = request.json or {}
    sign_password = data.get("sign_password", "")
    if not sign_password:
        return jsonify({"success": False, "error": "الرقم السري مطلوب"}), 400
    if not verify_sign_password(session['user_id'], sign_password):
        return jsonify({"success": False, "error": "الرقم السري غير صحيح"}), 401
    return jsonify({"success": True})


@app.route("/api/signature-assets/<int:asset_id>/visibility", methods=["POST"])
@login_required
def api_signature_asset_visibility(asset_id):
    data = request.json or {}
    visibility_scope = str(data.get('visibility_scope') or 'self').strip().lower()
    if visibility_scope not in ('all', 'managers', 'specific', 'self'):
        return jsonify({"success": False, "error": "نطاق الظهور غير صالح"}), 400

    visible_to_user_id = _safe_int(data.get('visible_to_user_id'), 0)

    conn = get_db()
    asset = conn.execute(
        """
        SELECT sa.id, sa.user_id, sa.signature_name, sa.signature_path,
             COALESCE(sa.visibility_scope, 'self') AS visibility_scope,
               sa.visible_to_user_id,
               u.name AS owner_name
        FROM signature_assets sa
        LEFT JOIN users u ON u.id = sa.user_id
        WHERE sa.id=?
        LIMIT 1
        """,
        (asset_id,)
    ).fetchone()

    if not asset:
        conn.close()
        return jsonify({"success": False, "error": "التوقيع غير موجود"}), 404

    viewer_role = str(session.get('user_role') or '').lower()
    viewer_id = session.get('user_id')
    if asset['user_id'] != viewer_id and viewer_role not in ('admin', 'sys_admin'):
        conn.close()
        return jsonify({"success": False, "error": "ليس لديك صلاحية تعديل هذا التوقيع"}), 403

    target_user = None
    if visibility_scope == 'self':
        visible_to_user_id = None
    elif visibility_scope == 'specific':
        if visible_to_user_id <= 0:
            conn.close()
            return jsonify({"success": False, "error": "اختر المستخدم المحدد لعرض التوقيع"}), 400
        target_user = conn.execute(
            "SELECT id, name, role FROM users WHERE id=? AND is_active=1 LIMIT 1",
            (visible_to_user_id,)
        ).fetchone()
        if not target_user:
            conn.close()
            return jsonify({"success": False, "error": "المستخدم المحدد غير موجود"}), 400
        if str(target_user['role'] or '').lower() not in ('manager', 'admin', 'sys_admin'):
            conn.close()
            return jsonify({"success": False, "error": "يمكن التخصيص فقط للمدير أو المدير العام"}), 400
    else:
        visible_to_user_id = None

    conn.execute(
        "UPDATE signature_assets SET visibility_scope=?, visible_to_user_id=? WHERE id=?",
        (visibility_scope, visible_to_user_id, asset_id)
    )
    conn.commit()
    conn.close()

    visibility_label = {
        'all': 'جميع المستخدمين',
        'managers': 'المدراء والمدير العام',
        'specific': f"مستخدم محدد: {target_user['name']}" if target_user else 'مستخدم محدد',
        'self': 'لا أحد',
    }.get(visibility_scope, 'جميع المستخدمين')

    return jsonify({
        "success": True,
        "asset_id": asset_id,
        "visibility_scope": visibility_scope,
        "visible_to_user_id": visible_to_user_id,
        "visibility_label": visibility_label
    })


@app.route("/api/stamp-assets/<int:asset_id>/visibility", methods=["POST"])
@login_required
def api_stamp_asset_visibility(asset_id):
    data = request.json or {}
    visibility_scope = str(data.get('visibility_scope') or 'self').strip().lower()
    if visibility_scope not in ('all', 'managers', 'specific', 'self'):
        return jsonify({"success": False, "error": "نطاق الظهور غير صالح"}), 400

    visible_to_user_id = _safe_int(data.get('visible_to_user_id'), 0)

    conn = get_db()
    asset = conn.execute(
        "SELECT sa.id, sa.user_id, sa.stamp_name, u.name AS owner_name FROM stamp_assets sa "
        "LEFT JOIN users u ON u.id = sa.user_id WHERE sa.id=? LIMIT 1",
        (asset_id,)
    ).fetchone()

    if not asset:
        conn.close()
        return jsonify({"success": False, "error": "الختم غير موجود"}), 404

    viewer_role = str(session.get('user_role') or '').lower()
    viewer_id = session.get('user_id')
    if asset['user_id'] != viewer_id and viewer_role not in ('admin', 'sys_admin'):
        conn.close()
        return jsonify({"success": False, "error": "ليس لديك صلاحية تعديل هذا الختم"}), 403

    target_user = None
    if visibility_scope == 'self':
        visible_to_user_id = None
    elif visibility_scope == 'specific':
        if visible_to_user_id <= 0:
            conn.close()
            return jsonify({"success": False, "error": "اختر المستخدم المحدد"}), 400
        target_user = conn.execute(
            "SELECT id, name, role FROM users WHERE id=? AND is_active=1 LIMIT 1",
            (visible_to_user_id,)
        ).fetchone()
        if not target_user:
            conn.close()
            return jsonify({"success": False, "error": "المستخدم غير موجود"}), 400
        if str(target_user['role'] or '').lower() not in ('manager', 'admin', 'sys_admin'):
            conn.close()
            return jsonify({"success": False, "error": "يمكن التخصيص فقط للمدير أو المدير العام"}), 400
    else:
        visible_to_user_id = None

    conn.execute(
        "UPDATE stamp_assets SET visibility_scope=?, visible_to_user_id=? WHERE id=?",
        (visibility_scope, visible_to_user_id, asset_id)
    )
    conn.commit()
    conn.close()

    visibility_label = {
        'all': 'جميع المستخدمين',
        'managers': 'المدراء والمدير العام',
        'specific': f"مستخدم محدد: {target_user['name']}" if target_user else 'مستخدم محدد',
        'self': 'لا أحد',
    }.get(visibility_scope, 'جميع المستخدمين')

    return jsonify({
        "success": True,
        "asset_id": asset_id,
        "visibility_scope": visibility_scope,
        "visible_to_user_id": visible_to_user_id,
        "visibility_label": visibility_label
    })


@app.route("/api/users/<int:user_id>/stamp-visibility", methods=["POST"])
@login_required
def api_text_stamp_visibility(user_id):
    """Update visibility for text-based stamp (stamp_text column)."""
    data = request.json or {}
    visibility_scope = str(data.get('visibility_scope') or 'self').strip().lower()
    if visibility_scope not in ('all', 'managers', 'specific', 'self'):
        return jsonify({"success": False, "error": "نطاق الظهور غير صالح"}), 400

    visible_to_user_id = _safe_int(data.get('visible_to_user_id'), 0)

    conn = get_db()
    user = conn.execute(
        "SELECT id, name, role, stamp_text FROM users WHERE id=? LIMIT 1",
        (user_id,)
    ).fetchone()

    if not user:
        conn.close()
        return jsonify({"success": False, "error": "المستخدم غير موجود"}), 404

    if not user['stamp_text']:
        conn.close()
        return jsonify({"success": False, "error": "لا يوجد ختم نصي لهذا المستخدم"}), 400

    viewer_role = str(session.get('user_role') or '').lower()
    viewer_id = session.get('user_id')
    if user['id'] != viewer_id and viewer_role not in ('admin', 'sys_admin'):
        conn.close()
        return jsonify({"success": False, "error": "ليس لديك صلاحية تعديل هذا الختم"}), 403

    target_user = None
    if visibility_scope == 'self':
        visible_to_user_id = None
    elif visibility_scope == 'specific':
        if visible_to_user_id <= 0:
            conn.close()
            return jsonify({"success": False, "error": "اختر المستخدم المحدد"}), 400
        target_user = conn.execute(
            "SELECT id, name, role FROM users WHERE id=? AND is_active=1 LIMIT 1",
            (visible_to_user_id,)
        ).fetchone()
        if not target_user:
            conn.close()
            return jsonify({"success": False, "error": "المستخدم غير موجود"}), 400
        if str(target_user['role'] or '').lower() not in ('manager', 'admin', 'sys_admin'):
            conn.close()
            return jsonify({"success": False, "error": "يمكن التخصيص فقط للمدير أو المدير العام"}), 400
    else:
        visible_to_user_id = None

    conn.execute(
        "UPDATE users SET stamp_visibility_scope=?, stamp_visible_to_user_id=? WHERE id=?",
        (visibility_scope, visible_to_user_id, user_id)
    )
    conn.commit()
    conn.close()

    visibility_label = {
        'all': 'جميع المستخدمين',
        'managers': 'المدراء والمدير العام',
        'specific': f"مستخدم محدد: {target_user['name']}" if target_user else 'مستخدم محدد',
        'self': 'لا أحد',
    }.get(visibility_scope, 'جميع المستخدمين')

    return jsonify({
        "success": True,
        "user_id": user_id,
        "visibility_scope": visibility_scope,
        "visible_to_user_id": visible_to_user_id,
        "visibility_label": visibility_label
    })


@app.route("/api/stamp-preview")
@login_required
def api_stamp_preview():
    owner_id = _safe_int(request.args.get('owner_id'), 0)
    mode = str(request.args.get('mode') or '').strip().lower()
    force_text_mode = mode in ('text', 'template', 'dynamic', '1', 'true')
    if owner_id <= 0:
        return _svg_response(_svg_placeholder(["لا يوجد ختم"], width=200, height=200, subtitle="ختم"), status=400)

    conn = get_db()
    visible_assets = _get_stamp_assets(conn, viewer_id=session.get('user_id'), viewer_role=session.get('user_role'))
    allowed_owner_ids = { _safe_int(a.get('owner_id'), 0) for a in visible_assets }
    if owner_id not in allowed_owner_ids:
        conn.close()
        return jsonify({'error': 'غير مصرح بعرض هذا الختم'}), 403

    stamp_asset = conn.execute(
        """
        SELECT stamp_path
        FROM stamp_assets
        WHERE user_id=? AND is_active=1 AND COALESCE(stamp_path,'')!=''
        ORDER BY id DESC
        LIMIT 1
        """,
        (owner_id,)
    ).fetchone()
    user_row = conn.execute(
        """
        SELECT u.name, u.stamp_text, u.role, u.employee_id,
               s.section_code
        FROM users u
        LEFT JOIN archive_sections s ON s.id = u.archive_section_id
        WHERE u.id=? LIMIT 1
        """,
        (owner_id,)
    ).fetchone()
    conn.close()

    if (not force_text_mode) and stamp_asset and stamp_asset['stamp_path']:
        resolved = _resolve_document_file_path(stamp_asset['stamp_path'])
        if resolved and os.path.exists(resolved):
            return send_file(resolved)

    if user_row:
        u_role = str(user_row['role'] or '').strip().lower()
        if u_role in ('manager', 'admin', 'sys_admin'):
            sec = str(user_row['section_code'] or '').strip().upper()
            emp = str(user_row['employee_id'] or '').strip()
            tpl_conn = get_db()
            tpl_row = _get_default_stamp_tpl(tpl_conn)
            tpl_conn.close()
            if tpl_row:
                tpl_file = _render_stamp_from_tpl(
                    tpl_row['file_name'], tpl_row['text_x_ratio'], tpl_row['text_y_ratio'],
                    sec, emp, '', as_file=True
                )
                if tpl_file:
                    return send_file(tpl_file, mimetype='image/png')

    return jsonify({'error': 'لا يوجد ختم'}), 404

# ════════════════════════════════════════════
# DASHBOARD
# ════════════════════════════════════════════

@app.route("/dashboard")
@login_required
def dashboard():
    conn  = get_db()
    uid   = session['user_id']
    stats = {
        "total_docs":    conn.execute("SELECT COUNT(*) FROM documents WHERE created_by=?", (uid,)).fetchone()[0],
        "pending_signs": conn.execute("SELECT COUNT(*) FROM signature_requests WHERE requested_from=? AND status='pending'", (uid,)).fetchone()[0],
        "approved":      conn.execute("SELECT COUNT(*) FROM documents WHERE status='approved'").fetchone()[0],
        "pending_docs":  conn.execute("SELECT COUNT(*) FROM documents WHERE status='pending'").fetchone()[0],
    }
    conn.close()
    return render_template("dashboard.html", stats=stats, user=dict(get_user_by_id(uid)))

@app.route("/documents")
@login_required
def documents_page():
    return render_template("documents.html")

@app.route("/archive")
@login_required
def archive_page():
    # فقط المديرين يقدرون يشوفوا صفحة الأرشيف
    role = (session.get('user_role') or '').lower()
    if role not in ('admin', 'sys_admin', 'manager'):
        return redirect(url_for('dashboard'))
    return render_template("archive.html")

@app.route("/admin/panel")
@login_required
@role_required("admin", "sys_admin")
def admin_panel():
    return render_template("admin/panel.html")

@app.route("/admin/logs")
@login_required
@role_required("admin", "sys_admin")
def admin_logs():
    return render_template("admin/logs.html")


@app.route("/admin/templates")
@login_required
@role_required("admin", "sys_admin")
def admin_templates_page():
    return render_template("admin/templates.html")

# ════════════════════════════════════════════
# إدارة المستخدمين
# ════════════════════════════════════════════

@app.route("/admin/users")
@login_required
@role_required("admin", "sys_admin")
def admin_users():
    return render_template("admin/users.html")

@app.route("/api/users", methods=["GET"])
@login_required
@role_required("admin", "sys_admin")
def api_get_users():
    return jsonify(get_all_users())


@app.route("/api/admin/self-registered-count", methods=["GET"])
@login_required
@role_required("admin", "sys_admin")
def api_self_registered_count():
    """عدد الموظفين الذين سجّلوا بأنفسهم ولم يُراجَعوا بعد."""
    conn = get_db()
    count = conn.execute(
        "SELECT COUNT(*) FROM users WHERE created_by IS NULL AND is_active=1"
    ).fetchone()[0]
    conn.close()
    return jsonify({"count": count})

@app.route("/api/users/mentions", methods=["GET"])
@login_required
def api_users_mentions():
    conn = get_db()
    rows = conn.execute(
        "SELECT id, name, email, role FROM users WHERE is_active=1 ORDER BY name ASC"
    ).fetchall()
    conn.close()
    return jsonify([dict(r) for r in rows])

@app.route("/api/archive-sections", methods=["GET"])
@login_required
def api_archive_sections():
    role = (session.get('user_role') or '').strip().lower()
    if role in ('admin', 'sys_admin'):
        return jsonify(get_all_archive_sections())

    if role == 'manager':
        conn = get_db()
        user_id = session.get('user_id')
        main_code_row = conn.execute(
            "SELECT UPPER(TRIM(COALESCE(s.section_code, 'GN'))) AS section_code FROM users u LEFT JOIN archive_sections s ON s.id=u.archive_section_id WHERE u.id=? LIMIT 1",
            (user_id,)
        ).fetchone()
        main_code = str((main_code_row['section_code'] if main_code_row else 'GN') or 'GN').strip().upper() or 'GN'

        perm_rows = conn.execute(
            """
            SELECT UPPER(TRIM(COALESCE(s.section_code, ''))) AS section_code
            FROM user_section_permissions usp
            JOIN archive_sections s ON s.id = usp.section_id
            WHERE usp.user_id=? AND s.is_active=1 AND (usp.can_view_archive=1 OR usp.can_stamp=1)
            """,
            (user_id,)
        ).fetchall()

        allowed_codes = {main_code}
        for row in perm_rows:
            code = str(row['section_code'] or '').strip().upper()
            if code:
                allowed_codes.add(code)

        placeholders = ','.join('?' * len(allowed_codes))
        rows = conn.execute(
            f"""
            SELECT id, section_name, section_code, is_active
            FROM archive_sections
            WHERE is_active=1 AND UPPER(COALESCE(section_code, '')) IN ({placeholders})
            ORDER BY section_name ASC
            """,
            tuple(sorted(allowed_codes))
        ).fetchall()
        conn.close()
        return jsonify([dict(r) for r in rows])

    return jsonify([])

@app.route("/api/archive-sections", methods=["POST"])
@login_required
@role_required("admin", "sys_admin")
def api_create_archive_section():
    data = request.json or {}
    section_name = data.get("section_name", "")
    section_code = data.get("section_code", "")
    ok, section, err = create_archive_section(section_name, section_code, session.get('user_id'))
    if not ok:
        return jsonify({"success": False, "error": err}), 400
    log_action(
        session['user_id'],
        "CREATE_ARCHIVE_SECTION",
        f"{section.get('section_name')} ({section.get('section_code')})",
        ip=request.remote_addr,
        user_agent=request.user_agent.string
    )
    return jsonify({"success": True, "section": section})


@app.route("/api/admin/cleanup-local-docs", methods=["POST"])
@login_required
@role_required("admin", "sys_admin")
def api_admin_cleanup_local_docs():
    """يحذف الملفات المحلية للوثائق التي رُفعت بنجاح للأرشيف (archive_storage_path موجود)."""
    conn = get_db()
    rows = conn.execute(
        "SELECT id, file_path, archive_storage_path FROM documents WHERE file_path IS NOT NULL AND archive_storage_path IS NOT NULL"
    ).fetchall()
    deleted = 0
    errors = 0
    for row in rows:
        local_path = _resolve_document_file_path(row['file_path'])
        if local_path and os.path.exists(local_path):
            try:
                os.remove(local_path)
                deleted += 1
            except OSError:
                errors += 1
        conn.execute("UPDATE documents SET file_path=NULL WHERE id=?", (row['id'],))
    conn.commit()
    conn.close()
    return jsonify({"success": True, "deleted_files": deleted, "errors": errors})


@app.route("/api/admin/activity")
@login_required
@role_required("admin", "sys_admin")
def api_admin_activity():
    limit = max(1, min(_safe_int(request.args.get('limit', 300), 300), 1000))
    conn = get_db()
    rows = conn.execute(
        """
        SELECT a.*, u.name AS user_name
        FROM audit_log a
        LEFT JOIN users u ON u.id = a.user_id
        ORDER BY a.id DESC
        LIMIT ?
        """,
        (limit,)
    ).fetchall()
    conn.close()
    return jsonify({"logs": [dict(r) for r in rows]})


@app.route("/api/admin/logins")
@login_required
@role_required("admin", "sys_admin")
def api_admin_logins():
    limit = max(1, min(_safe_int(request.args.get('limit', 300), 300), 1000))
    conn = get_db()
    rows = conn.execute(
        """
        SELECT l.*, u.name AS user_name
        FROM login_history l
        LEFT JOIN users u ON u.id = l.user_id
        ORDER BY l.id DESC
        LIMIT ?
        """,
        (limit,)
    ).fetchall()
    conn.close()
    return jsonify({"logs": [dict(r) for r in rows]})


@app.route("/api/admin/security-alerts")
@login_required
@role_required("admin", "sys_admin")
def api_admin_security_alerts():
    conn = get_db()

    pending_devices = conn.execute(
        "SELECT COUNT(*) FROM trusted_devices WHERE trusted=0"
    ).fetchone()[0]
    new_devices_today = conn.execute(
        "SELECT COUNT(*) FROM login_history WHERE event='new_device' AND DATE(timestamp)=DATE('now')"
    ).fetchone()[0]
    failed_auth_today = conn.execute(
        "SELECT COUNT(*) FROM login_history WHERE event IN ('failed','otp_fail') AND DATE(timestamp)=DATE('now')"
    ).fetchone()[0]
    blocked_accounts_today = conn.execute(
        "SELECT COUNT(*) FROM audit_log WHERE action='ACCOUNT_BLOCKED' AND DATE(timestamp)=DATE('now')"
    ).fetchone()[0]

    alerts = []

    failed_rows = conn.execute(
        """
        SELECT l.user_id, u.name AS user_name, l.ip_address, l.event, l.timestamp
        FROM login_history l
        LEFT JOIN users u ON u.id = l.user_id
        WHERE l.event IN ('failed','otp_fail')
        ORDER BY l.id DESC
        LIMIT 120
        """
    ).fetchall()
    for row in failed_rows:
        event_label = 'كلمة مرور خاطئة' if row['event'] == 'failed' else 'OTP خاطئ'
        alerts.append({
            "title": "محاولة دخول فاشلة",
            "user_name": row['user_name'] or 'مجهول',
            "details": f"{event_label} • IP: {row['ip_address'] or '—'}",
            "severity": "high",
            "occurred_at": row['timestamp'],
        })

    new_device_rows = conn.execute(
        """
        SELECT l.user_id, u.name AS user_name, l.ip_address, l.device_type, l.os_name, l.browser, l.timestamp
        FROM login_history l
        LEFT JOIN users u ON u.id = l.user_id
        WHERE l.event='new_device'
        ORDER BY l.id DESC
        LIMIT 120
        """
    ).fetchall()
    for row in new_device_rows:
        details = f"{row['device_type'] or 'جهاز'} {row['os_name'] or ''} {row['browser'] or ''} • IP: {row['ip_address'] or '—'}"
        alerts.append({
            "title": "محاولة دخول من جهاز جديد",
            "user_name": row['user_name'] or 'مجهول',
            "details": details.strip(),
            "severity": "medium",
            "occurred_at": row['timestamp'],
        })

    blocked_rows = conn.execute(
        """
        SELECT a.user_id, u.name AS user_name, a.details, a.timestamp
        FROM audit_log a
        LEFT JOIN users u ON u.id = a.user_id
        WHERE a.action='ACCOUNT_BLOCKED'
        ORDER BY a.id DESC
        LIMIT 120
        """
    ).fetchall()
    for row in blocked_rows:
        alerts.append({
            "title": "تم تجميد حساب",
            "user_name": row['user_name'] or 'مجهول',
            "details": row['details'] or 'تم التجميد لأسباب أمنية',
            "severity": "critical",
            "occurred_at": row['timestamp'],
        })

    pending_device_rows = conn.execute(
        """
        SELECT td.user_id, u.name AS user_name, td.device_name, td.ip_address, td.created_at
        FROM trusted_devices td
        LEFT JOIN users u ON u.id = td.user_id
        WHERE td.trusted=0
        ORDER BY td.id DESC
        LIMIT 120
        """
    ).fetchall()
    for row in pending_device_rows:
        alerts.append({
            "title": "جهاز بانتظار التأكيد",
            "user_name": row['user_name'] or 'مجهول',
            "details": f"{row['device_name'] or 'جهاز'} • IP: {row['ip_address'] or '—'}",
            "severity": "medium",
            "occurred_at": row['created_at'],
        })

    conn.close()

    alerts.sort(key=lambda item: str(item.get('occurred_at') or ''), reverse=True)
    return jsonify({
        "alerts": alerts[:400],
        "summary": {
            "pending_devices": pending_devices,
            "new_devices_today": new_devices_today,
            "failed_auth_today": failed_auth_today,
            "blocked_accounts_today": blocked_accounts_today,
        }
    })

@app.route("/api/users", methods=["POST"])
@login_required
@role_required("admin", "sys_admin")
def api_create_user():
    data      = request.json
    name      = data.get("name","").strip()
    email     = data.get("email","").strip().lower()
    phone     = data.get("phone","").strip()
    job_title = data.get("job_title","").strip()
    role      = data.get("role","user")

    if not all([name, email, phone, role]):
        return jsonify({"success": False, "error": "جميع الحقول مطلوبة"})

    section_id_raw = data.get("archive_section_id")
    section_id = int(section_id_raw) if str(section_id_raw or '').strip().isdigit() else None
    employee_id = str(data.get("employee_id") or '').strip() or None
    section_permissions = data.get("section_permissions") or []

    ok, temp_pass, err, new_uid = create_user(name, email, phone, job_title, role, session['user_id'],
                                               archive_section_id=section_id, employee_id=employee_id)
    if ok:
        conn = get_db()
        auto_stamp_created = False
        try:
            _replace_user_section_permissions(conn, new_uid, section_permissions, session['user_id'])

            # Auto-create default-template stamp for manager-level accounts.
            role_value = str(role or '').strip().lower()
            if role_value in ('manager', 'admin', 'sys_admin'):
                auto_stamp_created = _create_default_stamp_asset_for_user(new_uid, conn=conn)

            conn.commit()
        finally:
            conn.close()

        log_action(session['user_id'], "CREATE_USER", f"{name} — {email}",
                   ip=request.remote_addr, user_agent=request.user_agent.string)
        result = {"success": True, "user_id": new_uid, "auto_stamp_created": auto_stamp_created,
                  "temp_password": temp_pass,
                  "message": f"تم إنشاء الحساب بنجاح"}
        if str(role or '').strip().lower() in ('manager', 'admin', 'sys_admin') and not auto_stamp_created:
            result["warning"] = "لم يتم إنشاء الختم التلقائي لعدم توفر قالب ختم افتراضي"
        return jsonify(result)
    return jsonify({"success": False, "error": err})

@app.route("/api/users/<int:uid>/permanent-delete", methods=["DELETE"])
@login_required
@role_required("admin", "sys_admin")
def api_permanent_delete_user(uid):
    """حذف نهائي من قاعدة البيانات."""
    if uid == session['user_id']:
        return jsonify({"success": False, "error": "لا تستطيع حذف حسابك"}), 400
    conn = get_db()
    try:
        # حذف البيانات المرتبطة أولاً (ترتيب مهم بسبب Foreign Keys)
        conn.execute("DELETE FROM trusted_devices WHERE user_id=?", (uid,))
        conn.execute("DELETE FROM otp_codes WHERE user_id=?", (uid,))
        conn.execute("DELETE FROM user_section_permissions WHERE user_id=?", (uid,))
        conn.execute("DELETE FROM stamp_assets WHERE user_id=?", (uid,))
        conn.execute("DELETE FROM signature_assets WHERE user_id=?", (uid,))
        conn.execute("DELETE FROM users WHERE id=?", (uid,))
        conn.commit()
        return jsonify({"success": True})
    except Exception as e:
        conn.rollback()
        return jsonify({"success": False, "error": str(e)}), 500
    finally:
        conn.close()

@app.route("/api/users/<int:uid>", methods=["DELETE"])
@login_required
@role_required("admin", "sys_admin")
def api_delete_user(uid):
    """إيقاف (تعطيل) الحساب."""
    if uid == session['user_id']:
        return jsonify({"success": False, "error": "لا تستطيع إيقاف حسابك"}), 400
    conn = get_db()
    conn.execute("UPDATE users SET is_active=0 WHERE id=?", (uid,))
    conn.commit(); conn.close()
    log_action(session['user_id'], "DEACTIVATE_USER", f"user_id={uid}", ip=request.remote_addr)
    return jsonify({"success": True})


@app.route("/api/users/<int:uid>/activate", methods=["POST"])
@login_required
@role_required("admin", "sys_admin")
def api_activate_user(uid):
    """إعادة تفعيل حساب موقوف."""
    if uid == session['user_id']:
        return jsonify({"success": False, "error": "حسابك نشط بالفعل"}), 400
    conn = get_db()
    conn.execute("UPDATE users SET is_active=1 WHERE id=?", (uid,))
    conn.commit(); conn.close()
    log_action(session['user_id'], "ACTIVATE_USER", f"user_id={uid}", ip=request.remote_addr)
    return jsonify({"success": True})


@app.route("/api/users/<int:uid>", methods=["PUT"])
@login_required
@role_required("admin", "sys_admin")
def api_update_user(uid):
    data = request.json
    section_id_raw = data.get("archive_section_id")
    section_id = int(section_id_raw) if str(section_id_raw or '').strip().isdigit() else None
    employee_id = str(data.get("employee_id") or '').strip() or None
    section_permissions = data.get("section_permissions") or []
    conn = get_db()
    conn.execute(
        "UPDATE users SET name=?,phone=?,job_title=?,role=?,stamp_text=?,archive_section_id=?,employee_id=? WHERE id=?",
        (data.get("name"), data.get("phone"), data.get("job_title"),
         data.get("role"), data.get("stamp_text",""), section_id, employee_id, uid)
    )
    _replace_user_section_permissions(conn, uid, section_permissions, session['user_id'])
    _sync_manager_section_stamp_assets(uid, conn=conn)
    conn.commit(); conn.close()
    log_action(session['user_id'], "UPDATE_USER", f"user_id={uid}", ip=request.remote_addr)
    return jsonify({"success": True})


def _replace_user_section_permissions(conn, user_id, permissions, granted_by):
    """Replace all section permissions for a user with validated rows."""
    uid = _safe_int(user_id, 0)
    if uid <= 0:
        return

    conn.execute("DELETE FROM user_section_permissions WHERE user_id=?", (uid,))
    for p in permissions or []:
        sid = _safe_int((p or {}).get("section_id"), 0)
        if sid <= 0:
            continue

        sec = conn.execute(
            "SELECT id FROM archive_sections WHERE id=? AND is_active=1",
            (sid,)
        ).fetchone()
        if not sec:
            continue

        can_view = 1 if (p or {}).get("can_view_archive") else 0
        can_stamp = 1 if (p or {}).get("can_stamp") else 0
        if not can_view and not can_stamp:
            continue

        conn.execute(
            """
            INSERT INTO user_section_permissions (user_id, section_id, can_view_archive, can_stamp, granted_by)
            VALUES (?,?,?,?,?)
            ON CONFLICT(user_id, section_id) DO UPDATE SET
              can_view_archive=excluded.can_view_archive,
              can_stamp=excluded.can_stamp,
              granted_by=excluded.granted_by
            """,
            (uid, sid, can_view, can_stamp, _safe_int(granted_by, 0) or None)
        )


@app.route("/api/users/<int:uid>/section-perms", methods=["GET"])
@login_required
@role_required("admin", "sys_admin")
def api_get_user_section_perms(uid):
    conn = get_db()
    rows = conn.execute(
        """
        SELECT usp.section_id, usp.can_view_archive, usp.can_stamp,
               s.section_name, s.section_code
        FROM user_section_permissions usp
        JOIN archive_sections s ON s.id = usp.section_id
        WHERE usp.user_id = ?
        ORDER BY s.section_name
        """,
        (uid,)
    ).fetchall()
    conn.close()
    return jsonify({"success": True, "permissions": [dict(r) for r in rows]})


@app.route("/api/users/<int:uid>/section-perms", methods=["PUT"])
@login_required
@role_required("admin", "sys_admin")
def api_set_user_section_perms(uid):
    """Replace all section permissions for a user."""
    data = request.json or {}
    perms = data.get("permissions", [])  # [{section_id, can_view_archive, can_stamp}, ...]

    conn = get_db()
    _replace_user_section_permissions(conn, uid, perms, session['user_id'])
    _sync_manager_section_stamp_assets(uid, conn=conn)
    conn.commit(); conn.close()
    log_action(session['user_id'], "UPDATE_SECTION_PERMS", f"user_id={uid}", ip=request.remote_addr)
    return jsonify({"success": True})


# ════════════════════════════════════════════
# الملف الشخصي + الأجهزة
# ════════════════════════════════════════════

@app.route("/profile")
@login_required
def profile():
    user    = get_user_by_id(session['user_id'])
    devices = get_trusted_devices(session['user_id'])
    user_dict = dict(user)
    # Get text stamp visibility
    conn = get_db()
    row = conn.execute(
        "SELECT COALESCE(stamp_visibility_scope, 'self') AS stamp_visibility_scope, stamp_visible_to_user_id FROM users WHERE id=? LIMIT 1",
        (session['user_id'],)
    ).fetchone()
    conn.close()
    if row:
        user_dict['stamp_visibility_scope'] = row['stamp_visibility_scope']
        user_dict['stamp_visible_to_user_id'] = row['stamp_visible_to_user_id']
    else:
        user_dict['stamp_visibility_scope'] = 'self'
        user_dict['stamp_visible_to_user_id'] = None
    return render_template("profile.html", user=user_dict, devices=devices)

@app.route("/api/profile/signature", methods=["POST"])
@login_required
def upload_signature():
    file = request.files.get("signature")
    if not file: return jsonify({"error": "لا يوجد ملف"}), 400
    ext = file.filename.rsplit('.', 1)[-1].lower()
    if ext not in ['png','jpg','jpeg']:
        return jsonify({"error": "PNG أو JPG فقط"}), 400
    fname = f"sig_{session['user_id']}_{int(datetime.utcnow().timestamp())}.{ext}"
    file.save(os.path.join(UPLOAD_FOLDER, 'signatures', fname))
    update_signature(session['user_id'], f"uploads/signatures/{fname}")
    log_action(session['user_id'], "UPLOAD_SIGNATURE", ip=request.remote_addr)
    return jsonify({"success": True, "path": f"uploads/signatures/{fname}"})

@app.route("/api/profile/sign-password", methods=["POST"])
@login_required
def set_sign_password():
    data     = request.json
    new_pass = data.get("sign_password","")
    if len(new_pass) < 6:
        return jsonify({"success": False, "error": "الرقم السري يجب أن يكون 6 أحرف على الأقل"})
    update_sign_password(session['user_id'], new_pass)
    log_action(session['user_id'], "SET_SIGN_PASSWORD", ip=request.remote_addr)
    return jsonify({"success": True})

@app.route("/api/profile/signature-visibility-targets")
@login_required
def profile_signature_visibility_targets():
    conn = get_db()
    targets = _get_signature_visibility_targets(conn)
    conn.close()
    return jsonify({"success": True, "targets": targets})

@app.route("/api/profile/signatures", methods=["GET"])
@login_required
def profile_list_signatures():
    conn = get_db()
    rows = conn.execute(
        "SELECT id, signature_name, signature_path, COALESCE(visibility_scope, 'self') AS visibility_scope, visible_to_user_id, created_at FROM signature_assets "
        "WHERE user_id=? AND is_active=1 ORDER BY id DESC",
        (session['user_id'],)
    ).fetchall()
    conn.close()
    return jsonify({"success": True, "signatures": [dict(r) for r in rows]})

@app.route("/api/profile/signatures", methods=["POST"])
@login_required
def profile_upload_signature():
    file = request.files.get("signature")
    if not file:
        return jsonify({"success": False, "error": "لا يوجد ملف"}), 400
    ext = file.filename.rsplit('.', 1)[-1].lower()
    if ext not in ['png', 'jpg', 'jpeg']:
        return jsonify({"success": False, "error": "PNG أو JPG فقط"}), 400
    fname = f"sig_{session['user_id']}_{int(datetime.utcnow().timestamp())}.{ext}"
    fpath = os.path.join(UPLOAD_FOLDER, 'signatures', fname)
    file.save(fpath)
    sig_name = request.form.get("name", "").strip() or None
    conn = get_db()
    conn.execute(
        "INSERT INTO signature_assets (user_id, signature_name, signature_path, visibility_scope, is_active) VALUES (?,?,?,?,1)",
        (session['user_id'], sig_name, f"uploads/signatures/{fname}", 'self')
    )
    conn.commit(); conn.close()
    log_action(session['user_id'], "UPLOAD_SIGNATURE", ip=request.remote_addr)
    return jsonify({"success": True, "path": f"uploads/signatures/{fname}"})

@app.route("/api/profile/signatures/<int:sig_id>", methods=["DELETE"])
@login_required
def profile_delete_signature(sig_id):
    conn = get_db()
    row = conn.execute(
        "SELECT id, signature_path FROM signature_assets WHERE id=? AND user_id=? AND is_active=1",
        (sig_id, session['user_id'])
    ).fetchone()
    if not row:
        conn.close()
        return jsonify({"success": False, "error": "التوقيع غير موجود"}), 404
    conn.execute("UPDATE signature_assets SET is_active=0 WHERE id=?", (sig_id,))
    conn.commit(); conn.close()
    log_action(session['user_id'], "DELETE_SIGNATURE", details=str(sig_id), ip=request.remote_addr)
    return jsonify({"success": True})

@app.route("/api/profile/stamps", methods=["GET"])
@login_required
def profile_list_stamps():
    conn = get_db()
    rows = conn.execute(
        "SELECT id, stamp_name, stamp_path, COALESCE(visibility_scope, 'self') AS visibility_scope, visible_to_user_id, created_at FROM stamp_assets "
        "WHERE user_id=? AND is_active=1 ORDER BY id DESC",
        (session['user_id'],)
    ).fetchall()
    conn.close()
    return jsonify({"success": True, "stamps": [dict(r) for r in rows]})

@app.route("/api/profile/stamps", methods=["POST"])
@login_required
def profile_upload_stamp():
    file = request.files.get("stamp")
    if not file:
        return jsonify({"success": False, "error": "لا يوجد ملف"}), 400
    ext = file.filename.rsplit('.', 1)[-1].lower()
    if ext not in ['png', 'jpg', 'jpeg', 'webp']:
        return jsonify({"success": False, "error": "PNG أو JPG أو WebP فقط"}), 400
    fname = f"stamp_{session['user_id']}_{int(datetime.utcnow().timestamp())}.{ext}"
    fpath = os.path.join(UPLOAD_FOLDER, 'stamps', fname)
    os.makedirs(os.path.dirname(fpath), exist_ok=True)
    file.save(fpath)
    stamp_name = request.form.get("name", "").strip() or None
    conn = get_db()
    conn.execute(
        "INSERT INTO stamp_assets (user_id, stamp_name, stamp_path, visibility_scope, is_active) VALUES (?,?,?,?,1)",
        (session['user_id'], stamp_name, f"uploads/stamps/{fname}", 'self')
    )
    conn.commit(); conn.close()
    log_action(session['user_id'], "UPLOAD_STAMP", ip=request.remote_addr)
    return jsonify({"success": True, "path": f"uploads/stamps/{fname}"})

@app.route("/api/profile/stamps/<int:stamp_id>", methods=["DELETE"])
@login_required
def profile_delete_stamp(stamp_id):
    conn = get_db()
    row = conn.execute(
        "SELECT id FROM stamp_assets WHERE id=? AND user_id=? AND is_active=1",
        (stamp_id, session['user_id'])
    ).fetchone()
    if not row:
        conn.close()
        return jsonify({"success": False, "error": "الختم غير موجود"}), 404
    conn.execute("UPDATE stamp_assets SET is_active=0 WHERE id=?", (stamp_id,))
    conn.commit(); conn.close()
    log_action(session['user_id'], "DELETE_STAMP", details=str(stamp_id), ip=request.remote_addr)
    return jsonify({"success": True})

@app.route("/api/devices/<int:device_id>", methods=["DELETE"])
@login_required
def remove_device(device_id):
    """المستخدم يحذف جهاز من أجهزته الموثوقة"""
    conn = get_db()
    conn.execute(
        "DELETE FROM trusted_devices WHERE id=? AND user_id=?",
        (device_id, session['user_id'])
    )
    conn.commit(); conn.close()
    return jsonify({"success": True})

@app.route("/api/me")
@login_required
def api_me():
    user = get_user_by_id(session['user_id'])
    return jsonify({
        "id":             user['id'],
        "name":           user['name'],
        "email":          user['email'],
        "role":           user['role'],
        "job_title":      user['job_title'],
        "signature_path": user['signature_path'],
        "stamp_text":     user['stamp_text'],
        "can_stamp":      can_use_stamp(user['id'])
    })

# ════════════════════════════════════════════
# SCHEDULER
# ════════════════════════════════════════════

def rotate_admin_passwords():
    conn   = get_db()
    admins = conn.execute("SELECT * FROM users WHERE role='admin' AND is_active=1").fetchall()
    for admin in admins:
        new_pass = secrets.token_hex(4).upper()
        update_sign_password(admin['id'], new_pass)
        send_weekly_password(admin['email'], admin['name'], new_pass)
        log_action(admin['id'], "WEEKLY_PASSWORD_ROTATED")
        print(f"[NAJM] تم تجديد رقم {admin['name']}")
    conn.close()

# ════════════════════════════════════════════
# INIT
# ════════════════════════════════════════════

@app.route("/api/my-pending-signatures")
@login_required
def api_pending_signatures():
    conn = get_db()
    rows = conn.execute("""
        SELECT sr.*, d.title as doc_title, u.name as requester_name
        FROM signature_requests sr
        LEFT JOIN documents d ON sr.document_id = d.id
        LEFT JOIN users u ON sr.requested_by = u.id
        WHERE sr.requested_from=? AND sr.status='pending'
        ORDER BY sr.id DESC
    """, (session['user_id'],)).fetchall()
    conn.close()
    return jsonify([dict(r) for r in rows])

@app.route("/api/documents", methods=["GET", "POST"])
@login_required
def api_documents():
    # ── POST: إنشاء وثيقة جديدة ──────────────────────────────────────
    if request.method == "POST":
        role = (session.get('user_role') or '').strip().lower()
        current_user = get_user_by_id(session.get('user_id'))

        title = (request.form.get('title') or '').strip()
        if not title:
            return jsonify({"success": False, "error": "عنوان الوثيقة مطلوب"}), 400

        doc_mode          = (request.form.get('doc_mode') or 'upload').strip()
        archive_section_id_raw = (request.form.get('archive_section_id') or '').strip()
        archive_section   = (request.form.get('archive_section') or 'عام').strip()
        archive_section_code = (request.form.get('archive_section_code') or 'GN').strip().upper()

        main_section_name = str((current_user['archive_section'] if current_user else '') or 'عام').strip()
        main_section_code = str((current_user['archive_section_code'] if current_user else '') or 'GN').strip().upper() or 'GN'

        # Resolve selected section from DB when available.
        selected_section = None
        selected_id = _safe_int(archive_section_id_raw, 0)
        if selected_id > 0:
            _conn_sel = get_db()
            selected_section = _conn_sel.execute(
                "SELECT id, section_name, section_code FROM archive_sections WHERE id=? AND is_active=1 LIMIT 1",
                (selected_id,)
            ).fetchone()
            _conn_sel.close()

        # Admin and sys_admin can save to any active section.
        if role in ('admin', 'sys_admin'):
            if selected_section:
                archive_section = str(selected_section['section_name'] or archive_section or 'عام').strip()
                archive_section_code = str(selected_section['section_code'] or archive_section_code or 'GN').strip().upper() or 'GN'
        # Manager can save only to his scoped sections (can_view_archive + main section).
        elif role == 'manager':
            conn_perms = get_db()
            perm_rows = conn_perms.execute(
                """
                SELECT UPPER(TRIM(COALESCE(s.section_code, ''))) AS section_code
                FROM user_section_permissions usp
                JOIN archive_sections s ON s.id = usp.section_id
                WHERE usp.user_id=? AND s.is_active=1 AND (usp.can_view_archive=1 OR usp.can_stamp=1)
                """,
                (session.get('user_id'),)
            ).fetchall()
            conn_perms.close()

            allowed_codes = {main_section_code}
            for row in perm_rows:
                code = str(row['section_code'] or '').strip().upper()
                if code:
                    allowed_codes.add(code)

            if selected_section and str(selected_section['section_code'] or '').strip().upper() in allowed_codes:
                archive_section = str(selected_section['section_name'] or main_section_name).strip() or main_section_name
                archive_section_code = str(selected_section['section_code'] or main_section_code).strip().upper() or main_section_code
            else:
                archive_section = main_section_name
                archive_section_code = main_section_code
        # Employee/user always saves to main section.
        else:
            archive_section = main_section_name
            archive_section_code = main_section_code

        note              = (request.form.get('note') or '').strip()
        text_template_name = (request.form.get('text_template_name') or '').strip() or None
        # إذا لم يُحدَّد قالب، استخدم القالب الافتراضي
        if not text_template_name:
            text_template_name = _get_default_docx_template_name() or None

        try:
            mentions = json.loads(request.form.get('mentions') or '[]')
            mentions = [int(m) for m in mentions if m]
        except Exception:
            mentions = []

        file_path = None
        archive_storage_path = None
        content_json = None
        archived_at = None
        document_fingerprint = None
        upload_blob = None
        upload_ext = ''

        if doc_mode == 'upload':
            file = request.files.get('file')
            if not file or not file.filename:
                return jsonify({"success": False, "error": "اختر ملفاً للرفع"}), 400
            ext = file.filename.rsplit('.', 1)[-1].lower()
            if ext not in ('pdf', 'png', 'jpg', 'jpeg', 'docx', 'doc', 'xlsx', 'xls'):
                return jsonify({"success": False, "error": "نوع الملف غير مدعوم"}), 400
            try:
                file_blob = file.read()
            except Exception:
                file_blob = b''
            if not file_blob:
                return jsonify({"success": False, "error": "تعذر قراءة الملف المرفوع"}), 400
            document_fingerprint = _sha256_hex_from_bytes(file_blob)
            upload_blob = file_blob
            upload_ext = ext
        else:
            content_html = (request.form.get('content') or '').strip()
            if not content_html:
                return jsonify({"success": False, "error": "أدخل محتوى الوثيقة"}), 400
            content_json = json.dumps({"html": content_html}, ensure_ascii=False)
            document_fingerprint = _build_text_document_fingerprint(content_html, text_template_name)

        conn = get_db()
        duplicate_doc = _find_approved_duplicate_document(conn, document_fingerprint)
        if duplicate_doc:
            conn.rollback()
            conn.close()
            duplicate_ref = str(duplicate_doc.get('archive_number') or duplicate_doc.get('id') or '').strip()
            duplicate_label = duplicate_ref or 'وثيقة معتمدة'
            return jsonify({
                "success": False,
                "error": f"لا يمكن رفع نفس المستند مرة أخرى لأنه معتمد مسبقاً برقم {duplicate_label}"
            }), 409

        # ─── توليد رقم أرشفة ───────────────────────────────────────
        _now = datetime.utcnow()
        _year  = _now.year % 100   # آخر خانتين من السنة (مثلاً 2026 → 26)
        _month = _now.month
        _prefix = f"{archive_section_code}{_month:02d}{_year:02d}"
        last = conn.execute(
            "SELECT archive_number FROM documents "
            "WHERE archive_number LIKE ? ORDER BY id DESC LIMIT 1",
            (f"{_prefix}%",)
        ).fetchone()
        seq = 1
        if last:
            try:
                seq = int(str(last['archive_number'])[len(_prefix):]) + 1
            except Exception:
                seq = 1
        archive_number = f"{_prefix}{seq:03d}"
        # تأكد أنه فريد
        attempt = 0
        while conn.execute("SELECT 1 FROM documents WHERE archive_number=?", (archive_number,)).fetchone():
            attempt += 1
            archive_number = f"{_prefix}{seq + attempt:03d}"

        if doc_mode == 'upload':
            sub_path = f"{ARCHIVE_REMOTE_PREFIX}/{archive_section_code}"
            remote_name = f"{archive_number}.{upload_ext or 'bin'}"
            remote_rel_path = f"{sub_path}/{remote_name}".replace('\\', '/').strip('/')
            if not storage.upload_bytes(upload_blob or b'', remote_name, sub_path=sub_path):
                conn.rollback()
                conn.close()
                return jsonify({"success": False, "error": "تعذر رفع الملف إلى الأرشيف مباشرة"}), 503
            archive_storage_path = remote_rel_path
            archived_at = datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S')

        # ─── إدراج الوثيقة ──────────────────────────────────────────
        conn.execute(
            """
            INSERT INTO documents
              (title, archive_number, archive_section, archive_section_code,
                    file_path, archive_storage_path, content_json, template_name, notes, status, created_by, file_sha256, archived_at)
                VALUES (?,?,?,?,?,?,?,?,?, 'draft', ?, ?, ?)
            """,
            (title, archive_number, archive_section, archive_section_code,
             file_path, archive_storage_path, content_json, text_template_name, note or None,
                 session['user_id'], document_fingerprint, archived_at)
        )
        doc_id = conn.execute("SELECT last_insert_rowid()").fetchone()[0]
        conn.commit()

        # ─── طلبات التوقيع للمذكورين ────────────────────────────────
        warning = None
        doc_row = conn.execute("SELECT * FROM documents WHERE id=? LIMIT 1", (doc_id,)).fetchone()
        if doc_mode == 'upload':
            saved_to_storage = True
        else:
            saved_to_storage = _upload_doc_to_storage(dict(doc_row), conn) if doc_row else False
            if not saved_to_storage:
                warning = "تم إنشاء الوثيقة لكن تعذر رفعها لمجلد الأرشيف"

        creator_name = session.get('user_name', '')
        for uid in mentions:
            user_row = conn.execute(
                "SELECT id, name, email FROM users WHERE id=? AND is_active=1 LIMIT 1",
                (uid,)
            ).fetchone()
            if not user_row:
                continue
            serial = generate_serial("SIG", "signature_requests", "serial_number")
            conn.execute(
                """
                INSERT INTO signature_requests
                  (document_id, requested_from, requested_by, serial_number, sign_type, status)
                VALUES (?,?,?,?,'signature','pending')
                """,
                (doc_id, uid, session['user_id'], serial)
            )
            conn.commit()
            try:
                send_sign_request(user_row['email'], user_row['name'],
                                  title, archive_number, creator_name, note or "")
            except Exception:
                if warning:
                    warning += " • فشل إرسال بعض الإشعارات بالبريد"
                else:
                    warning = "تم إنشاء الوثيقة لكن فشل إرسال بعض الإشعارات بالبريد"

        log_action(session['user_id'], "CREATE_DOCUMENT",
                   details=f"{archive_number} — {title}", ip=request.remote_addr)
        conn.close()

        result = {"success": True, "doc_id": doc_id, "archive_number": archive_number}
        if warning:
            result["warning"] = warning
        return jsonify(result)

    # ── GET: قائمة الوثائق ────────────────────────────────────────────
    limit = max(1, min(_safe_int(request.args.get('limit', 50), 50), 200))
    uid = session['user_id']
    role = (session.get('user_role') or '').lower()

    conn = get_db()
    if role in ('admin', 'manager', 'sys_admin'):
        rows = conn.execute(
            """
            SELECT d.*, u.name AS creator_name
            FROM documents d
            LEFT JOIN users u ON u.id = d.created_by
            ORDER BY d.id DESC
            LIMIT ?
            """,
            (limit,)
        ).fetchall()
    else:
        rows = conn.execute(
            """
            SELECT DISTINCT d.*, u.name AS creator_name
            FROM documents d
            LEFT JOIN users u ON u.id = d.created_by
            LEFT JOIN signature_requests sr ON sr.document_id = d.id
            WHERE d.created_by=? OR (sr.requested_from=? AND sr.status='pending')
            ORDER BY d.id DESC
            LIMIT ?
            """,
            (uid, uid, limit)
        ).fetchall()

    docs = [_decorate_doc(conn, dict(r)) for r in rows]
    conn.close()
    return jsonify(docs)


@app.route("/api/documents/<int:doc_id>")
@login_required
def api_document_details(doc_id):
    conn = get_db()
    row = conn.execute(
        """
        SELECT d.*, u.name AS creator_name
        FROM documents d
        LEFT JOIN users u ON u.id = d.created_by
        WHERE d.id=?
        LIMIT 1
        """,
        (doc_id,)
    ).fetchone()
    if not row:
        conn.close()
        return jsonify({"error": "الوثيقة غير موجودة"}), 404

    doc = _decorate_doc(conn, dict(row))
    conn.close()
    return jsonify(doc)


@app.route("/api/documents/<int:doc_id>/signatures")
@login_required
def api_document_signatures(doc_id):
    conn = get_db()
    doc = conn.execute(
        "SELECT id, status, approved_by, approved_at FROM documents WHERE id=? LIMIT 1",
        (doc_id,)
    ).fetchone()
    if not doc:
        conn.close()
        return jsonify([])

    rows = conn.execute(
        """
        SELECT sr.id,
               sr.sign_type,
               sr.status,
               sr.serial_number,
               sr.signed_at,
               u.name AS signer_name
        FROM signature_requests sr
        LEFT JOIN users u ON u.id = sr.requested_from
        WHERE sr.document_id=?
        ORDER BY sr.id ASC
        """,
        (doc_id,)
    ).fetchall()

    payload = [dict(r) for r in rows]
    if doc['approved_by'] and doc['approved_at']:
        approver = conn.execute("SELECT name FROM users WHERE id=? LIMIT 1", (doc['approved_by'],)).fetchone()
        payload.append({
            "entry_type": "approval",
            "status": doc['status'],
            "event_at": doc['approved_at'],
            "action_label": "اعتماد" if doc['status'] == 'approved' else "رفض",
            "signer_name": approver['name'] if approver else "—",
        })

    conn.close()
    return jsonify(payload)


@app.route("/api/documents/<int:doc_id>/download")
@login_required
def api_document_download(doc_id):
    conn = get_db()
    row = conn.execute(
        "SELECT id, title, archive_number, template_name, file_path, archive_storage_path, content_json, created_at, archive_section, archive_section_code FROM documents WHERE id=? LIMIT 1",
        (doc_id,)
    ).fetchone()
    conn.close()

    if not row:
        return jsonify({"error": "الوثيقة غير موجودة"}), 404

    resolved_path = _resolve_document_file_path(row['file_path'])
    if not resolved_path:
        resolved_path = _resolve_document_file_path(row['archive_storage_path'])
    if not resolved_path:
        if not row['file_path'] and row['content_json']:
            inline = str(request.args.get('inline', '')).lower() in ('1', 'true', 'yes')
            generated_path = _build_text_document_archive_file(dict(row))
            fallback_name = (row['archive_number'] or row['title'] or f"document-{doc_id}").replace('/', '-')

            if generated_path and os.path.exists(generated_path):
                generated_ext = os.path.splitext(generated_path)[1].lower()
                if generated_ext == '.docx':
                    mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                    download_name = f"{fallback_name}.docx"
                elif generated_ext == '.pdf':
                    mimetype = 'application/pdf'
                    download_name = f"{fallback_name}.pdf"
                else:
                    mimetype = 'text/html; charset=utf-8'
                    download_name = f"{fallback_name}.html"

                @after_this_request
                def _cleanup_generated_file(response):
                    try:
                        os.remove(generated_path)
                    except OSError:
                        pass
                    return response

                return send_file(
                    generated_path,
                    mimetype=mimetype,
                    as_attachment=not inline,
                    download_name=download_name
                )

            html_markup = _build_text_document_archive_markup(dict(row))
            payload = io.BytesIO(html_markup.encode('utf-8'))
            return send_file(
                payload,
                mimetype='text/html; charset=utf-8',
                as_attachment=not inline,
                download_name=f"{fallback_name}.html"
            )
        return jsonify({"error": "الملف غير موجود"}), 404

    inline = str(request.args.get('inline', '')).lower() in ('1', 'true', 'yes')
    return send_file(
        resolved_path,
        as_attachment=not inline,
        download_name=os.path.basename(resolved_path)
    )


@app.route("/api/documents/<int:doc_id>/print")
@login_required
def api_document_print(doc_id):
    """عرض الوثيقة بصفحة HTML جاهزة للطباعة / حفظ كـ PDF من المتصفح"""
    conn = get_db()
    row = conn.execute(
        "SELECT id, title, archive_number, template_name, file_path, archive_storage_path, content_json, created_at, archive_section, archive_section_code FROM documents WHERE id=? LIMIT 1",
        (doc_id,)
    ).fetchone()
    conn.close()

    if not row:
        return jsonify({"error": "الوثيقة غير موجودة"}), 404

    doc = dict(row)
    if not doc.get('content_json'):
        return redirect(url_for('api_document_download', doc_id=doc_id, inline='1'))

    markup = _build_text_document_archive_markup(doc)
    bare = str(request.args.get('bare', '')).lower() in ('1', 'true', 'yes')
    if not bare:
        # حقن script طباعة تلقائية + تعليمات للمستخدم
        print_script = """
<script>
window.onload = function() {
    var bar = document.createElement('div');
    bar.id = 'print-bar';
    bar.style.cssText = 'position:fixed;top:0;left:0;right:0;background:#1e40af;color:#fff;padding:10px 20px;font-family:Arial;font-size:14px;z-index:9999;display:flex;gap:12px;align-items:center;direction:rtl';
    bar.innerHTML = '<span>📄 لحفظ كـ PDF: اضغط زر الطباعة ثم اختر <b>"حفظ كـ PDF"</b></span>'
      + '<button onclick="window.print()" style="background:#fff;color:#1e40af;border:none;padding:6px 14px;border-radius:6px;font-weight:700;cursor:pointer">🖨️ طباعة / PDF</button>'
      + '<button onclick="document.getElementById(\\'print-bar\\').style.display=\\'none\\'" style="background:rgba(255,255,255,.2);color:#fff;border:none;padding:6px 10px;border-radius:6px;cursor:pointer">✕</button>';
    document.body.insertBefore(bar, document.body.firstChild);
};
</script>
<style>
@media print { #print-bar { display: none !important; } }
</style>
"""
        markup = markup.replace('</head>', print_script + '</head>', 1)
    return markup, 200, {'Content-Type': 'text/html; charset=utf-8'}


@app.route("/api/documents/<int:doc_id>/download-signed")
@login_required
def api_document_download_signed(doc_id):
    conn = get_db()
    row = conn.execute(
        "SELECT id, title, archive_number, template_name, file_path, archive_storage_path, content_json, created_at, archive_section, archive_section_code FROM documents WHERE id=? LIMIT 1",
        (doc_id,)
    ).fetchone()
    conn.close()

    if not row:
        return jsonify({"error": "الوثيقة غير موجودة"}), 404

    doc = dict(row)
    inline = str(request.args.get('inline', '')).lower() in ('1', 'true', 'yes')
    fallback_name = (doc.get('archive_number') or doc.get('title') or f"document-{doc_id}").replace('/', '-')

    signed_path = _build_signed_document_file(doc)
    if not signed_path or not os.path.exists(signed_path):
        return api_document_download(doc_id)

    source_path = _resolve_document_file_path(doc.get('file_path'))
    source_ext = _get_document_extension(doc)
    generated_ext = os.path.splitext(signed_path)[1].lower().lstrip('.')
    if not generated_ext:
        generated_ext = source_ext or 'docx'

    mimetype_map = {
        'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'html': 'text/html; charset=utf-8',
        'png': 'image/png',
        'jpg': 'image/jpeg',
        'jpeg': 'image/jpeg',
        'webp': 'image/webp',
        'gif': 'image/gif',
    }

    is_temporary = os.path.abspath(signed_path) != os.path.abspath(source_path or '')

    if is_temporary:
        @after_this_request
        def _cleanup_signed_file(response):
            try:
                os.remove(signed_path)
            except OSError:
                pass
            return response

    return send_file(
        signed_path,
        mimetype=mimetype_map.get(generated_ext),
        as_attachment=not inline,
        download_name=f"{fallback_name}.{generated_ext}"
    )


@app.route("/api/documents/<int:doc_id>/preview")
@login_required
def api_document_preview(doc_id):
    page = max(0, _safe_int(request.args.get('page', 0), 0))

    conn = get_db()
    row = conn.execute(
        "SELECT id, title, file_path, archive_storage_path, archive_section_code, content_json, template_name FROM documents WHERE id=? LIMIT 1",
        (doc_id,)
    ).fetchone()
    conn.close()
    if not row:
        return _svg_response(_svg_placeholder(["الوثيقة غير موجودة"], subtitle="المعاينة"), status=404)

    doc = dict(row)
    
    # ── Try to use signed/built version first if signed ─────────
    # This ensures QR and signatures are visible in preview
    signed_doc_path = None
    try:
        signed_version_path = _build_signed_document_file(doc, conn=conn)
        if signed_version_path and os.path.exists(signed_version_path):
            signed_doc_path = signed_version_path
            resolved_path = signed_version_path
            print(f"[INFO] Preview using signed document version for doc {doc_id}: {signed_version_path}", file=__import__('sys').stderr)
        else:
            resolved_path = _resolve_document_file_path(doc.get('file_path'))
            if not resolved_path:
                resolved_path = _resolve_document_file_path(doc.get('archive_storage_path'))
    except Exception as e:
        print(f"[WARNING] Failed to build signed version for preview doc {doc_id}: {e}", file=__import__('sys').stderr)
        resolved_path = _resolve_document_file_path(doc.get('file_path'))
        if not resolved_path:
            resolved_path = _resolve_document_file_path(doc.get('archive_storage_path'))

    # ── For text documents without a file: build temporary PDF ────
    _temp_generated_path = None
    if not resolved_path and doc.get('content_json') and not doc.get('file_path'):
        try:
            temp_file_path = _build_text_document_archive_file(doc)
            if temp_file_path and os.path.exists(temp_file_path):
                if temp_file_path.lower().endswith('.pdf'):
                    resolved_path = temp_file_path
                    _temp_generated_path = temp_file_path
        except Exception:
            pass

    ext = _get_document_extension(doc)
    if not ext and resolved_path and resolved_path.lower().endswith('.pdf'):
        ext = 'pdf'

    if resolved_path and ext in ('png', 'jpg', 'jpeg', 'webp', 'gif'):
        return send_file(resolved_path)

    if ext == 'pdf':
        rendered_stream, total_pages = _render_pdf_page_png(resolved_path, page_index=page, scale=1.5)
        # تنظيف ملفات PDF المؤقتة بعد تحميلها في الذاكرة
        if _temp_generated_path:
            try:
                os.remove(_temp_generated_path)
            except OSError:
                pass
        if signed_doc_path and signed_doc_path != resolved_path:
            try:
                os.remove(signed_doc_path)
            except OSError:
                pass
            
        if rendered_stream:
            final_stream = rendered_stream
            # Note: Since we now use the signed/built document version above,
            # overlay is only needed if preview fell back to unsigned original
            base_image = None
            try:
                # Only overlay if NOT using a pre-built signed document
                if not signed_doc_path:
                    signed_assets = _get_latest_signed_assets_preview(doc_id)
                    if signed_assets:
                        rendered_stream.seek(0)
                        from PIL import Image
                        base_image = Image.open(rendered_stream).convert('RGBA')
                        overlaid = _overlay_signed_assets_on_pil_image(
                            base_image,
                            signed_assets,
                            section_code=doc.get('archive_section_code')
                        )
                        out_stream = io.BytesIO()
                        overlaid.save(out_stream, format='PNG', optimize=True)
                        out_stream.seek(0)
                        final_stream = out_stream
            except Exception as e:
                print(f"[WARNING] Failed to overlay signed assets in preview: {e}", file=__import__('sys').stderr)
                final_stream = rendered_stream
            finally:
                try:
                    if base_image is not None:
                        base_image.close()
                except Exception:
                    pass

            response = send_file(
                final_stream,
                mimetype='image/png',
                download_name=f"preview-{doc_id}-{page + 1}.png",
                max_age=0
            )
            response.headers['Cache-Control'] = 'no-store, no-cache, must-revalidate, max-age=0'
            response.headers['Pragma'] = 'no-cache'
            response.headers['Expires'] = '0'
            return response

        total_pages = _count_pdf_pages(resolved_path)
        svg = _svg_placeholder(
            [
                row['title'] or 'وثيقة PDF',
                f"صفحة {page + 1} من {max(1, total_pages)}",
                "تعذر إنشاء معاينة الصفحة حالياً"
            ],
            subtitle="PDF Preview"
        )
        return _svg_response(svg)

    if resolved_path:
        svg = _svg_placeholder([
            row['title'] or 'وثيقة',
            "هذا النوع لا يدعم معاينة الصفحة",
            "يمكنك المتابعة بالتوقيع ثم تنزيل الملف"
        ], subtitle="Document Preview")
        return _svg_response(svg)

    if _temp_generated_path:
        try:
            os.remove(_temp_generated_path)
        except OSError:
            pass

    return _svg_response(_svg_placeholder([
        row['title'] or 'وثيقة',
        "الملف غير متوفر حالياً",
        "تحقق من التخزين أو ارفع نسخة جديدة"
    ], subtitle="Document Preview"), status=404)


@app.route("/api/documents/<int:doc_id>/preview-office")
@login_required
def api_document_preview_office(doc_id):
    conn = get_db()
    row = conn.execute("SELECT id, title, file_path FROM documents WHERE id=? LIMIT 1", (doc_id,)).fetchone()
    conn.close()
    if not row:
        return "<div style='padding:24px;font-family:Arial'>الوثيقة غير موجودة</div>", 404
    if not row['file_path']:
        return "<div style='padding:24px;font-family:Arial'>لا يوجد ملف للمعاينة</div>", 404
    return f"""
    <html><body style='margin:0;background:#0f172a;color:#e2e8f0;font-family:Arial,sans-serif'>
      <div style='padding:24px;text-align:center'>
        <div style='font-size:18px;margin-bottom:8px'>المعاينة المباشرة غير متاحة لهذا النوع داخل المتصفح</div>
        <div style='opacity:.8;margin-bottom:14px'>{row['title']}</div>
        <a href='/api/documents/{doc_id}/download' style='display:inline-block;background:#1d4ed8;color:#fff;text-decoration:none;padding:10px 16px;border-radius:8px'>تحميل الملف</a>
      </div>
    </body></html>
    """


def _upload_doc_to_storage(doc, conn):
    """يرفع الوثيقة إلى التخزين حسب القسم ويحدث archive_storage_path. يدعم الملفات و وثائق النص."""
    if str(doc.get('archive_storage_path') or '').strip():
        return True

    archive_num = str(doc.get('archive_number') or doc.get('id') or 'doc').replace('/', '-').strip() or 'doc'
    section_code = str(doc.get('archive_section_code') or 'GN').strip().upper() or 'GN'
    sub_path = f"{ARCHIVE_REMOTE_PREFIX}/{section_code}"

    source_file_path = str(doc.get('file_path') or '').strip()
    local_path = None
    generated_temp = None

    if source_file_path:
        local_path = _resolve_document_file_path(source_file_path)
    else:
        try:
            generated_temp = _build_text_document_archive_file(doc, conn=conn)
            local_path = generated_temp
        except Exception:
            local_path = None

    if not local_path or not os.path.exists(local_path):
        if generated_temp and os.path.exists(generated_temp):
            try:
                os.remove(generated_temp)
            except OSError:
                pass
        return False

    if source_file_path:
        ext = source_file_path.split('?')[0].split('.')[-1].lower()
    elif generated_temp:
        ext = os.path.splitext(generated_temp)[1].lower().lstrip('.')
    else:
        ext = 'html'

    if not ext:
        ext = 'html'

    remote_name = f"{archive_num}.{ext}"
    remote_rel_path = f"{sub_path}/{remote_name}".replace('\\', '/').strip('/')

    try:
        ok = storage.upload(local_path, remote_name, sub_path=sub_path)
        if ok:
            conn.execute(
                "UPDATE documents SET archive_storage_path=?, file_path=NULL, archived_at=COALESCE(archived_at, datetime('now')) WHERE id=?",
                (remote_rel_path, doc['id'])
            )
            conn.commit()
            # احذف النسخة المحلية بعد الرفع الناجح
            if source_file_path and local_path and os.path.exists(local_path):
                try:
                    os.remove(local_path)
                except OSError:
                    pass
        return ok
    except Exception:
        return False
    finally:
        if generated_temp and os.path.exists(generated_temp):
            try:
                os.remove(generated_temp)
            except OSError:
                pass


@app.route("/api/documents/<int:doc_id>", methods=["PUT"])
@login_required
def api_document_edit(doc_id):
    conn = get_db()
    row = conn.execute("SELECT * FROM documents WHERE id=? LIMIT 1", (doc_id,)).fetchone()
    if not row:
        conn.close()
        return jsonify({"success": False, "error": "الوثيقة غير موجودة"}), 404

    uid = session['user_id']
    role = (session.get('user_role') or '').lower()
    if row['created_by'] != uid and role not in ('admin', 'sys_admin'):
        conn.close()
        return jsonify({"success": False, "error": "ليس لديك صلاحية التعديل"}), 403
    if row['status'] == 'approved':
        conn.close()
        return jsonify({"success": False, "error": "لا يمكن تعديل وثيقة معتمدة"}), 400

    data  = request.json or {}
    title = (data.get('title') or '').strip()
    note  = (data.get('note') or '').strip() or None
    content_html = (data.get('content') or '').strip()

    if not title:
        conn.close()
        return jsonify({"success": False, "error": "العنوان مطلوب"}), 400

    updates = {"title": title, "notes": note}
    regenerated = False
    if content_html and not row['file_path']:
        updates['content_json'] = json.dumps({"html": content_html}, ensure_ascii=False)
        regenerated = True

    conn.execute(
        "UPDATE documents SET title=?, notes=?" + (", content_json=?" if regenerated else "") + " WHERE id=?",
        ([title, note, updates['content_json'], doc_id] if regenerated else [title, note, doc_id])
    )
    conn.commit(); conn.close()
    log_action(uid, "EDIT_DOCUMENT", details=str(doc_id), ip=request.remote_addr)
    return jsonify({"success": True, "regenerated": regenerated})


@app.route("/api/documents/<int:doc_id>", methods=["DELETE"])
@login_required
def api_document_archive(doc_id):
    conn = get_db()
    row = conn.execute("SELECT * FROM documents WHERE id=? LIMIT 1", (doc_id,)).fetchone()
    if not row:
        conn.close()
        return jsonify({"success": False, "error": "الوثيقة غير موجودة"}), 404

    uid  = session['user_id']
    role = (session.get('user_role') or '').lower()
    if row['created_by'] != uid and role not in ('admin', 'manager', 'sys_admin'):
        conn.close()
        return jsonify({"success": False, "error": "ليس لديك صلاحية"}), 403

    doc = dict(row)
    if doc.get('archive_storage_path'):
        conn.close()
        return jsonify({"success": True, "already_archived": True})

    saved = _upload_doc_to_storage(doc, conn)
    conn.execute(
        "UPDATE documents SET archived_at=COALESCE(archived_at, datetime('now')), archived_by=? WHERE id=?",
        (uid, doc_id)
    )
    conn.commit(); conn.close()
    log_action(uid, "ARCHIVE_DOCUMENT", details=str(doc_id), ip=request.remote_addr)
    warning = None if saved else "تم تسجيل الأرشفة لكن لم يتم رفع الملف إلى التخزين"
    return jsonify({"success": True, "saved_to_storage": saved, "warning": warning})


@app.route("/api/documents/<int:doc_id>/approve", methods=["POST"])
@login_required
def api_document_approve(doc_id):
    role = (session.get('user_role') or '').lower()
    if role not in ('admin', 'manager', 'sys_admin'):
        return jsonify({"success": False, "error": "ليس لديك صلاحية الاعتماد"}), 403

    conn = get_db()
    row = conn.execute("SELECT * FROM documents WHERE id=? LIMIT 1", (doc_id,)).fetchone()
    if not row:
        conn.close()
        return jsonify({"success": False, "error": "الوثيقة غير موجودة"}), 404
    if row['status'] == 'approved':
        conn.close()
        return jsonify({"success": False, "error": "الوثيقة معتمدة مسبقاً"}), 400

    conn.execute(
        "UPDATE documents SET status='approved', approved_by=?, approved_at=datetime('now') WHERE id=?",
        (session['user_id'], doc_id)
    )
    _persist_document_fingerprint(conn, doc_id)
    conn.commit()

    doc = dict(conn.execute("SELECT * FROM documents WHERE id=? LIMIT 1", (doc_id,)).fetchone())
    saved = _upload_doc_to_storage(doc, conn)

    # إشعار المنشئ
    warning = None
    try:
        creator = conn.execute("SELECT name, email FROM users WHERE id=? LIMIT 1", (row['created_by'],)).fetchone()
        if creator:
            send_approved(creator['email'], creator['name'],
                          row['title'], row['archive_number'] or '—')
    except Exception:
        warning = "تم الاعتماد لكن فشل إرسال الإشعار بالبريد"

    conn.close()
    log_action(session['user_id'], "APPROVE_DOCUMENT", details=str(doc_id), ip=request.remote_addr)
    result = {"success": True, "saved_to_storage": saved}
    if warning:
        result["warning"] = warning
    return jsonify(result)


@app.route("/api/documents/<int:doc_id>/reject", methods=["POST"])
@login_required
def api_document_reject(doc_id):
    role = (session.get('user_role') or '').lower()
    if role not in ('admin', 'manager', 'sys_admin'):
        return jsonify({"success": False, "error": "ليس لديك صلاحية الرفض"}), 403

    data   = request.json or {}
    reason = (data.get('reason') or '').strip()

    conn = get_db()
    row  = conn.execute("SELECT * FROM documents WHERE id=? LIMIT 1", (doc_id,)).fetchone()
    if not row:
        conn.close()
        return jsonify({"success": False, "error": "الوثيقة غير موجودة"}), 404

    conn.execute(
        "UPDATE documents SET status='draft', notes=COALESCE(NULLIF(?,''), notes) WHERE id=?",
        (f"[مرفوض]: {reason}" if reason else None, doc_id)
    )
    conn.commit()

    warning = None
    try:
        creator = conn.execute("SELECT name, email FROM users WHERE id=? LIMIT 1", (row['created_by'],)).fetchone()
        if creator:
            send_rejected(creator['email'], creator['name'],
                          row['title'], row['archive_number'] or '—', reason)
    except Exception:
        warning = "تم الرفض لكن فشل إرسال الإشعار بالبريد"

    conn.close()
    log_action(session['user_id'], "REJECT_DOCUMENT", details=str(doc_id), ip=request.remote_addr)
    result = {"success": True}
    if warning:
        result["warning"] = warning
    return jsonify(result)


@app.route("/api/documents/<int:doc_id>/convert-pdf", methods=["POST"])
@login_required
def api_document_convert_pdf(doc_id):
    conn = get_db()
    row  = conn.execute("SELECT * FROM documents WHERE id=? LIMIT 1", (doc_id,)).fetchone()
    conn.close()
    if not row:
        return jsonify({"success": False, "error": "الوثيقة غير موجودة"}), 404

    file_path = row['file_path'] or ''
    ext = file_path.split('?')[0].split('.')[-1].lower() if file_path else ''

    if ext == 'pdf':
        return jsonify({"success": True, "already_pdf": True})

    # محاولة التحويل باستخدام LibreOffice إذا كان متاحاً
    import subprocess
    local = _resolve_document_file_path(file_path)
    if not local or not os.path.exists(local):
        return jsonify({"success": False, "error": "الملف غير موجود"}), 404

    try:
        out_dir = os.path.dirname(local)
        result  = subprocess.run(
            ['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', out_dir, local],
            capture_output=True, timeout=60
        )
        pdf_name = os.path.splitext(os.path.basename(local))[0] + '.pdf'
        pdf_path = os.path.join(out_dir, pdf_name)
        if os.path.exists(pdf_path):
            new_rel = 'uploads/documents/' + pdf_name
            conn2 = get_db()
            conn2.execute("UPDATE documents SET file_path=? WHERE id=?", (new_rel, doc_id))
            conn2.commit(); conn2.close()
            log_action(session['user_id'], "CONVERT_PDF", details=str(doc_id), ip=request.remote_addr)
            return jsonify({"success": True, "already_pdf": False})
        return jsonify({"success": False, "error": "فشل التحويل — LibreOffice غير متاح أو الملف غير مدعوم",
                        "warning": "يمكنك تحميل الملف يدوياً وتحويله"}), 400
    except (FileNotFoundError, subprocess.TimeoutExpired):
        return jsonify({"success": False,
                        "error": "LibreOffice غير مثبت على الخادم",
                        "warning": "يمكنك تحميل الملف وتحويله خارجياً ثم رفعه من جديد"}), 400


@app.route("/api/docx-templates", methods=["GET"])
@login_required
def api_docx_templates():
    default_template = _get_default_docx_template_name()
    items = _list_docx_templates()
    for item in items:
        item['is_default'] = str(item.get('file_name') or '') == default_template
    return jsonify(items)


@app.route("/api/docx-templates/default", methods=["GET"])
@login_required
def api_docx_templates_default():
    return jsonify({
        "success": True,
        "default_template": _get_default_docx_template_name()
    })


@app.route("/api/admin/docx-templates", methods=["GET"])
@login_required
@role_required("admin", "sys_admin")
def api_admin_docx_templates():
    default_template = _get_default_docx_template_name()
    items = _list_docx_templates()
    for item in items:
        item['is_default'] = str(item.get('file_name') or '') == default_template
    return jsonify({
        "success": True,
        "default_template": default_template,
        "templates": items
    })


@app.route("/api/admin/docx-templates/default", methods=["POST"])
@login_required
@role_required("admin", "sys_admin")
def api_admin_set_default_docx_template():
    data = request.json or {}
    template_name = str(data.get('template_name') or '').strip()

    ok, result = _set_default_docx_template_name(template_name)
    if not ok:
        return jsonify({"success": False, "error": result}), 400

    chosen = str(result or '')
    log_action(
        session['user_id'],
        "SET_DEFAULT_DOCX_TEMPLATE",
        chosen or "clear-default-template",
        ip=request.remote_addr,
        user_agent=request.user_agent.string
    )

    return jsonify({
        "success": True,
        "default_template": chosen
    })


# ════════════════════════════════════════════
# STAMP TEMPLATES (multi-template CRUD)
# ════════════════════════════════════════════

@app.route("/api/admin/stamp-templates", methods=["GET"])
@login_required
@role_required("admin", "sys_admin")
def api_list_stamp_templates():
    conn = get_db()
    rows = conn.execute(
        "SELECT * FROM stamp_templates ORDER BY is_default DESC, id DESC"
    ).fetchall()
    conn.close()
    stamps_dir = os.path.join(UPLOAD_FOLDER, 'stamps')
    result = []
    for r in rows:
        path = os.path.join(stamps_dir, r['file_name'])
        stat = os.stat(path) if os.path.isfile(path) else None
        result.append({
            "id": r['id'],
            "name": r['name'],
            "file_name": r['file_name'],
            "text_x_ratio": r['text_x_ratio'],
            "text_y_ratio": r['text_y_ratio'],
            "is_default": bool(r['is_default']),
            "created_at": r['created_at'],
            "size_kb": int(stat.st_size / 1024) if stat else 0,
            "preview_url": f"/api/admin/stamp-templates/{r['id']}/preview",
        })
    return jsonify(result)


@app.route("/api/admin/stamp-templates/upload", methods=["POST"])
@login_required
@role_required("admin", "sys_admin")
def api_upload_stamp_template():
    f = request.files.get('template')
    name = (request.form.get('name') or '').strip()
    if not f:
        return jsonify({"success": False, "error": "لا يوجد ملف"}), 400
    raw_name = secure_filename(f.filename or '')
    ext = raw_name.rsplit('.', 1)[-1].lower() if '.' in raw_name else ''
    if ext not in ('png', 'jpg', 'jpeg', 'webp'):
        return jsonify({"success": False, "error": "يُقبل PNG أو JPG فقط"}), 400
    stamps_dir = os.path.join(UPLOAD_FOLDER, 'stamps')
    os.makedirs(stamps_dir, exist_ok=True)
    import uuid as _uuid
    file_name = f"stpl_{_uuid.uuid4().hex[:8]}.{ext}"
    f.save(os.path.join(stamps_dir, file_name))
    conn = get_db()
    is_first = conn.execute("SELECT COUNT(*) FROM stamp_templates").fetchone()[0] == 0
    cur = conn.execute(
        "INSERT INTO stamp_templates (name, file_name, text_x_ratio, text_y_ratio, is_default) VALUES (?,?,0.25,0.08,?)",
        (name or file_name, file_name, 1 if is_first else 0)
    )
    tpl_id = cur.lastrowid
    conn.commit()
    conn.close()
    log_action(session['user_id'], "UPLOAD_STAMP_TEMPLATE", file_name,
               ip=request.remote_addr, user_agent=request.user_agent.string)
    return jsonify({"success": True, "id": tpl_id, "file_name": file_name})


@app.route("/api/admin/stamp-templates/<int:tpl_id>/position", methods=["PUT"])
@login_required
@role_required("admin", "sys_admin")
def api_update_stamp_template_position(tpl_id):
    data = request.json or {}
    try:
        x = max(0.0, min(1.0, float(data.get('text_x_ratio', 0.25))))
        y = max(0.0, min(1.0, float(data.get('text_y_ratio', 0.08))))
    except (TypeError, ValueError):
        return jsonify({"success": False, "error": "قيم غير صالحة"}), 400
    name = str(data.get('name') or '').strip()
    conn = get_db()
    row = conn.execute("SELECT id FROM stamp_templates WHERE id=?", (tpl_id,)).fetchone()
    if not row:
        conn.close()
        return jsonify({"success": False, "error": "القالب غير موجود"}), 404
    if name:
        conn.execute("UPDATE stamp_templates SET text_x_ratio=?, text_y_ratio=?, name=? WHERE id=?",
                     (x, y, name, tpl_id))
    else:
        conn.execute("UPDATE stamp_templates SET text_x_ratio=?, text_y_ratio=? WHERE id=?",
                     (x, y, tpl_id))
    conn.commit()
    conn.close()
    return jsonify({"success": True})


@app.route("/api/admin/stamp-templates/<int:tpl_id>/default", methods=["POST"])
@login_required
@role_required("admin", "sys_admin")
def api_set_default_stamp_template(tpl_id):
    conn = get_db()
    if not conn.execute("SELECT id FROM stamp_templates WHERE id=?", (tpl_id,)).fetchone():
        conn.close()
        return jsonify({"success": False, "error": "القالب غير موجود"}), 404
    conn.execute("UPDATE stamp_templates SET is_default=0")
    conn.execute("UPDATE stamp_templates SET is_default=1 WHERE id=?", (tpl_id,))
    conn.commit()
    conn.close()
    return jsonify({"success": True})


@app.route("/api/admin/stamp-templates/<int:tpl_id>/preview")
@login_required
def api_preview_stamp_template(tpl_id):
    conn = get_db()
    row = conn.execute("SELECT file_name FROM stamp_templates WHERE id=?", (tpl_id,)).fetchone()
    conn.close()
    if not row:
        return jsonify({"error": "غير موجود"}), 404
    path = os.path.join(UPLOAD_FOLDER, 'stamps', row['file_name'])
    if not os.path.isfile(path):
        return jsonify({"error": "الملف مفقود"}), 404
    return send_file(path)


@app.route("/api/admin/stamp-templates/<int:tpl_id>", methods=["DELETE"])
@login_required
@role_required("admin", "sys_admin")
def api_delete_stamp_template(tpl_id):
    conn = get_db()
    row = conn.execute("SELECT file_name FROM stamp_templates WHERE id=?", (tpl_id,)).fetchone()
    if not row:
        conn.close()
        return jsonify({"success": False, "error": "القالب غير موجود"}), 404
    path = os.path.join(UPLOAD_FOLDER, 'stamps', row['file_name'])
    try:
        if os.path.isfile(path):
            os.unlink(path)
    except Exception:
        pass
    conn.execute("DELETE FROM stamp_templates WHERE id=?", (tpl_id,))
    conn.commit()
    # إن كان هناك قالب آخر اجعله افتراضياً
    conn.execute(
        "UPDATE stamp_templates SET is_default=1 WHERE id=(SELECT id FROM stamp_templates ORDER BY id DESC LIMIT 1)"
    )
    conn.commit()
    conn.close()
    log_action(session['user_id'], "DELETE_STAMP_TEMPLATE", str(tpl_id), ip=request.remote_addr)
    return jsonify({"success": True})


@app.route("/api/admin/generate-manager-stamps", methods=["POST"])
@login_required
@role_required("admin", "sys_admin")
def api_generate_manager_stamps():
    """توليد ختم PNG لكل مدير ليس لديه ختم مرفوع، باستخدام قالب الختم المحدد."""
    data = request.json or {}
    conn = get_db()
    force = bool(data.get('force', False))
    tpl_id = _safe_int(data.get('template_id'), 0)
    if tpl_id:
        tpl_row = conn.execute("SELECT * FROM stamp_templates WHERE id=?", (tpl_id,)).fetchone()
    else:
        tpl_row = _get_default_stamp_tpl(conn)

    if not tpl_row:
        conn.close()
        return jsonify({"success": False,
                        "error": "يجب رفع قالب ختم أولاً من تبويب «قوالب الختم»"}), 400

    managers = conn.execute(
        """
        SELECT u.id, u.name, u.employee_id, s.section_code
        FROM users u
        LEFT JOIN archive_sections s ON s.id = u.archive_section_id
        WHERE u.is_active=1
          AND COALESCE(u.role,'') IN ('manager','admin','sys_admin')
        ORDER BY u.id ASC
        """
    ).fetchall()

    generated = 0
    already_exists = 0
    failed = 0
    stamps_dir = os.path.join(UPLOAD_FOLDER, 'stamps')
    os.makedirs(stamps_dir, exist_ok=True)

    import shutil as _shutil
    for m in managers:
        uid = m['id']
        existing = conn.execute(
            "SELECT id, stamp_path FROM stamp_assets WHERE user_id=? AND is_active=1 LIMIT 1", (uid,)
        ).fetchone()
        if existing:
            if not force:
                already_exists += 1
                continue
            # force=True: delete old stamp file and DB record
            old_path = os.path.join(UPLOAD_FOLDER, '..', 'static',
                                     str(existing['stamp_path']).lstrip('/').replace('/', os.sep))
            try:
                if os.path.isfile(old_path):
                    os.remove(old_path)
            except Exception:
                pass
            conn.execute("DELETE FROM stamp_assets WHERE id=?", (existing['id'],))

        sec = str(m['section_code'] or '').strip().upper()
        emp = str(m['employee_id'] or '').strip()
        tmp_file = _render_stamp_from_tpl(
            tpl_row['file_name'], tpl_row['text_x_ratio'], tpl_row['text_y_ratio'],
            sec, emp, '', as_file=True
        )
        if not tmp_file:
            failed += 1
            continue

        fname = f"auto_mgr_{uid}_{int(datetime.utcnow().timestamp())}.png"
        dest = os.path.join(stamps_dir, fname)
        try:
            _shutil.move(tmp_file, dest)
        except Exception:
            failed += 1
            continue

        conn.execute(
            "INSERT INTO stamp_assets (user_id, stamp_name, stamp_path, visibility_scope, is_active) VALUES (?,?,?,?,1)",
            (uid, 'ختم المدير (تلقائي)', f"uploads/stamps/{fname}", 'self')
        )
        generated += 1

    conn.commit()
    conn.close()
    log_action(session['user_id'], "GENERATE_MANAGER_STAMPS",
               f"generated={generated}, exists={already_exists}, failed={failed}",
               ip=request.remote_addr)
    msg = f"تم توليد {generated} ختم"
    if already_exists:
        msg += f" • موجود مسبقاً: {already_exists}"
    if failed:
        msg += f" • فشل: {failed}"
    return jsonify({"success": True, "generated": generated,
                    "already_exists": already_exists, "failed": failed, "message": msg})



@app.route("/api/admin/docx-templates/upload", methods=["POST"])
@login_required
@role_required("admin", "sys_admin")
def api_admin_upload_docx_template():
    template_file = request.files.get('template')
    if not template_file:
        return jsonify({"success": False, "error": "لا يوجد ملف قالب"}), 400

    raw_name = secure_filename(template_file.filename or '')
    # Support both DOCX and PDF templates
    filename = _sanitize_any_template_name(raw_name)
    if not filename:
        return jsonify({"success": False, "error": "صيغة القالب يجب أن تكون DOCX أو PDF"}), 400

    if request.content_length and request.content_length > 20 * 1024 * 1024:
        return jsonify({"success": False, "error": "حجم الملف كبير جداً"}), 400

    is_pdf = filename.lower().endswith('.pdf')
    if is_pdf:
        target_path = _resolve_pdf_template_file(filename)
    else:
        target_path = _resolve_docx_template_file(filename)
    if not target_path:
        return jsonify({"success": False, "error": "مسار القالب غير مسموح"}), 400

    template_file.save(target_path)

    # For new PDF templates, write default config JSON
    if is_pdf and not os.path.exists(target_path + '.json'):
        _save_pdf_template_config(filename, _PDF_TEMPLATE_DEFAULT_CONFIG)

    stat_info = os.stat(target_path)
    template_meta = {
        "file_name": filename,
        "file_type": "pdf" if is_pdf else "docx",
        "size_kb": int((stat_info.st_size or 0) / 1024),
        "updated_at": datetime.fromtimestamp(stat_info.st_mtime).strftime('%Y-%m-%d %H:%M'),
    }

    log_action(
        session['user_id'],
        "UPLOAD_TEMPLATE",
        filename,
        ip=request.remote_addr,
        user_agent=request.user_agent.string
    )

    return jsonify({
        "success": True,
        "template": template_meta,
        "default_template": _get_default_docx_template_name()
    })


@app.route("/api/admin/docx-templates/<path:file_name>", methods=["PUT"])
@login_required
@role_required("admin", "sys_admin")
def api_admin_rename_docx_template(file_name):
    source_name = _sanitize_any_template_name(file_name)
    if not source_name:
        return jsonify({"success": False, "error": "اسم القالب غير صالح"}), 400

    is_pdf = source_name.lower().endswith('.pdf')
    if is_pdf:
        source_path = _resolve_pdf_template_file(source_name)
    else:
        source_path = _resolve_docx_template_file(source_name)
    if not source_path or not os.path.isfile(source_path):
        return jsonify({"success": False, "error": "القالب غير موجود"}), 404

    data = request.json or {}

    # Save PDF config positions if provided
    if is_pdf and 'config' in data:
        _save_pdf_template_config(source_name, data['config'])
        stat = os.stat(source_path)
        return jsonify({
            "success": True,
            "template": {
                "file_name": source_name,
                "file_type": "pdf",
                "size_kb": int(stat.st_size / 1024),
                "updated_at": datetime.fromtimestamp(stat.st_mtime).strftime('%Y-%m-%d %H:%M'),
                "config": _get_pdf_template_config(source_name),
            },
            "default_template": _get_default_docx_template_name()
        })

    target_name = _sanitize_any_template_name(data.get('new_name') or '')
    if not target_name:
        return jsonify({"success": False, "error": "الاسم الجديد غير صالح"}), 400
    if target_name.lower().endswith('.pdf') != is_pdf:
        return jsonify({"success": False, "error": "لا يمكن تغيير نوع القالب"}), 400

    if target_name == source_name:
        current_default = _get_default_docx_template_name()
        stat = os.stat(source_path)
        template_meta = {
            "file_name": source_name,
            "file_type": "pdf" if is_pdf else "docx",
            "size_kb": int(stat.st_size / 1024),
            "updated_at": datetime.fromtimestamp(stat.st_mtime).strftime('%Y-%m-%d %H:%M'),
            "is_default": source_name == current_default,
        }
        return jsonify({"success": True, "template": template_meta, "default_template": current_default})

    if is_pdf:
        target_path = _resolve_pdf_template_file(target_name)
    else:
        target_path = _resolve_docx_template_file(target_name)
    if not target_path:
        return jsonify({"success": False, "error": "مسار الاسم الجديد غير مسموح"}), 400
    if os.path.exists(target_path):
        return jsonify({"success": False, "error": "يوجد قالب بنفس الاسم"}), 409

    os.replace(source_path, target_path)
    # Move config JSON if exists
    if is_pdf and os.path.isfile(source_path + '.json'):
        try:
            os.replace(source_path + '.json', target_path + '.json')
        except Exception:
            pass

    current_default = _get_default_docx_template_name()
    if current_default == source_name:
        _set_default_docx_template_name(target_name)
        current_default = target_name

    log_action(
        session['user_id'],
        "RENAME_TEMPLATE",
        f"{source_name} -> {target_name}",
        ip=request.remote_addr,
        user_agent=request.user_agent.string
    )

    stat = os.stat(target_path)
    template_meta = {
        "file_name": target_name,
        "file_type": "pdf" if is_pdf else "docx",
        "size_kb": int(stat.st_size / 1024),
        "updated_at": datetime.fromtimestamp(stat.st_mtime).strftime('%Y-%m-%d %H:%M'),
        "is_default": target_name == current_default,
    }

    return jsonify({"success": True, "template": template_meta, "default_template": current_default})


@app.route("/api/admin/docx-templates/<path:file_name>", methods=["DELETE"])
@login_required
@role_required("admin", "sys_admin")
def api_admin_delete_docx_template(file_name):
    safe_name = _sanitize_any_template_name(file_name)
    if not safe_name:
        return jsonify({"success": False, "error": "اسم القالب غير صالح"}), 400

    is_pdf = safe_name.lower().endswith('.pdf')
    if is_pdf:
        target_path = _resolve_pdf_template_file(safe_name)
    else:
        target_path = _resolve_docx_template_file(safe_name)
    if not target_path or not os.path.isfile(target_path):
        return jsonify({"success": False, "error": "القالب غير موجود"}), 404

    os.remove(target_path)
    # Remove config JSON if exists
    if is_pdf and os.path.isfile(target_path + '.json'):
        try:
            os.remove(target_path + '.json')
        except Exception:
            pass

    current_default = _get_default_docx_template_name()
    if current_default == safe_name:
        _set_default_docx_template_name('')
        current_default = ''

    log_action(
        session['user_id'],
        "DELETE_TEMPLATE",
        safe_name,
        ip=request.remote_addr,
        user_agent=request.user_agent.string
    )

    return jsonify({
        "success": True,
        "deleted": safe_name,
        "default_template": current_default
    })


@app.route("/api/docx-templates/<path:file_name>/preview", methods=["GET"])
@login_required
def api_docx_template_preview(file_name):
    requested_name = _sanitize_any_template_name(file_name)
    if not requested_name:
        return jsonify({"success": False, "error": "اسم القالب غير صالح"}), 400

    is_pdf = requested_name.lower().endswith('.pdf')

    if is_pdf:
        target_path = _resolve_pdf_template_file(requested_name)
        if not target_path or not os.path.isfile(target_path):
            return jsonify({"success": False, "error": "القالب غير موجود"}), 404
        # Return a data-URI of the first page rendered as PNG
        try:
            import pypdfium2 as pdfium
            _pdf = pdfium.PdfDocument(target_path)
            _page = _pdf[0]
            _bmp = _page.render(scale=1.5)
            _img = _bmp.to_pil()
            _buf = io.BytesIO()
            _img.save(_buf, format='PNG', optimize=True)
            _buf.seek(0)
            page_uri = 'data:image/png;base64,' + base64.b64encode(_buf.read()).decode('ascii')
            _pdf.close()
        except Exception:
            page_uri = None
        cfg = _get_pdf_template_config(requested_name)
        return jsonify({
            "success": True,
            "template": requested_name,
            "file_type": "pdf",
            "page_image": page_uri,
            "config": cfg,
        })

    target_path = _resolve_docx_template_file(requested_name)
    if not target_path:
        return jsonify({"success": False, "error": "مسار القالب غير مسموح"}), 400
    if not os.path.isfile(target_path):
        return jsonify({"success": False, "error": "القالب غير موجود"}), 404

    try:
        preview = _build_docx_template_preview(target_path)
    except Exception:
        return jsonify({"success": False, "error": "تعذر إنشاء معاينة القالب"}), 500

    return jsonify({
        "success": True,
        "template": requested_name,
        "file_type": "docx",
        "preview": preview,
    })


@app.route("/api/docx-templates", methods=["POST"])
@login_required
def api_upload_docx_template():
    template_file = request.files.get('template')
    if not template_file:
        return jsonify({"success": False, "error": "لا يوجد ملف قالب"}), 400

    raw_name = secure_filename(template_file.filename or '')
    filename = _sanitize_docx_template_name(raw_name)
    if not filename:
        return jsonify({"success": False, "error": "صيغة القالب يجب أن تكون DOCX"}), 400

    if request.content_length and request.content_length > 12 * 1024 * 1024:
        return jsonify({"success": False, "error": "حجم الملف كبير جداً"}), 400

    target_path = _resolve_docx_template_file(filename)
    if not target_path:
        return jsonify({"success": False, "error": "مسار القالب غير مسموح"}), 400
    template_file.save(target_path)

    template_meta = _docx_template_meta(filename)
    return jsonify({
        "success": True,
        "template": template_meta or {"file_name": filename}
    })


@app.route("/api/archive")
@login_required
def api_archive():
    # تحقق من الصلاحيات
    role = (session.get('user_role') or '').lower()
    user_id = session.get('user_id')
    
    # الموظفون لا يستطيعون الوصول للأرشيف
    if role not in ('admin', 'sys_admin', 'manager'):
        return jsonify({"success": False, "error": "ليس لديك صلاحية الوصول للأرشيف"}), 403
    
    scope = _get_archive_scope_for_current_user()
    conn = get_db()

    # المديرين يشوفوا فقط ملفات أقسامهم المُصرح بها
    if scope['is_scoped']:
        allowed_codes = scope.get('section_codes') or [scope['section_code']]
        placeholders = ','.join('?' * len(allowed_codes))
        rows = conn.execute(
            f"""
            SELECT d.*, u.name AS creator_name,
                   (SELECT GROUP_CONCAT(DISTINCT u2.name)
                    FROM signature_requests sr2
                    LEFT JOIN users u2 ON u2.id = sr2.requested_from
                    WHERE sr2.document_id = d.id AND sr2.status='signed') AS signer_names,
                   (SELECT sr3.serial_number
                    FROM signature_requests sr3
                    WHERE sr3.document_id=d.id AND COALESCE(sr3.serial_number,'')!=''
                    ORDER BY sr3.id DESC LIMIT 1) AS serial
            FROM documents d
            LEFT JOIN users u ON u.id = d.created_by
            WHERE UPPER(COALESCE(d.archive_section_code, 'GN')) IN ({placeholders})
            ORDER BY d.id DESC
            """,
            allowed_codes
        ).fetchall()
    else:
        # Admin/Sys_admin يشوفوا كل الوثائق
        rows = conn.execute(
            """
            SELECT d.*, u.name AS creator_name,
                   (SELECT GROUP_CONCAT(DISTINCT u2.name)
                    FROM signature_requests sr2
                    LEFT JOIN users u2 ON u2.id = sr2.requested_from
                    WHERE sr2.document_id = d.id AND sr2.status='signed') AS signer_names,
                   (SELECT sr3.serial_number
                    FROM signature_requests sr3
                    WHERE sr3.document_id=d.id AND COALESCE(sr3.serial_number,'')!=''
                    ORDER BY sr3.id DESC LIMIT 1) AS serial
            FROM documents d
            LEFT JOIN users u ON u.id = d.created_by
            ORDER BY d.id DESC
            """
        ).fetchall()

    docs = [_decorate_doc(conn, dict(r)) for r in rows]
    conn.close()

    stats = {
        "total": len(docs),
        "approved": sum(1 for d in docs if d.get('status') == 'approved'),
        "pending": sum(1 for d in docs if d.get('status') == 'pending'),
        "draft": sum(1 for d in docs if d.get('status') == 'draft'),
        "rejected": sum(1 for d in docs if d.get('status') == 'rejected'),
    }

    return jsonify({"success": True, "docs": docs, "stats": stats, "scope": scope})


@app.route("/api/archive/check")
@login_required
def api_archive_check():
    archive_number = (request.args.get('archive_number') or '').strip()
    if not archive_number:
        return jsonify({"success": False, "error": "رقم الأرشيف مطلوب"}), 400

    conn = get_db()
    doc = conn.execute(
        """
        SELECT d.id, d.title, d.archive_number, d.archive_section, d.archive_section_code, d.status
        FROM documents d
        WHERE UPPER(TRIM(COALESCE(d.archive_number, ''))) = UPPER(TRIM(?))
        LIMIT 1
        """,
        (archive_number,)
    ).fetchone()

    if not doc:
        conn.close()
        return jsonify({"success": True, "exists": False})

    signatures = conn.execute(
        """
        SELECT sr.serial_number, sr.signed_at, u.name AS signer_name
        FROM signature_requests sr
        LEFT JOIN users u ON u.id = sr.requested_from
        WHERE sr.document_id=?
          AND sr.status='signed'
          AND COALESCE(sr.sign_type, 'signature')='signature'
          AND COALESCE(u.role, '') IN ('manager', 'admin', 'sys_admin')
        ORDER BY sr.id DESC
        """,
        (doc['id'],)
    ).fetchall()

    conn.close()
    return jsonify({
        "success": True,
        "exists": True,
        "doc": dict(doc),
        "manager_signatures": [dict(r) for r in signatures]
    })


@app.route("/api/search-serial")
@login_required
def api_search_serial():
    q = (request.args.get('q') or '').strip().upper()
    if not q:
        return jsonify({"success": False, "error": "رقم التسلسل مطلوب"}), 400

    conn = get_db()
    rows = conn.execute(
        """
        SELECT sr.serial_number, sr.sign_type, sr.status, sr.signed_at,
               d.title, d.archive_number, d.archive_section,
               u.name AS signer_name
        FROM signature_requests sr
        LEFT JOIN documents d ON d.id = sr.document_id
        LEFT JOIN users u ON u.id = sr.requested_from
        WHERE UPPER(TRIM(COALESCE(sr.serial_number, ''))) LIKE ?
        ORDER BY sr.id DESC
        LIMIT 20
        """,
        (f"%{q}%",)
    ).fetchall()
    conn.close()

    return jsonify({"success": True, "results": [dict(r) for r in rows]})


@app.route("/api/archive/storage")
@login_required
def api_archive_storage():
    scope = _get_archive_scope_for_current_user()
    active_sections = get_all_archive_sections()

    if scope['is_scoped']:
        allowed_codes = set(scope.get('section_codes') or [scope['section_code']])
        sections = [
            {
                "id": s.get('id'),
                "section_name": s.get('section_name'),
                "section_code": str(s.get('section_code') or 'GN').upper(),
            }
            for s in active_sections
            if str(s.get('section_code') or '').upper() in allowed_codes
        ]
        # إذا لم يُطابق أي قسم، أضف القسم الأساسي
        if not sections:
            scoped_row = next(
                (s for s in active_sections if str(s.get('section_code', '')).upper() == scope['section_code']),
                None
            )
            sections = [{
                "id": scoped_row.get('id') if scoped_row else None,
                "section_name": scope['section_name'],
                "section_code": scope['section_code'],
            }]
    else:
        sections = [{
            "id": s.get('id'),
            "section_name": s.get('section_name'),
            "section_code": str(s.get('section_code') or 'GN').upper(),
        } for s in active_sections]

    requested_code = (request.args.get('section_code') or '').strip().upper()
    if not requested_code:
        return jsonify({
            "success": True,
            "sections": sections,
            "scope": scope,
        })

    allowed_codes = {str(s.get('section_code') or '').upper() for s in sections}
    if requested_code not in allowed_codes:
        return jsonify({"success": False, "error": "غير مصرح لهذا القسم"}), 403

    selected_section = next(
        (s for s in sections if str(s.get('section_code') or '').upper() == requested_code),
        {"section_name": scope['section_name'], "section_code": requested_code}
    )

    relative_path = _normalize_relative_path(request.args.get('path', ''))
    root_path = f"{ARCHIVE_REMOTE_PREFIX}/{requested_code}"
    sub_path = f"{root_path}/{relative_path}" if relative_path else root_path

    raw_listing = storage.list_dir(sub_path)
    listing = {
        "folders": [dict(item) for item in (raw_listing.get('folders') or [])],
        "files": [dict(item) for item in (raw_listing.get('files') or [])],
    }

    conn = get_db()
    doc_rows = conn.execute(
        """
        SELECT d.id, d.title, d.archive_number, d.status, d.file_path, d.archive_storage_path, u.name AS creator_name
        FROM documents d
        LEFT JOIN users u ON u.id = d.created_by
        WHERE UPPER(COALESCE(d.archive_section_code, 'GN')) = ?
        """,
        (requested_code,)
    ).fetchall()
    conn.close()

    file_map = {}
    for row in doc_rows:
        row_data = dict(row)
        candidates = []
        fp = str(row_data.get('file_path') or '').replace('\\', '/').strip()
        ap = str(row_data.get('archive_storage_path') or '').replace('\\', '/').strip()
        if fp:
            candidates.append(os.path.basename(fp).lower())
        if ap:
            candidates.append(os.path.basename(ap).lower())
        for key in candidates:
            if key and key not in file_map:
                file_map[key] = row_data

    enriched_files = []
    for file_item in listing['files']:
        current = dict(file_item)
        name = str(current.get('name') or '')
        meta = file_map.get(name.lower())
        if meta:
            current['title'] = meta.get('title') or name
            current['archive_number'] = meta.get('archive_number') or ''
            current['creator_name'] = meta.get('creator_name') or ''
            current['status'] = meta.get('status') or ''
            current['download_url'] = f"/api/documents/{meta.get('id')}/download-signed"
        else:
            current.setdefault('title', name)
            current.setdefault('archive_number', '')
            current.setdefault('creator_name', '')
            current.setdefault('status', '')
            current.setdefault('download_url', '')
        enriched_files.append(current)

    resolved_path = f"{root_path}/{relative_path}" if relative_path else root_path

    return jsonify({
        "success": True,
        "scope": scope,
        "sections": sections,
        "section": {
            "section_name": selected_section.get('section_name') or scope['section_name'],
            "section_code": selected_section.get('section_code') or requested_code,
        },
        "root_path": root_path,
        "relative_path": relative_path,
        "resolved_path": resolved_path,
        "can_go_up": bool(relative_path),
        "listing": {
            "folders": listing['folders'],
            "files": enriched_files,
        }
    })


@app.route("/api/signature-request/<req_id>/sign-pdf", methods=["POST"])
@login_required
def sign_document_pdf(req_id):
    data = request.json or {}
    sign_type = str(data.get("sign_type", "signature") or "signature").strip().lower()
    if sign_type not in ("signature", "stamp", "both"):
        sign_type = "signature"

    sign_password = str(data.get("sign_password", "") or "")
    if sign_password and not verify_sign_password(session['user_id'], sign_password):
        return jsonify({"success": False, "error": "الرقم السري غير صحيح"}), 401

    signature_owner_id = _safe_int(data.get("signature_owner_id"), session['user_id'])
    signature_asset_id = _safe_int(data.get("signature_asset_id"), 0)
    if signature_asset_id <= 0:
        signature_asset_id = None

    stamp_owner_id_raw = data.get("stamp_owner_id")
    stamp_owner_id = _safe_int(stamp_owner_id_raw, 0) if stamp_owner_id_raw is not None else None
    stamp_asset_raw = data.get("stamp_asset_id")
    stamp_asset_str = str(stamp_asset_raw or '').strip().lower()
    if stamp_asset_str.startswith('text-'):
        stamp_asset_id = -1
        if not stamp_owner_id:
            stamp_owner_id = _safe_int(stamp_asset_str.split('-', 1)[-1], 0) or None
    else:
        stamp_asset_id = _safe_int(stamp_asset_raw, 0)
        if stamp_asset_id <= 0:
            stamp_asset_id = None

    include_qr = bool(data.get("include_qr", True))
    doc_id = _safe_int(data.get("doc_id"), 0)

    # للملفات المرفوعة: لا QR بغض النظر عن اختيار المستخدم
    if doc_id:
        _qr_conn = get_db()
        _qr_check_doc = _qr_conn.execute("SELECT file_path FROM documents WHERE id=? LIMIT 1", (doc_id,)).fetchone()
        _qr_conn.close()
        if _qr_check_doc and _qr_check_doc['file_path']:
            include_qr = False

    positions = _normalize_sign_positions(data.get("positions"), include_qr=include_qr)
    has_sig_position = any(item.get('type') == 'sig' for item in positions)
    has_stamp_position = any(item.get('type') == 'stamp' for item in positions)
    if sign_type == 'signature' and has_stamp_position:
        sign_type = 'both'
    if sign_type == 'stamp' and has_sig_position:
        sign_type = 'both'
    positions_json = json.dumps(positions, ensure_ascii=False)

    conn = get_db()
    role = (session.get('user_role') or '').lower()
    uid = session['user_id']

    visible_stamp_assets = _get_stamp_assets(conn, viewer_id=uid, viewer_role=session.get('user_role'))
    allowed_stamp_asset_ids = {
        _safe_int(a.get('asset_id'), 0)
        for a in visible_stamp_assets
        if str(a.get('kind') or '') == 'image' and _safe_int(a.get('asset_id'), 0) > 0
    }
    allowed_stamp_owner_ids = {
        _safe_int(a.get('owner_id'), 0)
        for a in visible_stamp_assets
        if _safe_int(a.get('owner_id'), 0) > 0
    }

    if req_id == 'direct':
        if not doc_id:
            conn.close()
            return jsonify({"success": False, "error": "doc_id مطلوب"}), 400

        req = conn.execute(
            """
            SELECT *
            FROM signature_requests
            WHERE document_id=? AND requested_from=? AND status='pending'
            ORDER BY id DESC
            LIMIT 1
            """,
            (doc_id, uid)
        ).fetchone()

        if not req:
            try:
                conn.execute(
                    """
                    INSERT INTO signature_requests
                      (document_id, requested_from, requested_by,
                       signature_owner_id, stamp_owner_id,
                       signature_asset_id, stamp_asset_id,
                       sign_type, status, message)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, 'pending', ?)
                    """,
                    (
                        doc_id, uid, uid,
                        signature_owner_id, stamp_owner_id,
                        signature_asset_id, stamp_asset_id,
                        sign_type,
                        'توقيع مباشر من شاشة التوقيع'
                    )
                )
            except Exception:
                conn.execute(
                    """
                    INSERT INTO signature_requests
                      (document_id, requested_from, requested_by,
                       signature_owner_id, stamp_owner_id,
                       sign_type, status, message)
                    VALUES (?, ?, ?, ?, ?, ?, 'pending', ?)
                    """,
                    (
                        doc_id, uid, uid,
                        signature_owner_id, stamp_owner_id,
                        sign_type,
                        'توقيع مباشر من شاشة التوقيع'
                    )
                )
            conn.commit()
            req = conn.execute("SELECT * FROM signature_requests WHERE id=last_insert_rowid() LIMIT 1").fetchone()
    else:
        req_id_int = _safe_int(req_id, 0)
        req = conn.execute("SELECT * FROM signature_requests WHERE id=? LIMIT 1", (req_id_int,)).fetchone()
        if not req:
            conn.close()
            return jsonify({"success": False, "error": "طلب التوقيع غير موجود"}), 404
        if req['requested_from'] != uid and role not in ('admin', 'manager', 'sys_admin'):
            conn.close()
            return jsonify({"success": False, "error": "غير مصرح"}), 403
        doc_id = req['document_id']

    doc = conn.execute("SELECT * FROM documents WHERE id=? LIMIT 1", (doc_id,)).fetchone()
    if not doc:
        conn.close()
        return jsonify({"success": False, "error": "الوثيقة غير موجودة"}), 404

    if sign_type in ('stamp', 'both'):
        if stamp_asset_id and stamp_asset_id > 0 and stamp_asset_id not in allowed_stamp_asset_ids:
            conn.close()
            return jsonify({"success": False, "error": "الختم المختار غير مصرح لك"}), 403

        if stamp_asset_id == -1:
            if not stamp_owner_id or stamp_owner_id not in allowed_stamp_owner_ids:
                conn.close()
                return jsonify({"success": False, "error": "الختم النصي غير مصرح لك"}), 403

        if (not stamp_asset_id) and stamp_owner_id and stamp_owner_id not in allowed_stamp_owner_ids:
            conn.close()
            return jsonify({"success": False, "error": "المالك المختار للختم غير مصرح لك"}), 403

    if req['status'] == 'signed' and req['serial_number']:
        serial_existing = req['serial_number']
        include_qr_existing = bool(_safe_int(req['include_qr'], 1)) if 'include_qr' in req.keys() else include_qr
        conn.close()
        return jsonify({"success": True, "serial": serial_existing, "include_qr": include_qr_existing})

    serial = _build_signature_serial(conn, doc_id)
    hash_seed = sign_password if sign_password else f"unlock-{uid}-{doc_id}"
    sign_hash = generate_sign_hash(uid, doc_id, serial, hash_seed)

    try:
        conn.execute(
            """
            UPDATE signature_requests
            SET status='signed',
                serial_number=?,
                sign_hash=?,
                sign_type=?,
                signed_at=?,
                signature_owner_id=?,
                stamp_owner_id=?,
                signature_asset_id=?,
                stamp_asset_id=?,
                positions_json=?,
                include_qr=?
            WHERE id=?
            """,
            (
                serial,
                sign_hash,
                sign_type,
                datetime.utcnow().isoformat(),
                signature_owner_id,
                stamp_owner_id,
                signature_asset_id,
                stamp_asset_id,
                positions_json,
                1 if include_qr else 0,
                req['id']
            )
        )
    except Exception:
        try:
            conn.execute(
                """
                UPDATE signature_requests
                SET status='signed',
                    serial_number=?,
                    sign_hash=?,
                    sign_type=?,
                    signed_at=?,
                    signature_owner_id=?,
                    stamp_owner_id=?,
                    signature_asset_id=?,
                    stamp_asset_id=?
                WHERE id=?
                """,
                (
                    serial,
                    sign_hash,
                    sign_type,
                    datetime.utcnow().isoformat(),
                    signature_owner_id,
                    stamp_owner_id,
                    signature_asset_id,
                    stamp_asset_id,
                    req['id']
                )
            )
        except Exception:
            conn.execute(
                """
                UPDATE signature_requests
                SET status='signed',
                    serial_number=?,
                    sign_hash=?,
                    sign_type=?,
                    signed_at=?,
                    signature_owner_id=?,
                    stamp_owner_id=?
                WHERE id=?
                """,
                (
                    serial,
                    sign_hash,
                    sign_type,
                    datetime.utcnow().isoformat(),
                    signature_owner_id,
                    stamp_owner_id,
                    req['id']
                )
            )

    if doc['status'] != 'approved':
        conn.execute(
            "UPDATE documents SET status='approved', approved_by=COALESCE(approved_by, ?), approved_at=COALESCE(approved_at, datetime('now')) WHERE id=?",
            (uid, doc_id)
        )

    _persist_document_fingerprint(conn, doc_id)

    conn.commit()
    signed_doc = conn.execute("SELECT * FROM documents WHERE id=? LIMIT 1", (doc_id,)).fetchone()
    saved_to_storage = _upload_doc_to_storage(dict(signed_doc), conn) if signed_doc else False
    conn.close()

    user = get_user_by_id(uid)
    send_sign_used(
        user['email'],
        user['name'],
        doc['title'],
        doc['archive_number'] or '—',
        serial,
        session.get('user_name', user['name']),
        sign_type=sign_type,
    )
    log_action(uid, "SIGNED_PDF", f"SIG:{serial}", doc_id, request.remote_addr, request.user_agent.string)

    return jsonify({
        "success": True,
        "serial": serial,
        "sign_hash": sign_hash,
        "include_qr": include_qr,
        "approved": True,
        "saved_to_storage": saved_to_storage,
        "warning": None if saved_to_storage else "تم التوقيع والاعتماد لكن تعذر رفع الملف إلى الأرشيف"
    })

@app.route("/api/signature-request/<int:req_id>/sign", methods=["POST"])
@login_required
def sign_document(req_id):
    data      = request.json or {}
    sign_type = data.get("sign_type","signature")
    uid = session['user_id']

    conn = get_db()
    req  = conn.execute("SELECT * FROM signature_requests WHERE id=? AND requested_from=? AND status='pending'",
                        (req_id, uid)).fetchone()
    if not req:
        conn.close()
        return jsonify({"success": False, "error": "الطلب غير موجود"})

    doc    = conn.execute("SELECT * FROM documents WHERE id=?", (req['document_id'],)).fetchone()
    serial = _build_signature_serial(conn, req['document_id'])
    hash_seed = f"nopass-{uid}-{req['document_id']}"
    shash  = generate_sign_hash(uid, req['document_id'], serial, hash_seed)

    default_positions = _normalize_sign_positions([], include_qr=True)
    try:
        conn.execute(
            "UPDATE signature_requests SET status='signed', serial_number=?, sign_hash=?, sign_type=?, signed_at=?, include_qr=1, positions_json=? WHERE id=?",
            (serial, shash, sign_type, datetime.utcnow().isoformat(), json.dumps(default_positions, ensure_ascii=False), req_id)
        )
    except Exception:
        conn.execute(
            "UPDATE signature_requests SET status='signed', serial_number=?, sign_hash=?, sign_type=?, signed_at=? WHERE id=?",
            (serial, shash, sign_type, datetime.utcnow().isoformat(), req_id)
        )

    if doc and doc['status'] != 'approved':
        conn.execute(
            "UPDATE documents SET status='approved', approved_by=COALESCE(approved_by, ?), approved_at=COALESCE(approved_at, datetime('now')) WHERE id=?",
            (uid, req['document_id'])
        )

    _persist_document_fingerprint(conn, req['document_id'])

    conn.commit()
    signed_doc = conn.execute("SELECT * FROM documents WHERE id=? LIMIT 1", (req['document_id'],)).fetchone()
    saved_to_storage = _upload_doc_to_storage(dict(signed_doc), conn) if signed_doc else False

    # إشعار صاحب التوقيع
    user = get_user_by_id(uid)
    send_sign_used(user['email'], user['name'], doc['title'],
                   doc['archive_number'] or '—', serial, session['user_name'])

    log_action(uid, "SIGNED", f"SIG:{serial}", req['document_id'],
               request.remote_addr, request.user_agent.string)
    conn.close()
    result = {"success": True, "serial": serial, "approved": True, "saved_to_storage": saved_to_storage}
    if not saved_to_storage:
        result["warning"] = "تم التوقيع والاعتماد لكن تعذر رفع الملف إلى الأرشيف"
    return jsonify(result)

if __name__ == "__main__":
    init_db()

    admin_email = DEFAULT_ADMIN_EMAIL if DEFAULT_ADMIN_EMAIL else "admin@najm.sa"
    legacy_admin_email = "admin@najm.sa"
    target_admin = get_user_by_email(admin_email)

    if not target_admin:
        legacy_admin = get_user_by_email(legacy_admin_email) if admin_email != legacy_admin_email else None
        if legacy_admin:
            conn = get_db()
            conn.execute("UPDATE users SET email=? WHERE id=?", (admin_email, legacy_admin['id']))
            conn.commit(); conn.close()
            print(f"[NAJM] OK - تم تحديث إيميل الأدمن إلى: {admin_email}")
        else:
            create_user("المدير العام", admin_email, "+966500000000", "مدير عام", "admin", None)
            conn = get_db()
            conn.execute("UPDATE users SET password_hash=? WHERE email=?",
                         (hash_password("Admin@1234"), admin_email))
            conn.commit(); conn.close()
            print("=" * 45)
            print(f"[NAJM] OK - حساب المدير العام: {admin_email} / Admin@1234")
            print("[NAJM] WARNING - غيّر الباسورد فور الدخول!")
            print("=" * 45)

    try:
        from apscheduler.schedulers.background import BackgroundScheduler
        scheduler = BackgroundScheduler()
        scheduler.add_job(rotate_admin_passwords, 'interval', weeks=1)
        scheduler.start()
        print("[NAJM] Scheduler running")
    except ImportError:
        print("[NAJM] WARNING - pip install apscheduler")

    app.run(debug=True, host='0.0.0.0', port=5000)
